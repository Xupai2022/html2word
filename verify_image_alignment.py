#!/usr/bin/env python
"""验证Word文档中图片的对齐方式"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def verify_image_alignment(docx_file):
    """检查Word文档中所有图片的对齐方式"""
    print(f"正在检查文件: {docx_file}\n")

    doc = Document(docx_file)

    image_count = 0
    centered_count = 0
    left_count = 0
    right_count = 0
    none_count = 0

    # 查找所有包含图片的段落
    for i, para in enumerate(doc.paragraphs):
        if para._element.xpath('.//pic:pic'):  # 检查是否包含图片
            image_count += 1

            # 获取段落文本(用于识别上下文)
            prev_text = ""
            if i > 0:
                prev_text = doc.paragraphs[i-1].text[:50]

            # 检查对齐方式
            alignment = para.alignment
            if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                centered_count += 1
                print(f"✓ 图片 #{image_count}: 居中对齐")
                if "Security Rating" in prev_text or "security rating" in prev_text.lower():
                    print(f"  → 这是Security Rating下方的图片! ✓")
            elif alignment == WD_ALIGN_PARAGRAPH.LEFT:
                left_count += 1
                print(f"✗ 图片 #{image_count}: 左对齐")
                if "Security Rating" in prev_text or "security rating" in prev_text.lower():
                    print(f"  → 这是Security Rating下方的图片! ✗ (应该居中)")
            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                right_count += 1
                print(f"  图片 #{image_count}: 右对齐")
            else:
                none_count += 1
                print(f"? 图片 #{image_count}: 无对齐设置 (默认左对齐)")
                if "Security Rating" in prev_text or "security rating" in prev_text.lower():
                    print(f"  → 这是Security Rating下方的图片! ✗ (应该居中)")

            if prev_text:
                print(f"  上文: {prev_text}...")
            print()

    # 统计信息
    print("=" * 60)
    print(f"总计图片数量: {image_count}")
    print(f"  - 居中对齐: {centered_count}")
    print(f"  - 左对齐: {left_count}")
    print(f"  - 右对齐: {right_count}")
    print(f"  - 无对齐: {none_count}")
    print("=" * 60)

if __name__ == "__main__":
    verify_image_alignment("test_image_center_output.docx")
