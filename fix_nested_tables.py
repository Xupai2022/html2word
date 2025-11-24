"""
Fix nested tables by extracting them to top level.
"""
from docx import Document
from docx.oxml.ns import qn
import sys

def extract_nested_tables(input_file, output_file):
    """Extract nested tables to top level."""
    doc = Document(input_file)

    # Create new document
    new_doc = Document()

    # Process each table
    for table in doc.tables:
        rows = len(table.rows)
        cols = len(table.columns)

        # If it's a wrapper table (1x1) with nested content
        if rows == 1 and cols == 1:
            cell = table.rows[0].cells[0]

            # Check for paragraphs to copy
            for para in cell.paragraphs:
                if para.text.strip():
                    new_para = new_doc.add_paragraph()
                    # Copy text and basic formatting
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        if run.bold:
                            new_run.bold = True
                        if run.italic:
                            new_run.italic = True
                        if run.underline:
                            new_run.underline = True

            # Check for nested tables
            nested_tables = cell.tables
            if nested_tables:
                # Extract nested tables to top level
                for nested_table in nested_tables:
                    # Create new table at top level
                    n_rows = len(nested_table.rows)
                    n_cols = len(nested_table.columns)
                    new_table = new_doc.add_table(rows=n_rows, cols=n_cols)

                    # Copy table style
                    try:
                        new_table.style = nested_table.style
                    except:
                        new_table.style = 'Table Grid'

                    # Copy cell contents
                    for r_idx in range(n_rows):
                        for c_idx in range(n_cols):
                            old_cell = nested_table.rows[r_idx].cells[c_idx]
                            new_cell = new_table.rows[r_idx].cells[c_idx]

                            # Copy text
                            new_cell.text = old_cell.text

                            # Copy column width if it's set
                            try:
                                if nested_table.columns[c_idx].width:
                                    new_table.columns[c_idx].width = nested_table.columns[c_idx].width
                            except:
                                pass

                    print(f"Extracted {n_rows}x{n_cols} table to top level")
        else:
            # Direct table (not wrapped), copy as-is
            n_rows = len(table.rows)
            n_cols = len(table.columns)
            new_table = new_doc.add_table(rows=n_rows, cols=n_cols)

            # Copy table style
            try:
                new_table.style = table.style
            except:
                new_table.style = 'Table Grid'

            # Copy contents
            for r_idx in range(n_rows):
                for c_idx in range(n_cols):
                    old_cell = table.rows[r_idx].cells[c_idx]
                    new_cell = new_table.rows[r_idx].cells[c_idx]
                    new_cell.text = old_cell.text

                    # Copy column width
                    try:
                        if table.columns[c_idx].width:
                            new_table.columns[c_idx].width = table.columns[c_idx].width
                    except:
                        pass

            print(f"Copied {n_rows}x{n_cols} table")

    # Save new document
    new_doc.save(output_file)
    print(f"\nFixed document saved to: {output_file}")

    # Verify the result
    verify_doc = Document(output_file)
    print(f"\nVerification - Total tables: {len(verify_doc.tables)}")
    for i, table in enumerate(verify_doc.tables[:10]):
        print(f"  Table {i+1}: {len(table.rows)}x{len(table.columns)}")

if __name__ == "__main__":
    extract_nested_tables('final_fixed.docx', 'final_fixed_no_nesting.docx')

    # Also process test file
    extract_nested_tables('test_merged_result.docx', 'test_merged_no_nesting.docx')