#!/usr/bin/env python3
"""
Verification script to check improvements in security_quarterly_report.html conversion.
"""

from docx import Document
import sys

def verify_document(docx_path):
    """Verify that the document has expected styling improvements."""

    print(f"\n{'='*60}")
    print(f"Verifying: {docx_path}")
    print(f"{'='*60}\n")

    try:
        doc = Document(docx_path)

        # Count elements
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)

        print(f"[OK] Document loaded successfully")
        print(f"  - Total paragraphs: {total_paragraphs}")
        print(f"  - Total tables: {total_tables}")

        # Check for colored backgrounds (from gradients)
        paragraphs_with_bg = 0
        for para in doc.paragraphs:
            # Check paragraph shading
            pPr = para._element.get_or_add_pPr()
            shd = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            if shd is not None and shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'):
                paragraphs_with_bg += 1

        print(f"\n[OK] Background colors detected:")
        print(f"  - Paragraphs with background: {paragraphs_with_bg}")

        if paragraphs_with_bg > 0:
            print(f"  [OK] Linear gradient degradation working!")

        # Check tables
        cells_with_color = 0
        if total_tables > 0:
            print(f"\n[OK] Tables found: {total_tables}")

            # Check for colored cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        tcPr = cell._element.get_or_add_tcPr()
                        shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                        if shd is not None and shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'):
                            cells_with_color += 1

            print(f"  - Table cells with background: {cells_with_color}")

            if cells_with_color > 0:
                print(f"  [OK] Table cell coloring working!")

        # Check for text formatting
        runs_with_color = 0
        runs_with_bold = 0
        total_runs = 0

        for para in doc.paragraphs:
            for run in para.runs:
                total_runs += 1
                if run.font.color and run.font.color.rgb:
                    runs_with_color += 1
                if run.font.bold:
                    runs_with_bold += 1

        print(f"\n[OK] Text formatting:")
        print(f"  - Total text runs: {total_runs}")
        print(f"  - Runs with color: {runs_with_color}")
        print(f"  - Runs with bold: {runs_with_bold}")

        # Overall assessment
        print(f"\n{'='*60}")
        print("ASSESSMENT:")
        print(f"{'='*60}")

        issues = []

        if paragraphs_with_bg == 0:
            issues.append("[WARN] No paragraph backgrounds detected (gradients may not be converted)")
        else:
            print(f"[OK] Gradient backgrounds converted: {paragraphs_with_bg} instances")

        if total_tables == 0:
            issues.append("[WARN] No tables found (expected some tables)")
        else:
            print(f"[OK] Tables present: {total_tables}")

        if cells_with_color == 0 and total_tables > 0:
            issues.append("[WARN] Tables have no colored cells")
        elif total_tables > 0:
            print(f"[OK] Table cells with colors: {cells_with_color}")

        if runs_with_color == 0:
            issues.append("[WARN] No colored text found")
        else:
            print(f"[OK] Colored text runs: {runs_with_color}")

        if issues:
            print(f"\nPotential issues found:")
            for issue in issues:
                print(f"  {issue}")
        else:
            print(f"\n[SUCCESS] All checks passed! Document has good styling.")

        print(f"\n{'='*60}\n")

        return len(issues) == 0

    except Exception as e:
        print(f"[ERROR] Error: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    # Verify the improved security report
    print("\n" + "="*60)
    print("HTML2WORD IMPROVEMENT VERIFICATION")
    print("="*60)

    success = verify_document('test_security_fixed.docx')

    # Also verify baseline
    print("\n\nBaseline comparison:")
    verify_document('test_complex_baseline.docx')

    sys.exit(0 if success else 1)
