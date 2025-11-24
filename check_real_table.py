"""
Check real table widths in Word document.
"""
from docx import Document
from docx.oxml.ns import qn

def check_all_tables(docx_path):
    """Check all tables in the Word document."""
    doc = Document(docx_path)

    print(f"Document: {docx_path}")
    print(f"Total tables: {len(doc.tables)}")
    print("=" * 80)

    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}:")
        print(f"  Rows: {len(table.rows)}, Columns: {len(table.columns)}")

        # Check if this could be the target table (3 columns)
        if len(table.columns) == 3:
            print("  >>> This is a 3-column table! <<<")

            # Check table element properties
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is not None:
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is not None:
                    w_val = tblW.get(qn('w:w'))
                    w_type = tblW.get(qn('w:type'))
                    print(f"  Table width: {w_val} ({w_type})")
                    if w_type == 'dxa':
                        width_pt = int(w_val) / 20
                        width_inches = width_pt / 72
                        print(f"  Table width: {width_pt:.1f}pt ({width_inches:.2f} inches)")

            # Check column widths
            print("  Column widths:")
            total_width = 0
            for j, col in enumerate(table.columns):
                # Try to get width
                try:
                    if col.width:
                        width_pt = col.width.inches * 72
                        total_width += width_pt
                        print(f"    Column {j+1}: {width_pt:.1f}pt ({col.width.inches:.2f} inches)")
                    else:
                        print(f"    Column {j+1}: Auto")
                except:
                    print(f"    Column {j+1}: Error getting width")

            if total_width > 0:
                print(f"  Total column width: {total_width:.1f}pt ({total_width/72:.2f} inches)")
                if total_width > 468:
                    print(f"  !!! WARNING: Table exceeds page width (468pt / 6.5 inches) !!!")

            # Show content preview
            print("  Content preview:")
            for row_idx in range(min(3, len(table.rows))):
                row = table.rows[row_idx]
                row_text = []
                for cell in row.cells:
                    text = cell.text.strip()[:30]
                    row_text.append(text)
                print(f"    Row {row_idx+1}: {row_text}")

        elif len(table.columns) == 1:
            # Single column table - show first 50 chars
            if len(table.rows) > 0:
                text = table.rows[0].cells[0].text[:50]
                print(f"  Single cell content: {text}...")

# Check the generated documents
print("Checking test_merged.docx:")
check_all_tables('test_merged.docx')

print("\n" + "=" * 80)
print("\nChecking oversear_monthly_report_part1_test.docx:")
check_all_tables('oversear_monthly_report_part1_test.docx')