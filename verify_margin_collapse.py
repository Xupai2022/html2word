#!/usr/bin/env python
"""
Verify margin collapse implementation with simple test case.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent / 'src'))

from html2word.converter import HTML2WordConverter
from docx import Document

def verify():
    input_file = Path(__file__).parent / 'test_margin_collapse_simple.html'
    output_file = Path(__file__).parent / 'output_margin_test.docx'

    print("="*70)
    print("Margin Collapse Verification Test")
    print("="*70)

    # Convert
    print("\nConverting test HTML...")
    converter = HTML2WordConverter()
    converter.convert(str(input_file), str(output_file))
    print("[OK] Converted successfully")

    # Analyze
    print("\nAnalyzing spacing...")
    doc = Document(str(output_file))

    print("\n{:<5} {:<12} {:<12} {}".format("Para", "Before(pt)", "After(pt)", "Text"))
    print("-"*70)

    for i, para in enumerate(doc.paragraphs[:30]):
        fmt = para.paragraph_format
        before = fmt.space_before.pt if fmt.space_before else 0
        after = fmt.space_after.pt if fmt.space_after else 0
        text = para.text[:50] if para.text else "[empty]"

        if before > 0 or after > 0 or text.strip():
            print("{:<5} {:<12.1f} {:<12.1f} {}".format(i, before, after, text))

    print("\n"+"="*70)
    print("Analysis:")
    print("="*70)

    # Count patterns
    content_paras = [p for p in doc.paragraphs if p.text.strip()]
    spacer_paras = [p for p in doc.paragraphs if not p.text.strip()]

    content_with_after = sum(1 for p in content_paras
                             if p.paragraph_format.space_after
                             and p.paragraph_format.space_after.pt > 0)

    content_with_before = sum(1 for p in content_paras
                              if p.paragraph_format.space_before
                              and p.paragraph_format.space_before.pt > 0)

    spacers_with_before = sum(1 for p in spacer_paras
                              if p.paragraph_format.space_before
                              and p.paragraph_format.space_before.pt > 0)

    print(f"Content paragraphs with space_after:  {content_with_after}")
    print(f"Content paragraphs with space_before: {content_with_before}")
    print(f"Spacer paragraphs (after tables):     {spacers_with_before}")

    print("\nExpected behavior:")
    print("- Content paragraphs should use space_after ONLY (margin-bottom)")
    print("- Content paragraphs should NOT use space_before (margin-top)")
    print("- Empty spacers after tables use space_before (for table margin)")

    if content_with_before == 0:
        print("\n[PASS] No content paragraphs use space_before - margin collapse OK!")
    else:
        print(f"\n[FAIL] {content_with_before} content paragraphs use space_before")

    print(f"\n[OK] Check {output_file} to verify spacing visually")
    print("="*70)

if __name__ == '__main__':
    verify()
