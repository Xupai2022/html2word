#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Check image alignment in Word document"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def check_alignment(docx_file):
    """Check alignment of all images"""
    doc = Document(docx_file)

    image_count = 0

    # Find all paragraphs with images
    for i, para in enumerate(doc.paragraphs):
        if para._element.xpath('.//pic:pic'):
            image_count += 1

            # Get previous paragraph text
            prev_text = ""
            if i > 0:
                prev_text = doc.paragraphs[i-1].text[:80]

            # Check alignment
            alignment = para.alignment
            align_str = "NONE"
            if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                align_str = "CENTER"
            elif alignment == WD_ALIGN_PARAGRAPH.LEFT:
                align_str = "LEFT"
            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                align_str = "RIGHT"

            print(f"Image #{image_count}: alignment={align_str}")
            if prev_text:
                print(f"  Previous text: {prev_text}")
            print()

    print(f"\nTotal images: {image_count}")

if __name__ == "__main__":
    check_alignment("test_image_center_output.docx")
