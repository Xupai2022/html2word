"""
Compare regression test results.
"""
from docx import Document

def compare_documents(doc1_path, doc2_path):
    """Compare two documents for regression testing."""
    doc1 = Document(doc1_path)
    doc2 = Document(doc2_path)

    print(f"\nComparing:")
    print(f"  Expected: {doc1_path}")
    print(f"  Actual:   {doc2_path}")
    print("=" * 80)

    # Compare basic metrics
    metrics = [
        ("Paragraphs", len(doc1.paragraphs), len(doc2.paragraphs)),
        ("Tables", len(doc1.tables), len(doc2.tables)),
    ]

    # Count table cells with background
    def count_colored_cells(doc):
        count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                    if shd is not None and shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'):
                        count += 1
        return count

    colored_cells_expected = count_colored_cells(doc1)
    colored_cells_actual = count_colored_cells(doc2)
    metrics.append(("Colored cells", colored_cells_expected, colored_cells_actual))

    # Count colored text
    def count_colored_text(doc):
        count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    count += 1
        return count

    colored_text_expected = count_colored_text(doc1)
    colored_text_actual = count_colored_text(doc2)
    metrics.append(("Colored text runs", colored_text_expected, colored_text_actual))

    # Count bold text
    def count_bold_text(doc):
        count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.bold:
                    count += 1
        return count

    bold_expected = count_bold_text(doc1)
    bold_actual = count_bold_text(doc2)
    metrics.append(("Bold text runs", bold_expected, bold_actual))

    all_match = True
    for name, expected, actual in metrics:
        match = "OK" if expected == actual else "CHANGED"
        if match == "CHANGED":
            all_match = False
        print(f"{name:20s}: {expected:4d} -> {actual:4d}  [{match}]")

    print("=" * 80)
    if all_match:
        print("RESULT: All metrics match - NO REGRESSION")
    else:
        print("RESULT: Metrics changed - VERIFY manually if changes are expected")

    return all_match

if __name__ == "__main__":
    # Compare complex_tables_test.html output
    compare_documents("test_complex_baseline.docx", "REGRESSION_TEST.docx")
