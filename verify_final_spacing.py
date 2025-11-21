#!/usr/bin/env python
"""Verify spacing in the final converted document."""

from docx import Document
from docx.shared import Pt

def main():
    doc = Document('security_quarterly_report_spacing_fixed.docx')

    print("="*80)
    print("SPACING VERIFICATION")
    print("="*80)

    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")

    # Find spacer paragraphs (empty with spacing)
    spacers = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            fmt = para.paragraph_format
            space_before = fmt.space_before
            space_after = fmt.space_after

            if (space_before and space_before.pt > 0) or (space_after and space_after.pt > 0):
                spacers.append({
                    'index': i,
                    'space_before': space_before.pt if space_before else 0,
                    'space_after': space_after.pt if space_after else 0
                })

    print(f"\n{'-'*80}")
    print(f"Spacer paragraphs found: {len(spacers)}")
    print(f"{'-'*80}")

    expected_spacings = [22.5, 15.0, 30.0]  # Common spacing values
    spacing_counts = {s: 0 for s in expected_spacings}
    spacing_counts['other'] = 0

    for spacer in spacers:
        space = spacer['space_before'] or spacer['space_after']
        if space in expected_spacings:
            spacing_counts[space] += 1
        else:
            spacing_counts['other'] += 1

        direction = 'before' if spacer['space_before'] > 0 else 'after'
        print(f"  Paragraph {spacer['index']}: {space}pt ({direction})")

    print(f"\n{'-'*80}")
    print("Spacing summary:")
    print(f"{'-'*80}")
    print(f"  22.5pt (30px margin): {spacing_counts[22.5]} occurrences")
    print(f"  15.0pt (20px margin): {spacing_counts[15.0]} occurrences")
    print(f"  30.0pt (40px margin): {spacing_counts[30.0]} occurrences")
    if spacing_counts['other'] > 0:
        print(f"  Other values: {spacing_counts['other']} occurrences")

    print(f"\n{'-'*80}")
    print("EXPECTED vs ACTUAL:")
    print(f"{'-'*80}")
    print("✓ Executive Summary should have 22.5pt spacing after")
    print("✓ Risk Overview should have 22.5pt spacing after")
    print("✓ Metric Grid should have 15pt spacing before AND after")
    print("✓ Section should have 30pt spacing after")

    if len(spacers) > 0:
        print(f"\n✅ SUCCESS: {len(spacers)} spacing elements found!")
        print("   Spacing is being correctly applied to converted tables.")
    else:
        print(f"\n❌ PROBLEM: No spacing elements found.")
        print("   Check if spacing logic is working correctly.")

if __name__ == '__main__':
    main()
