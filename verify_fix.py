"""
Verify that the text between sections is properly preserved
"""

from docx import Document
import re

# Check if the specific text is in the Word document
doc = Document('final_test_output.docx')

# Collect all text from the document
all_text = []
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        all_text.append(paragraph.text.strip())

# Also check text in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    all_text.append(paragraph.text.strip())

# Search for the specific text
target_text = "Based on the analysis of all logs using multiple engines"
found = False
context = []

for i, text in enumerate(all_text):
    if target_text in text or "Based on the analysis" in text:
        found = True
        # Get context (previous and next lines)
        if i > 0:
            context.append(f"Previous: {all_text[i-1][:100]}")
        context.append(f">>> FOUND: {text[:200]}")
        if i < len(all_text) - 1:
            context.append(f"Next: {all_text[i+1][:100]}")
        break

print("=" * 60)
print("VERIFICATION RESULTS")
print("=" * 60)

if found:
    print("[SUCCESS] The text 'Based on the analysis...' was found!")
    print("\nContext:")
    for line in context:
        print(f"  {line}")
else:
    print("[FAILURE] The text 'Based on the analysis...' was NOT found!")
    print("\nSearching for similar patterns...")

    # Search for partial matches
    for text in all_text:
        if "analysis" in text.lower() or "logs" in text.lower() or "engines" in text.lower():
            print(f"  Partial match: {text[:150]}")

print("\n" + "=" * 60)

# Also check for the section titles
print("\nChecking for section titles:")
titles = ["Endpoint-Side Log Trend", "Security Alert Trend"]
for title in titles:
    found_title = any(title in text for text in all_text)
    if found_title:
        print(f"  [FOUND] {title}")
    else:
        print(f"  [MISSING] {title}")

print("\nTotal paragraphs in document:", len(doc.paragraphs))
print("Total tables in document:", len(doc.tables))