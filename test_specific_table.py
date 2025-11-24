"""
Test script to convert only the specific table section.
"""
from bs4 import BeautifulSoup
import re
from src.html2word.converter import HTML2WordConverter

# Read full HTML
with open('oversear_monthly_report_part1.html', 'r', encoding='utf-8') as f:
    full_html = f.read()

# Parse with BeautifulSoup
soup = BeautifulSoup(full_html, 'html.parser')

# Find the section containing the target text
text_element = soup.find(string=re.compile('Security devices are the first line of defense'))
if not text_element:
    print("Target text not found")
    exit(1)

# Go up to find containing section
section = text_element.parent
while section and section.name != 'section':
    section = section.parent

if not section:
    print("Section not found")
    exit(1)

# Extract just this section
# Create a new HTML document with just this section
test_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        .el-table__header {{ }}
        .el-table__body {{ }}
        table {{ border-collapse: collapse; }}
        td, th {{ border: 1px solid #ddd; padding: 8px; }}
    </style>
</head>
<body>
    {section}
</body>
</html>
"""

# Save test HTML
with open('test_section.html', 'w', encoding='utf-8') as f:
    f.write(test_html)

print("Test HTML saved to test_section.html")

# Convert using the converter
converter = HTML2WordConverter()
converter.convert('test_section.html', 'test_section.docx')

print("Converted to test_section.docx")

# Check the result
from docx import Document

doc = Document('test_section.docx')
print(f"\nResult: {len(doc.tables)} tables in document")

for i, table in enumerate(doc.tables):
    print(f"\nTable {i+1}:")
    print(f"  Rows: {len(table.rows)}, Columns: {len(table.columns)}")

    # Show first row
    if len(table.rows) > 0:
        first_row = []
        for cell in table.rows[0].cells[:4]:
            text = cell.text.strip()[:20]
            first_row.append(text)
        print(f"  First row: {first_row}")