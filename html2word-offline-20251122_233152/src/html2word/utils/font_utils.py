"""
Font utilities for consistent font handling across different character types.

This module provides utilities to ensure consistent font application across
ASCII, East Asian, and other character types in Word documents.
"""

import logging
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

logger = logging.getLogger(__name__)


def apply_uniform_font(run, font_name: str):
    """
    Apply uniform font to all character types in a text run.

    This function sets all Word font attributes (ascii, hAnsi, eastAsia, cs)
    to ensure consistent font display across different character types
    (numbers, letters, Chinese characters, etc.).

    Args:
        run: python-docx Run object
        font_name: Font name to apply uniformly

    Example:
        >>> from docx import Document
        >>> doc = Document()
        >>> para = doc.add_paragraph()
        >>> run = para.add_run("2024年第三季度网络安全态势分析报告")
        >>> apply_uniform_font(run, "Microsoft YaHei")
    """
    try:
        # Get the run properties element
        rPr = run._element.get_or_add_rPr()

        # Create the rFonts element with all font types set to the same font
        rFonts_xml = f'''
            <w:rFonts {nsdecls("w")}
                w:ascii="{font_name}"
                w:hAnsi="{font_name}"
                w:eastAsia="{font_name}"
                w:cs="{font_name}"/>
        '''

        # Parse the XML
        rFonts = parse_xml(rFonts_xml)

        # Remove existing rFonts element if present
        existing_rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        if existing_rFonts is not None:
            rPr.remove(existing_rFonts)

        # Add the new rFonts element
        rPr.append(rFonts)

        # Also set the basic font name for compatibility
        run.font.name = font_name

        logger.debug(f"Applied uniform font '{font_name}' to all character types")

    except Exception as e:
        logger.error(f"Failed to apply uniform font '{font_name}': {e}")
        # Fallback to basic font setting
        run.font.name = font_name


def get_run_font_info(run):
    """
    Get detailed font information from a run for debugging.

    Args:
        run: python-docx Run object

    Returns:
        dict: Font information for different character types
    """
    try:
        rPr = run._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is None:
            return {"error": "No run properties found"}

        rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        if rFonts is None:
            return {"error": "No rFonts element found"}

        return {
            "ascii": rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii'),
            "hAnsi": rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi'),
            "eastAsia": rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia'),
            "cs": rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs')
        }
    except Exception as e:
        return {"error": str(e)}