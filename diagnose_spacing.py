#!/usr/bin/env python3
"""
Diagnostic script to analyze spacing in converted DOCX files.
Compares actual spacing values with expected values from HTML.
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt

# Add src to path
sys.path.insert(0, str(Path(__file__).parent / "src"))

from html2word.parser.html_parser import HTMLParser
from html2word.style.box_model import BoxModel


def analyze_html_spacing(html_file):
    """Analyze expected spacing from HTML file."""
    print(f"\n{'='*80}")
    print(f"ANALYZING HTML: {html_file}")
    print(f"{'='*80}\n")

    parser = HTMLParser()
    dom_tree = parser.parse(html_file)

    # Find key elements and their spacing
    spacing_info = []

    def traverse(node, depth=0):
        """Traverse DOM tree and collect spacing info."""
        if not hasattr(node, 'tag'):
            return

        # Get class name for identification
        class_name = node.get_attribute('class', '')
        node_id = node.get_attribute('id', '')

        # Calculate box model
        if node.computed_styles:
            box_model = BoxModel(node)

            # Collect spacing for elements with margin or padding
            if (box_model.margin.top > 0 or box_model.margin.bottom > 0 or
                box_model.padding.top > 0 or box_model.padding.bottom > 0):

                info = {
                    'tag': node.tag,
                    'class': class_name,
                    'id': node_id,
                    'margin_top': box_model.margin.top,
                    'margin_bottom': box_model.margin.bottom,
                    'padding_top': box_model.padding.top,
                    'padding_bottom': box_model.padding.bottom,
                    'has_background': box_model.border.has_border() or
                                     'background-color' in node.computed_styles,
                }
                spacing_info.append(info)

                # Print info
                indent = "  " * depth
                print(f"{indent}<{node.tag}", end="")
                if class_name:
                    print(f" class='{class_name}'", end="")
                if node_id:
                    print(f" id='{node_id}'", end="")
                print(">")

                if box_model.margin.top > 0:
                    print(f"{indent}  margin-top: {box_model.margin.top}pt")
                if box_model.margin.bottom > 0:
                    print(f"{indent}  margin-bottom: {box_model.margin.bottom}pt")
                if box_model.padding.top > 0:
                    print(f"{indent}  padding-top: {box_model.padding.top}pt")
                if box_model.padding.bottom > 0:
                    print(f"{indent}  padding-bottom: {box_model.padding.bottom}pt")
                if info['has_background']:
                    print(f"{indent}  (has background/border - will wrap in table)")
                print()

        # Traverse children
        for child in node.children:
            traverse(child, depth + 1)

    traverse(dom_tree.root)

    return spacing_info


def analyze_docx_spacing(docx_file):
    """Analyze actual spacing in DOCX file."""
    print(f"\n{'='*80}")
    print(f"ANALYZING DOCX: {docx_file}")
    print(f"{'='*80}\n")

    doc = Document(docx_file)

    spacing_found = []

    # Analyze paragraphs
    print("PARAGRAPH SPACING:")
    print("-" * 80)

    for i, para in enumerate(doc.paragraphs):
        space_before = para.paragraph_format.space_before
        space_after = para.paragraph_format.space_after

        # Convert to pt
        before_pt = space_before.pt if space_before else 0
        after_pt = space_after.pt if space_after else 0

        if before_pt > 0 or after_pt > 0:
            text_preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            if not text_preview.strip():
                text_preview = "[EMPTY PARAGRAPH - SPACER]"

            print(f"Para {i}: {text_preview}")
            if before_pt > 0:
                print(f"  space_before: {before_pt}pt")
            if after_pt > 0:
                print(f"  space_after: {after_pt}pt")
            print()

            spacing_found.append({
                'type': 'paragraph',
                'index': i,
                'text': text_preview,
                'before': before_pt,
                'after': after_pt
            })

    # Analyze tables
    print("\nTABLE SPACING:")
    print("-" * 80)

    for i, table in enumerate(doc.tables):
        print(f"Table {i}: {len(table.rows)}x{len(table.columns)}")

        # Check cell padding
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Check paragraphs inside cells
                for p_idx, para in enumerate(cell.paragraphs):
                    space_before = para.paragraph_format.space_before
                    space_after = para.paragraph_format.space_after

                    before_pt = space_before.pt if space_before else 0
                    after_pt = space_after.pt if space_after else 0

                    if before_pt > 0 or after_pt > 0:
                        text_preview = para.text[:30] + "..." if len(para.text) > 30 else para.text
                        print(f"  Cell [{row_idx},{col_idx}] Para {p_idx}: {text_preview}")
                        if before_pt > 0:
                            print(f"    space_before: {before_pt}pt")
                        if after_pt > 0:
                            print(f"    space_after: {after_pt}pt")
        print()

    return spacing_found


def compare_spacing(html_spacing, docx_spacing):
    """Compare HTML expected vs DOCX actual spacing."""
    print(f"\n{'='*80}")
    print("SPACING COMPARISON & ISSUES")
    print(f"{'='*80}\n")

    print("EXPECTED FROM HTML (key elements):")
    print("-" * 80)

    total_expected = {}
    for info in html_spacing:
        key = f"{info['tag']}"
        if info['class']:
            key += f".{info['class']}"
        if info['id']:
            key += f"#{info['id']}"

        # Expected spacing after this element
        expected_after = info['margin_bottom']

        # If it has background/border, it will be wrapped in table
        # So spacing will be applied before and after table
        if info['has_background']:
            key += " [TABLE WRAPPER]"

        if expected_after > 0:
            total_expected[key] = expected_after
            print(f"{key}: {expected_after}pt (margin-bottom)")
            if info['padding_bottom'] > 0:
                print(f"  (also has padding-bottom: {info['padding_bottom']}pt - should NOT be external spacing)")

    print("\n\nACTUAL SPACING IN DOCX (empty paragraphs = spacers):")
    print("-" * 80)

    spacers = [s for s in docx_spacing if s['text'] == "[EMPTY PARAGRAPH - SPACER]"]
    print(f"Found {len(spacers)} spacer paragraphs:\n")

    for spacer in spacers:
        total_spacing = spacer['before'] + spacer['after']
        print(f"Spacer {spacer['index']}: {total_spacing}pt total")
        if spacer['before'] > 0:
            print(f"  space_before: {spacer['before']}pt")
        if spacer['after'] > 0:
            print(f"  space_after: {spacer['after']}pt")
        print()

    # Look for potential issues
    print("\n\nPOTENTIAL ISSUES:")
    print("-" * 80)

    # Check for unusually large spacing
    large_spacings = [s for s in spacers if (s['before'] + s['after']) > 35]
    if large_spacings:
        print(f"\n⚠️  {len(large_spacings)} spacers with >35pt spacing (may indicate padding being added to margin):")
        for s in large_spacings:
            print(f"   Spacer {s['index']}: {s['before'] + s['after']}pt")

    # Check for consecutive spacers (may indicate double spacing)
    consecutive = []
    for i in range(len(spacers) - 1):
        if spacers[i+1]['index'] == spacers[i]['index'] + 1:
            consecutive.append((spacers[i], spacers[i+1]))

    if consecutive:
        print(f"\n⚠️  {len(consecutive)} pairs of consecutive spacers (may indicate double spacing):")
        for s1, s2 in consecutive:
            total = s1['before'] + s1['after'] + s2['before'] + s2['after']
            print(f"   Spacers {s1['index']}-{s2['index']}: {total}pt total")

    # Statistics
    print(f"\n\nSTATISTICS:")
    print("-" * 80)

    if spacers:
        all_spacing_values = [s['before'] + s['after'] for s in spacers]
        avg_spacing = sum(all_spacing_values) / len(all_spacing_values)
        max_spacing = max(all_spacing_values)
        min_spacing = min([v for v in all_spacing_values if v > 0])

        print(f"Average spacer: {avg_spacing:.1f}pt")
        print(f"Max spacer: {max_spacing:.1f}pt")
        print(f"Min spacer: {min_spacing:.1f}pt")
        print(f"Total spacers: {len(spacers)}")


def main():
    """Main diagnostic function."""
    html_file = "security_quarterly_report.html"
    docx_file = "security_quarterly_report_spacing_fixed.docx"

    # Check files exist
    if not Path(html_file).exists():
        print(f"❌ HTML file not found: {html_file}")
        return

    if not Path(docx_file).exists():
        print(f"❌ DOCX file not found: {docx_file}")
        print(f"   Please convert the HTML file first")
        return

    # Analyze HTML
    html_spacing = analyze_html_spacing(html_file)

    # Analyze DOCX
    docx_spacing = analyze_docx_spacing(docx_file)

    # Compare
    compare_spacing(html_spacing, docx_spacing)

    print("\n" + "="*80)
    print("DIAGNOSIS COMPLETE")
    print("="*80 + "\n")


if __name__ == "__main__":
    main()
