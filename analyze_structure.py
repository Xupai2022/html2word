"""
Analyze the structure of generated Word document to verify proper nesting.
"""
from docx import Document

def analyze_document_structure(docx_path):
    """Analyze and print document structure."""
    doc = Document(docx_path)

    print(f"Document Analysis: {docx_path}")
    print("=" * 80)
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")

    print("\n" + "=" * 80)
    print("DOCUMENT STRUCTURE (First 30 elements):")
    print("=" * 80)

    # Get body elements to show actual structure
    element_count = 0
    for i, element in enumerate(doc.element.body):
        if element_count >= 30:
            print("\n... (showing first 30 elements)")
            break

        tag = element.tag.split('}')[-1]

        if tag == 'p':
            # Paragraph
            para = doc.paragraphs[sum(1 for e in doc.element.body[:i] if e.tag.split('}')[-1] == 'p')]
            text = para.text.strip()
            if len(text) > 60:
                text = text[:60] + "..."
            if text:
                print(f"{element_count:3d}. PARAGRAPH: {text}")
                element_count += 1

        elif tag == 'tbl':
            # Table
            table_idx = sum(1 for e in doc.element.body[:i] if e.tag.split('}')[-1] == 'tbl')
            table = doc.tables[table_idx]
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0

            # Get first cell text
            first_text = ""
            if table.rows and table.rows[0].cells:
                first_text = table.rows[0].cells[0].text.strip()
                if len(first_text) > 40:
                    first_text = first_text[:40] + "..."

            print(f"{element_count:3d}. TABLE ({rows}x{cols}): {first_text}")
            element_count += 1

    print("\n" + "=" * 80)
    print("DETAILED TABLE ANALYSIS:")
    print("=" * 80)

    for idx, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0

        print(f"\nTable {idx+1}: {rows} rows x {cols} columns")

        # Check if this is a wrapper table (1x1 containing nested content)
        if rows == 1 and cols == 1:
            cell = table.rows[0].cells[0]
            nested_tables = [e for e in cell._element if e.tag.split('}')[-1] == 'tbl']
            nested_paras = [e for e in cell._element if e.tag.split('}')[-1] == 'p']

            if nested_tables or len(nested_paras) > 3:
                print(f"  WARNING: WRAPPER TABLE detected!")
                print(f"     Contains: {len(nested_paras)} paragraphs, {len(nested_tables)} nested tables")

                # Show what's inside
                if nested_paras:
                    first_para_text = ""
                    for p_elem in nested_paras[:3]:
                        para = None
                        for p in cell.paragraphs:
                            if p._element == p_elem:
                                para = p
                                break
                        if para and para.text.strip():
                            first_para_text = para.text.strip()[:50]
                            print(f"     First content: {first_para_text}...")
                            break
        else:
            # Show first cell
            if table.rows and table.rows[0].cells:
                first_text = table.rows[0].cells[0].text.strip()
                if len(first_text) > 50:
                    first_text = first_text[:50] + "..."
                print(f"  First cell: {first_text}")

    print("\n" + "=" * 80)
    print("STRUCTURE CHECK:")
    print("=" * 80)

    # Check if first few elements follow expected pattern
    body_elements = list(doc.element.body)
    first_10_types = [e.tag.split('}')[-1] for e in body_elements[:10]]

    print(f"\nFirst 10 element types: {first_10_types}")

    # Look for big wrapper tables (indicator of structure problem)
    wrapper_tables = []
    for idx, table in enumerate(doc.tables):
        if len(table.rows) == 1 and len(table.columns) == 1:
            cell = table.rows[0].cells[0]
            nested_count = len([e for e in cell._element if e.tag.split('}')[-1] in ('p', 'tbl')])
            if nested_count > 5:
                wrapper_tables.append((idx, nested_count))

    if wrapper_tables:
        print(f"\nWARNING: Found {len(wrapper_tables)} large wrapper table(s):")
        for idx, count in wrapper_tables:
            print(f"   Table {idx+1}: wrapping {count} elements")
        print("\n   This suggests the .container div is being wrapped despite white background.")
    else:
        print("\nGOOD: No large wrapper tables detected - structure looks good!")

if __name__ == "__main__":
    analyze_document_structure("STRUCTURE_FINAL.docx")
