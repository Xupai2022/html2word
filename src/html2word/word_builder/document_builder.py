"""
Document builder for Word documents.

Main coordinator for building Word documents from DOM trees.
"""

import logging
import math
import re
from typing import Optional
from docx import Document

from html2word.parser.dom_tree import DOMNode, DOMTree
from html2word.word_builder.paragraph_builder import ParagraphBuilder
from html2word.word_builder.table_builder import TableBuilder
from html2word.word_builder.image_builder import ImageBuilder
from html2word.word_builder.header_footer_builder import HeaderFooterBuilder

logger = logging.getLogger(__name__)


class DocumentBuilder:
    """Builds complete Word documents from DOM trees."""

    def __init__(self, base_path: Optional[str] = None, enable_header_footer: bool = True):
        """
        Initialize document builder.

        Args:
            base_path: Base path for resolving relative paths
            enable_header_footer: Whether to enable headers and footers (default: True)
        """
        self.base_path = base_path
        self.document = Document()
        self.paragraph_builder = ParagraphBuilder(self.document, document_builder=self)
        self.table_builder = TableBuilder(self.document)
        self.image_builder = ImageBuilder(self.document, base_path)
        self.header_footer_builder = HeaderFooterBuilder(self.document, base_path)
        self.enable_header_footer = enable_header_footer
        self.in_table_cell = False  # Track if we're processing content inside a table cell
        self.processed_nodes = set()  # Track nodes that have been processed (for el-table merging)
        # 性能优化：缓存递归检查结果
        self._svg_cache = {}  # node id -> bool (是否包含SVG)
        self._bg_image_cache = {}  # node id -> bool (是否有背景图片)
        self._el_table_pairs = {}  # header node id -> body node (el-table配对缓存)

    def build(self, tree: DOMTree) -> Document:
        """
        Build Word document from DOM tree.

        Args:
            tree: DOM tree

        Returns:
            python-docx Document object
        """
        logger.info("Building Word document")

        # Apply cover image FIRST (before any other content)
        if self.enable_header_footer:
            try:
                logger.info("Adding cover image to document")
                self.header_footer_builder.apply_cover_image()
            except Exception as e:
                logger.error(f"Failed to add cover image: {e}", exc_info=True)
                # Continue even if cover image fails

        # Get body content
        from html2word.parser.html_parser import HTMLParser
        parser = HTMLParser()
        body = parser.get_body_content(tree)

        root_node = body if body else tree.root

        # 性能优化：启用SVG并行预处理
        # 现在的实现已修复序列化不一致问题，可以正确命中缓存
        self._preprocess_svg_nodes(root_node)

        # 性能优化：预索引el-table配对
        self._preindex_el_tables(root_node)

        if body:
            # Process body children
            self._process_children(body)
        else:
            # No body found, process root
            self._process_children(tree.root)

        # Apply headers and footers if enabled
        if self.enable_header_footer:
            try:
                logger.info("Applying headers and footers to document")
                self.header_footer_builder.apply_headers_footers()
            except Exception as e:
                logger.error(f"Failed to apply headers/footers: {e}", exc_info=True)
                # Continue even if headers/footers fail

        # FIXED: Keep table cell content left-aligned by default for better readability
        # This is a global safety net for all tables in the document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():  # Only process non-empty paragraphs
                            current_align = para.paragraph_format.alignment
                            # FIXED: Keep left alignment (0) for table cells, don't convert to justify
                            # Left alignment is more appropriate for tabular data
                            if current_align is None:
                                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        logger.info("Document built successfully")
        return self.document

    def _process_children(self, node: DOMNode):
        """
        Process child nodes recursively.

        Args:
            node: Parent DOM node
        """
        for child in node.children:
            self._process_node(child)

    def _process_node(self, node: DOMNode):
        """
        Process a single DOM node.

        Args:
            node: DOM node
        """
        if node.is_text:
            # Skip standalone text nodes (they should be inside elements)
            return

        if not node.is_element:
            return

        # CRITICAL: Skip hidden/template elements (contains templates, not visible content)
        # This must be checked BEFORE any other processing
        if self._should_skip_hidden_element(node):
            # Add more debug info for why element is being skipped
            if node.tag == 'p':
                logger.debug(f"Skipping paragraph with text: {node.get_text_content()[:50] if node.get_text_content() else 'empty'}")
                if node.computed_styles:
                    logger.debug(f"  Display: {node.computed_styles.get('display', 'not set')}")
                    logger.debug(f"  Visibility: {node.computed_styles.get('visibility', 'not set')}")
            logger.debug(f"Skipping hidden/template element: {node.tag} with class='{node.attributes.get('class', '')}'")
            return

        # Debug: Log when we encounter SVG nodes
        if node.tag == 'svg':
            logger.debug(f"_process_node encountered SVG element")

        # Skip elements with position: absolute or position: fixed
        # UNLESS they contain important content (cover pages, titles, tables, etc.)
        position = node.computed_styles.get('position', '')
        if position in ('absolute', 'fixed'):
            # SVG elements with absolute positioning are important (ECharts charts)
            if node.tag == 'svg':
                logger.debug(f"Processing SVG with position: {position}")
                # Continue processing the SVG element normally
            # Check if element has background-image (important visual content)
            elif self._has_background_image_in_inline_styles(node):
                logger.info(f"Processing {node.tag} with position: {position} because it has background-image")
                # Continue processing normally (don't return)
            # Check if this is a cover page or contains important structural content
            elif self._contains_important_content(node):
                logger.debug(f"Processing {node.tag} with position: {position} because it contains important content")
                # Process children directly, ignoring the positioning wrapper
                self._process_children(node)
                return
            else:
                logger.debug(f"Skipping {node.tag} with position: {position}")
                return

        # Route to appropriate builder based on tag
        tag = node.tag

        # Headings
        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag[1])
            self.paragraph_builder.build_heading(node, level)

        # Paragraphs
        elif tag in ('p', 'blockquote', 'pre'):
            logger.debug(f"Building paragraph for <{tag}>, text: {node.get_text_content()[:50] if hasattr(node, 'get_text_content') else 'N/A'}")
            self.paragraph_builder.build_paragraph(node)

        # Divs - intelligent handling based on layout and styles
        elif tag == 'div':
            # CRITICAL: Skip Element UI table wrapper divs - process children directly
            # These are structural divs that should not be converted to tables
            div_classes = node.attributes.get('class', '')
            if isinstance(div_classes, list):
                div_classes = ' '.join(div_classes)

            # Skip Element UI hidden-columns div (contains column templates, not visible content)
            if 'hidden-columns' in div_classes:
                logger.info(f"!!! SKIPPING Element UI hidden-columns div: {div_classes}")
                print(f"!!! SKIPPING hidden-columns div")
                return
            elif any(cls in div_classes for cls in ['el-table', 'el-table__header-wrapper', 'el-table__body-wrapper',
                                                     'el-table__fixed-wrapper', 'el-table__fixed-header-wrapper',
                                                     'el-table__fixed-body-wrapper', 'el-table__append-wrapper',
                                                     'el-table__empty-wrapper']):
                logger.debug(f"Skipping Element UI wrapper div: {div_classes}, processing children directly")
                self._process_children(node)
            # IMPORTANT: Check background-image FIRST, before SVG check
            # If div has background-image, it should be converted as a whole (including any SVG children)
            elif self._has_background_image(node):
                logger.debug(f"Div has background-image, converting to screenshot")
                success = self._convert_background_image_element(node)
                if not success:
                    # Fallback: process children if conversion fails
                    logger.warning("Background-image conversion failed, falling back to children processing")
                    self._process_children(node)
            # Check if this is an inline icon + text pattern (like tips/hints)
            # This should be checked BEFORE the general SVG check
            elif self._is_icon_text_pattern(node):
                logger.debug(f"Div is icon+text pattern, treating as single paragraph")
                self.paragraph_builder.build_paragraph(node)
            # Check if this div contains SVG (ECharts charts) - recursively
            elif self._contains_svg(node):
                logger.debug(f"Div contains SVG (possibly nested), processing children directly")
                self._process_children(node)
            # FIXED: display:grid now works correctly after removing default styles
            elif self._should_convert_to_grid_table(node):
                # Grid/flex layout - convert to table with horizontal layout
                self._convert_grid_to_table_smart(node)
            # IMPORTANT: Check if should be paragraph BEFORE checking styled table
            # Simple text divs with styling (like headers) should be paragraphs, not tables
            elif self._should_treat_div_as_paragraph(node):
                # Only inline content - treat as paragraph (paragraphs can have background/border styles too)
                self.paragraph_builder.build_paragraph(node)
            elif self._should_wrap_in_styled_table(node):
                # Has background/border - wrap in table to preserve styling
                # But skip if it's a root layout container (like .container)
                if self._is_root_layout_container(node):
                    # Skip wrapping for root layout containers
                    self._process_children(node)
                else:
                    self._wrap_div_in_styled_table(node)
            else:
                # Plain container - process children directly
                # NO spacing needed: child elements handle their own margins via space_after
                # This implements CSS margin collapse behavior
                self._process_children(node)

        # Lists
        elif tag in ('ul', 'ol'):
            self._process_list(node)

        # List items (shouldn't appear standalone, but handle it)
        elif tag == 'li':
            self.paragraph_builder.build_paragraph(node)

        # Tables
        elif tag == 'table':
            # Check if this table was already processed (as part of el-table merging)
            if id(node) in self.processed_nodes:
                logger.debug(f"Skipping table (already processed as part of el-table merge)")
                return

            # Debug: Log table class (use attributes.get for consistency)
            table_classes = node.attributes.get('class', '')
            if isinstance(table_classes, list):
                table_classes_str = ' '.join(table_classes)
            else:
                table_classes_str = table_classes
            logger.debug(f"Processing table with classes: {table_classes_str}")

            # Check for Element UI el-table pattern (header and body split into separate tables)
            is_header = self._is_el_table_header(node)
            logger.debug(f"Is el-table__header: {is_header}")

            if is_header:
                # Try to find and merge with el-table__body
                body_table = self._find_el_table_body(node)
                logger.debug(f"Found el-table__body: {body_table is not None}")

                if body_table:
                    logger.info("Detected Element UI el-table pattern, merging header and body tables")
                    self._build_merged_el_table(node, body_table)
                    # Mark body table as processed to avoid duplicate processing
                    self.processed_nodes.add(id(body_table))
                    return
                else:
                    logger.warning("Found el-table__header but no corresponding el-table__body")

            # Normal table processing
            self.table_builder.build_table(node)

        # Images
        elif tag == 'img':
            self.image_builder.build_image(node)

        # SVG elements
        elif tag == 'svg':
            self._process_svg(node)

        # Line break
        elif tag == 'br':
            self.document.add_paragraph()

        # Horizontal rule
        elif tag == 'hr':
            self._add_horizontal_rule()

        # Section elements - check for grid/flex layout or border/background styling
        elif tag == 'section':
            # Log section processing for debugging
            logger.debug(f"Processing section with classes: {node.attributes.get('class', [])}")

            # Check if section uses grid/flex layout (same logic as div)
            if self._should_convert_to_grid_table(node):
                logger.debug(f"Converting section to grid table (classes={node.attributes.get('class', [])})")
                self._convert_grid_to_table_smart(node)
            # Check if section has visual styling that requires table wrapper
            elif self._should_wrap_in_styled_table(node):
                logger.debug(f"Wrapping section in styled table (classes={node.attributes.get('class', [])})")
                self._wrap_div_in_styled_table(node)
            else:
                # Plain container - process children directly
                # Child sections with borders will be handled by their own processing
                logger.debug(f"Processing section children directly (classes={node.attributes.get('class', [])})")
                self._process_children(node)

        # Other container elements - process children directly
        elif tag in ('body', 'html', 'main', 'article', 'header', 'footer', 'nav', 'aside'):
            # NO spacing needed: child elements handle their own margins via space_after
            # This implements CSS margin collapse behavior naturally
            self._process_children(node)

        # Inline elements - skip (should be handled by paragraph builder)
        elif node.is_inline:
            # Create a paragraph for standalone inline elements
            para = self.document.add_paragraph()
            text = node.get_text_content()
            if text.strip():
                run = para.add_run(text)

        # Unknown elements - process children
        else:
            logger.debug(f"Unknown element: {tag}, processing children")
            self._process_children(node)

    def _process_list(self, list_node: DOMNode):
        """
        Process list (ul/ol) and its items.

        Args:
            list_node: List DOM node
        """
        is_ordered = list_node.tag == 'ol'

        # Get list style
        list_style = list_node.computed_styles.get('list-style-type', 'disc' if not is_ordered else 'decimal')

        for child in list_node.children:
            if child.tag == 'li':
                self._process_list_item(child, is_ordered, list_style)
            elif child.tag in ('ul', 'ol'):
                # Nested list
                self._process_list(child)

    def _process_list_item(self, li_node: DOMNode, is_ordered: bool, list_style: str):
        """
        Process list item.

        Args:
            li_node: List item DOM node
            is_ordered: Whether it's an ordered list
            list_style: List style type
        """
        paragraph = self.paragraph_builder.build_paragraph(li_node)

        if paragraph:
            # Apply list formatting
            if is_ordered:
                paragraph.style = 'List Number'
            else:
                paragraph.style = 'List Bullet'

    def _add_horizontal_rule(self):
        """Add a horizontal rule."""
        # Create a paragraph with bottom border
        paragraph = self.document.add_paragraph()
        fmt = paragraph.paragraph_format
        from docx.shared import Pt
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(6)

        # Add bottom border (simplified)
        paragraph.add_run()

    def _should_treat_div_as_paragraph(self, node: DOMNode) -> bool:
        """
        Determine if a div should be treated as a paragraph or a container.

        Args:
            node: DOM node

        Returns:
            True if should treat as paragraph, False if should treat as container
        """
        # Check if div has any block-level children
        for child in node.children:
            if child.is_element and not child.is_inline:
                # Has block-level children - treat as container
                return False

        # Only has text and inline elements - treat as paragraph
        # But only if it has actual content
        text_content = node.get_text_content().strip()
        if text_content:
            return True

        # Empty or only whitespace - treat as container (will be skipped)
        return False

    def _contains_important_content(self, node: DOMNode) -> bool:
        """
        Check if a node contains important structural content that should be preserved
        even if the node itself has position: absolute/fixed.

        This is used to preserve cover pages, titles, and tables that use absolute
        positioning for layout purposes.

        Args:
            node: DOM node to check

        Returns:
            True if node contains important content
        """
        # Check if node has classes that indicate it's a cover page or important container
        class_attr = node.attributes.get('class', '')
        if isinstance(class_attr, list):
            node_classes = class_attr
        else:
            node_classes = class_attr.split() if class_attr else []
        important_classes = [
            'cover-info--wrapper',  # Cover page wrapper
            'cover-info--title',    # Cover title
            'cover-info-table',     # Cover table
            'export-report__cover', # Export report cover
        ]

        if any(cls in node_classes for cls in important_classes):
            return True

        # Recursively check if any descendant contains important tags
        def has_important_descendants(n: DOMNode) -> bool:
            # Important tags that indicate real content (including SVG for charts)
            important_tags = {'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table', 'ul', 'ol', 'pre', 'blockquote', 'svg'}

            if n.tag in important_tags:
                # SVG always counts as important
                if n.tag == 'svg':
                    return True
                # Check if it has actual text content (not just whitespace)
                text_content = self._get_text_content(n).strip()
                if text_content:
                    return True

            # Check children
            for child in n.children:
                if child.is_element and has_important_descendants(child):
                    return True

            return False

        return has_important_descendants(node)

    def _get_text_content(self, node: DOMNode) -> str:
        """
        Get all text content from a node and its descendants.

        Args:
            node: DOM node

        Returns:
            Concatenated text content
        """
        if node.is_text:
            return node.text

        text_parts = []
        for child in node.children:
            if child.is_text:
                text_parts.append(child.text)
            elif child.is_element:
                text_parts.append(self._get_text_content(child))

        return ''.join(text_parts)

    def _should_convert_to_grid_table(self, node: DOMNode) -> bool:
        """
        Check if div/section uses grid/flex layout and should be converted to table.

        Args:
            node: DOM node

        Returns:
            True if should convert to table
        """
        if not node.computed_styles:
            return False

        display = node.computed_styles.get('display', '')

        # Check for grid or flex with multiple children
        if 'grid' in display or 'flex' in display:
            # FIXED: Count ALL element children, not just block-level
            # In flex/grid layouts, img and svg are inline but should be treated as flex/grid items
            child_elements = [c for c in node.children if c.is_element]
            # Only convert if has 2+ children (grid/flex items)
            if len(child_elements) >= 2:
                logger.debug(f"Detected {display} layout with {len(child_elements)} children, converting to table")
                return True
            elif len(child_elements) == 1:
                logger.debug(f"Skipping {display} layout with only 1 child (tag={node.tag})")

        return False

    def _has_significant_border(self, node: DOMNode) -> bool:
        """
        Check if element has a significant visible border.

        Args:
            node: DOM node

        Returns:
            True if has visible border
        """
        if not node.computed_styles:
            return False

        styles = node.computed_styles
        border_count = 0

        # Check each side for non-zero, non-white border
        for side in ['top', 'right', 'bottom', 'left']:
            width_key = f'border-{side}-width'
            color_key = f'border-{side}-color'
            style_key = f'border-{side}-style'

            # Check if border exists (non-zero width and not 'none' style)
            if width_key in styles and styles.get(width_key, '0') not in ('0', '0px', 'none'):
                if style_key in styles and styles.get(style_key, 'none') != 'none':
                    # Check if border color is not white/transparent
                    border_color = styles.get(color_key, '').lower()
                    is_white_border = (
                        border_color in ('', 'transparent', '#fff', '#ffffff', 'white',
                                       'rgb(255,255,255)', 'rgb(255, 255, 255)') or
                        border_color.startswith('#fff') or
                        'rgb(255,255,255)' in border_color.replace(' ', '')
                    )

                    if not is_white_border:
                        border_count += 1

        # Consider it significant if 3+ sides have visible borders
        return border_count >= 3

    def _should_wrap_in_styled_table(self, node: DOMNode) -> bool:
        """
        Check if div has significant styling that requires table wrapper.

        Args:
            node: DOM node

        Returns:
            True if needs table wrapper for styling
        """
        # CRITICAL: If already inside a table cell, generally do NOT wrap in another table
        # EXCEPT for section elements with significant borders that need to be displayed
        if self.in_table_cell:
            # For section elements with clear borders, allow nested table wrapping
            if node.tag == 'section' and self._has_significant_border(node):
                logger.debug(f"Allowing nested table wrap for section with significant border (classes={node.attributes.get('class', [])})")
                # Continue to check for styling
            else:
                logger.debug(f"Skipping table wrap for {node.tag} (already in cell)")
                return False

        if not node.computed_styles:
            return False

        # Check for important visual styles FIRST
        styles = node.computed_styles

        # Check background property first (for gradients)
        has_non_white_bg = False
        bg_prop = styles.get('background', '')
        if bg_prop and bg_prop not in ('', 'none', 'transparent'):
            bg_lower = bg_prop.lower()
            # Gradient always counts as meaningful background
            if 'gradient' in bg_lower:
                has_non_white_bg = True

        # Check background-color
        bg_color = styles.get('background-color', '')
        if bg_color and bg_color not in ('', 'transparent', 'rgba(0, 0, 0, 0)', 'rgba(0,0,0,0)'):
            bg_lower = bg_color.lower()
            # Check if it's white/very light (likely just layout container)
            is_white = (
                bg_lower in ('#fff', '#ffffff', 'white', 'rgb(255,255,255)', 'rgb(255, 255, 255)') or
                bg_lower.startswith('#fff') or
                'rgb(255,255,255)' in bg_lower.replace(' ', '')
            )

            if not is_white:
                has_non_white_bg = True
                logger.debug(f"Has non-white background: {bg_color}")
            else:
                logger.debug(f"Ignoring white background: {bg_color}")

        # Borders - only count if it's a full border with non-white color
        has_border = False
        border_count = 0
        for side in ['top', 'right', 'bottom', 'left']:
            width_key = f'border-{side}-width'
            color_key = f'border-{side}-color'

            # Check if border exists (non-zero width)
            if width_key in styles and styles.get(width_key, '0') not in ('0', '0px', 'none'):
                # Check if border color is not white/transparent
                border_color = styles.get(color_key, '').lower()
                is_white_border = (
                    border_color in ('', 'transparent', '#fff', '#ffffff', 'white',
                                   'rgb(255,255,255)', 'rgb(255, 255, 255)') or
                    border_color.startswith('#fff') or
                    'rgb(255,255,255)' in border_color.replace(' ', '')
                )

                # Only count non-white borders
                if not is_white_border:
                    border_count += 1

        # Only consider it styled if it has 3+ sides bordered (like a box)
        if border_count >= 3:
            has_border = True

        # Box shadow
        has_shadow = 'box-shadow' in styles and styles['box-shadow'] not in ('none', '')

        # Check if has meaningful visual styling
        has_visual_styling = has_non_white_bg or has_border or has_shadow

        # If has visual styling, always wrap (even for simple content)
        if has_visual_styling:
            return True

        # No visual styling - don't wrap plain containers
        # Tables and other block elements can handle their own layout
        return False

    def _has_background_image_in_inline_styles(self, node: DOMNode) -> bool:
        """
        Check if element has background-image in inline_styles.

        This is used for position: absolute elements before computed_styles are available.

        Args:
            node: DOM node

        Returns:
            True if has valid background-image in inline_styles
        """
        if not hasattr(node, 'inline_styles') or not node.inline_styles:
            return False

        bg_image = node.inline_styles.get('background-image', '')
        if bg_image and bg_image not in ('', 'none', 'initial', 'inherit') and 'url(' in bg_image:
            return True

        return False

    def _should_skip_hidden_element(self, node: DOMNode) -> bool:
        """
        Check if an element should be skipped because it's hidden or is a template.

        This is a UNIVERSAL method that handles:
        - Element UI hidden-columns (column templates)
        - Elements with display:none
        - Elements with visibility:hidden
        - Template elements from various frameworks
        - Other framework-specific hidden elements

        Args:
            node: DOM node

        Returns:
            True if element should be skipped
        """
        if not node.is_element:
            return False

        # Get element classes
        classes = node.attributes.get('class', '')
        if isinstance(classes, list):
            classes = ' '.join(classes)
        elif not isinstance(classes, str):
            classes = ''

        # Skip Element UI hidden-columns (column definition templates)
        if 'hidden-columns' in classes:
            return True

        # Skip other common template/hidden classes from various frameworks
        # Vue.js templates
        if 'v-show-false' in classes or 'v-if-false' in classes:
            return True

        # React/Angular hidden elements
        if 'hidden' in classes or 'd-none' in classes or 'display-none' in classes:
            return True

        # Check computed styles for display:none or visibility:hidden
        if node.computed_styles:
            display = node.computed_styles.get('display', '')
            visibility = node.computed_styles.get('visibility', '')

            # Special handling: if element has actual text content, don't skip it
            # even if display:none (might be wrongly applied CSS)
            if display == 'none':
                text_content = node.get_text_content()
                if text_content and text_content.strip():
                    # Has meaningful text content, check if this is a false positive
                    # (e.g., p:empty rule incorrectly applied to non-empty paragraph)
                    logger.debug(f"Element has display:none but contains text: {text_content[:50]}")
                    # Don't skip if it's a paragraph with actual content
                    if node.tag == 'p':
                        logger.debug(f"Not skipping non-empty paragraph despite display:none")
                        return False
                return True

            if visibility == 'hidden':
                return True

        # Check for template tags (Vue, Angular, etc.)
        if node.tag in ('template', 'script', 'style', 'noscript'):
            return True

        # Check for aria-hidden="true" (accessibility hidden elements)
        aria_hidden = node.attributes.get('aria-hidden', '')
        if aria_hidden == 'true' or aria_hidden is True:
            return True

        return False

    def _has_background_image(self, node: DOMNode) -> bool:
        """
        Check if element has a background-image that needs special handling.
        使用缓存避免重复检查。

        Args:
            node: DOM node

        Returns:
            True if has valid background-image
        """
        node_id = id(node)
        if node_id in self._bg_image_cache:
            return self._bg_image_cache[node_id]

        result = False

        if node.computed_styles:
            bg_image = node.computed_styles.get('background-image', '')

            # Fallback: check inline_styles if not in computed_styles
            if not bg_image and hasattr(node, 'inline_styles') and node.inline_styles:
                bg_image = node.inline_styles.get('background-image', '')

            if bg_image and bg_image not in ('', 'none', 'initial', 'inherit'):
                # Check if it's a valid image URL (data URI or external URL)
                if 'url(' in bg_image:
                    result = True

        self._bg_image_cache[node_id] = result
        return result

    def _preprocess_svg_nodes(self, root: DOMNode):
        """
        Pre-process all SVG nodes and convert them to PNG in parallel.
        Performance optimization: Batch process SVGs to enable parallel Chrome rendering.
        """
        # 收集所有SVG节点
        svg_nodes = []

        def collect_svg(node: DOMNode):
            if node.is_element:
                if node.tag == 'svg':
                    svg_nodes.append(node)
                for child in node.children:
                    collect_svg(child)

        collect_svg(root)

        if not svg_nodes:
            return

        logger.info(f"Found {len(svg_nodes)} SVG elements, preparing batch conversion...")

        # 准备批量转换数据
        svg_list = []
        for svg_node in svg_nodes:
            try:
                # 1. 获取宽高等属性 (逻辑与 ImageBuilder.build_svg 保持一致)
                width_str = svg_node.get_attribute('width') or svg_node.computed_styles.get('width', '100')
                height_str = svg_node.get_attribute('height') or svg_node.computed_styles.get('height', '100')

                # 2. 使用 ImageBuilder 序列化 SVG (确保字符串完全一致)
                svg_content = self.image_builder.serialize_svg_node(svg_node, width_str, height_str)
                
                # CRITICAL: 将序列化好的内容存回节点，供后续 ImageBuilder 使用
                # 这保证了预处理生成的 Key 与构建时生成的 Key 100% 匹配
                svg_node._preprocessed_svg_content = svg_content

                # 3. 计算像素尺寸 (逻辑与 ImageBuilder._convert_svg_with_browser 保持一致)
                # Get font-size for em/rem unit calculations
                font_size_pt = 12.0  # Default
                font_size_str = svg_node.computed_styles.get('font-size', '')
                if font_size_str:
                    import re
                    fs_match = re.match(r'([\d.]+)(px|pt)?', str(font_size_str))
                    if fs_match:
                        fs_num = float(fs_match.group(1))
                        fs_unit = fs_match.group(2) or 'px'
                        if fs_unit == 'px':
                            font_size_pt = fs_num * 0.75
                        elif fs_unit == 'pt':
                            font_size_pt = fs_num

                # 先转为 pt (with font-size context for em units)
                width_pt = self.image_builder._parse_dimension(width_str, font_size_pt)
                height_pt = self.image_builder._parse_dimension(height_str, font_size_pt)

                # 再转为 px (96 DPI)
                width_px = int(width_pt * 96 / 72)
                height_px = int(height_pt * 96 / 72)

                svg_list.append((svg_content, width_px, height_px))
            except Exception as e:
                logger.debug(f"Failed to prepare SVG for batch conversion: {e}")

        if svg_list:
            # 批量并行转换
            # Windows上Chrome进程开销较大，并发数过高容易卡死，降低到 2 以保证稳定性
            from html2word.utils.browser_svg_converter import get_browser_converter
            converter = get_browser_converter()
            # 这里生成的缓存 Key = MD5(svg_content + width_px + height_px)
            # 后续 ImageBuilder 会使用完全相同的参数查询缓存，从而命中
            converter.convert_batch(svg_list, max_workers=2)

    def _preindex_el_tables(self, root: DOMNode):
        """
        预索引所有el-table的header-body配对。
        性能优化：一次遍历构建配对映射，避免后续重复搜索。
        改进：使用结构化搜索代替全局列表匹配，确保配对的准确性。

        Args:
            root: DOM树根节点
        """
        self._el_table_pairs = {}
        headers = []

        def has_class(node: DOMNode, class_name: str) -> bool:
            if not node.is_element:
                return False
            classes = node.attributes.get('class', '')
            if isinstance(classes, list):
                return class_name in classes
            elif isinstance(classes, str):
                return class_name in classes.split()
            return False

        # 1. 收集所有 header
        def collect_headers(node: DOMNode):
            if node.is_element:
                if node.tag == 'table' and has_class(node, 'el-table__header'):
                    headers.append(node)
                for child in node.children:
                    collect_headers(child)

        collect_headers(root)

        if not headers:
            return

        logger.info(f"Found {len(headers)} el-table headers, building index...")

        count = 0
        for header in headers:
            # 模仿 _find_el_table_body 的逻辑进行结构化搜索
            # Header -> Parent (Wrapper) -> Grandparent (Container) -> Body Wrapper -> Body
            
            if not header.parent or not header.parent.parent:
                continue
                
            parent = header.parent
            grandparent = parent.parent
            
            # 在 Grandparent 的子节点中寻找 Body Wrapper
            target_body = None
            
            # 优化：只在 parent 之后的兄弟节点中寻找，因为 Body 肯定在 Header 后面
            try:
                if parent in grandparent.children:
                    start_idx = grandparent.children.index(parent) + 1
                else:
                    continue
            except ValueError:
                continue
                
            for i in range(start_idx, len(grandparent.children)):
                sibling = grandparent.children[i]
                
                # Case 1: Sibling IS the body table (rare)
                if sibling.tag == 'table' and has_class(sibling, 'el-table__body'):
                    target_body = sibling
                    break
                
                # Case 2: Sibling IS the body wrapper (common)
                # 检查 sibling 的 children 中是否有 el-table__body
                if sibling.is_element:
                    found_in_sibling = None
                    for child in sibling.children:
                        if child.tag == 'table' and has_class(child, 'el-table__body'):
                            found_in_sibling = child
                            break
                    
                    if found_in_sibling:
                        target_body = found_in_sibling
                        break
            
            if target_body:
                self._el_table_pairs[id(header)] = target_body
                count += 1
        
        logger.info(f"Indexed {count} el-table pairs")

    def _parse_dimension_to_px(self, value: str) -> int:
        """解析尺寸值为像素"""
        import re
        if not value:
            return 100

        value = str(value).strip().lower()

        # 纯数字
        if value.isdigit():
            return int(value)

        # 带单位
        match = re.match(r'^([\d.]+)(px|pt|em|rem|%)?$', value)
        if match:
            num = float(match.group(1))
            unit = match.group(2) or 'px'

            if unit == 'px':
                return int(num)
            elif unit == 'pt':
                return int(num * 1.333)  # pt to px
            elif unit in ('em', 'rem'):
                return int(num * 16)  # assume 16px base
            elif unit == '%':
                return int(num * 6)  # rough estimate

        return 100  # default

    def _serialize_element_to_html(self, node: DOMNode) -> str:
        """
        Serialize a DOM element and its children to HTML string with inline styles.

        Args:
            node: DOM node to serialize

        Returns:
            HTML string
        """
        if node.is_text:
            return node.text or ""

        if not node.is_element:
            return ""

        # Build opening tag with attributes and inline styles
        tag = node.tag
        attrs = []

        # Add existing attributes (excluding style, we'll rebuild it)
        for key, value in node.attributes.items():
            if key != 'style':
                attrs.append(f'{key}="{value}"')

        # Build inline style from inline_styles (NOT computed_styles to avoid bloat)
        # Only use inline_styles which contain the original style attribute
        if hasattr(node, 'inline_styles') and node.inline_styles:
            style_parts = []
            for key, value in node.inline_styles.items():
                style_parts.append(f"{key}: {value}")
            if style_parts:
                attrs.append(f'style="{"; ".join(style_parts)}"')

        attrs_str = f" {' '.join(attrs)}" if attrs else ""

        # Serialize children recursively
        children_html = ""
        for child in node.children:
            children_html += self._serialize_element_to_html(child)

        # Return complete HTML
        if not children_html and tag in ('img', 'br', 'hr', 'input', 'meta', 'link'):
            # Self-closing tags
            return f"<{tag}{attrs_str} />"
        else:
            return f"<{tag}{attrs_str}>{children_html}</{tag}>"

    def _composite_background_with_text(self, node: DOMNode, width_pt: float, height_pt: float) -> bytes:
        """
        Composite background image with text overlays using PIL.

        This creates a perfect visual representation by:
        1. Decoding the base64 background image
        2. Drawing text overlays on top with correct positioning

        Args:
            node: DOM node with background-image
            width_pt: Target width in points
            height_pt: Target height in points

        Returns:
            PNG image data as bytes, or None if failed
        """
        try:
            import base64
            import re
            import io
            from PIL import Image, ImageDraw, ImageFont

            # Extract background image
            bg_image_css = node.computed_styles.get('background-image', '') or \
                          (node.inline_styles.get('background-image', '') if hasattr(node, 'inline_styles') else '')

            match = re.search(r'url\(["\']?data:image/([^;]+);base64,([^)"\'\s]+)', bg_image_css)
            if not match:
                logger.warning("Cannot extract base64 image from background-image")
                return None

            base64_data = match.group(2)
            image_data = base64.b64decode(base64_data)

            # Load background image
            bg_img = Image.open(io.BytesIO(image_data))
            original_width, original_height = bg_img.size
            logger.debug(f"Loaded background image: {original_width}x{original_height}")

            # Resize if needed to match target dimensions
            target_width_px = int(width_pt * 96 / 72)
            target_height_px = int(height_pt * 96 / 72)

            # Calculate scaling factors for text positioning
            # IMPORTANT: HTML text positions are NOT relative to the image's actual pixel size (1600x1144)
            # They are relative to how the image is displayed in the browser
            # Based on the text positions (max left ~552px, max top ~328px),
            # we estimate the container display size is around 800x600

            # Fine-tuned display size for optimal text scaling
            # 660x460 gives scale factors around 0.945x0.970 (slightly larger)
            assumed_display_width = 665  # Fine-tuned for larger text
            assumed_display_height = 510  # Fine-tuned for larger text

            # Calculate proper scale based on assumed display size, not actual image size
            scale_x = target_width_px / assumed_display_width if assumed_display_width > 0 else 1.0
            scale_y = target_height_px / assumed_display_height if assumed_display_height > 0 else 1.0

            # Position compensation - shift left and up
            offset_x = -30  # Shift left by 15 pixels
            offset_y = +10  # Shift up by 15 pixels

            logger.debug(f"Text position scaling: assumed display {assumed_display_width}x{assumed_display_height}")
            logger.debug(f"Text scale factors: x={scale_x:.3f}, y={scale_y:.3f}")

            if bg_img.size != (target_width_px, target_height_px):
                bg_img = bg_img.resize((target_width_px, target_height_px), Image.Resampling.LANCZOS)
                logger.debug(f"Resized image from {original_width}x{original_height} to {target_width_px}x{target_height_px}")
                img_scale_x = target_width_px / original_width
                img_scale_y = target_height_px / original_height
                logger.debug(f"Image scale factors: x={img_scale_x:.3f}, y={img_scale_y:.3f}")

            # Convert to RGBA to support transparency
            if bg_img.mode != 'RGBA':
                bg_img = bg_img.convert('RGBA')

            # Draw text overlays
            draw = ImageDraw.Draw(bg_img)

            # Process text children
            for child in node.children:
                if child.is_element:
                    text_content = child.get_text_content() if hasattr(child, 'get_text_content') else ''
                    if not text_content.strip():
                        continue

                    # Get position from inline_styles first (priority), then computed_styles
                    top_str = '0'
                    left_str = '0'

                    if hasattr(child, 'inline_styles') and child.inline_styles:
                        top_str = child.inline_styles.get('top', top_str)
                        left_str = child.inline_styles.get('left', left_str)

                    # Fallback to computed_styles if not in inline_styles
                    if top_str == '0' and hasattr(child, 'computed_styles'):
                        top_str = child.computed_styles.get('top', '0')
                    if left_str == '0' and hasattr(child, 'computed_styles'):
                        left_str = child.computed_styles.get('left', '0')

                    # Parse positions - if they are in px, convert directly
                    # If they are percentages or other units, convert to pixels
                    if 'px' in top_str:
                        # Direct pixel value from HTML
                        top_px_original = float(top_str.replace('px', ''))
                        # Apply scaling to match the resized image and add offset
                        top_px = int(top_px_original * scale_y) + offset_y
                    else:
                        # Convert other units using UnitConverter
                        from html2word.utils.units import UnitConverter
                        top_pt = UnitConverter.to_pt(top_str)
                        # Convert to pixels and apply scaling and offset
                        top_px = int(top_pt * 96 / 72 * scale_y) + offset_y

                    if 'px' in left_str:
                        # Direct pixel value from HTML
                        left_px_original = float(left_str.replace('px', ''))
                        # Apply scaling to match the resized image and add offset
                        left_px = int(left_px_original * scale_x) + offset_x
                    else:
                        # Convert other units using UnitConverter
                        from html2word.utils.units import UnitConverter
                        left_pt = UnitConverter.to_pt(left_str)
                        # Convert to pixels and apply scaling and offset
                        left_px = int(left_pt * 96 / 72 * scale_x) + offset_x

                    # Get font size from styles
                    font_size_str = '14px'  # Default
                    if hasattr(child, 'inline_styles') and child.inline_styles:
                        font_size_val = child.inline_styles.get('font-size', font_size_str)
                        font_size_str = str(font_size_val) if font_size_val else font_size_str
                    if font_size_str == '14px' and hasattr(child, 'computed_styles'):
                        font_size_val = child.computed_styles.get('font-size', '14px')
                        font_size_str = str(font_size_val) if font_size_val else '14px'

                    # Parse font size
                    if isinstance(font_size_str, str) and 'px' in font_size_str:
                        base_font_size = int(float(font_size_str.replace('px', '')))
                    elif isinstance(font_size_str, str) and 'pt' in font_size_str:
                        base_font_size = int(float(font_size_str.replace('pt', '')) * 96 / 72)
                    else:
                        # Try to convert directly to int if it's a number
                        try:
                            base_font_size = int(float(font_size_str))
                        except (ValueError, TypeError):
                            base_font_size = 14

                    # Scale font size proportionally with image
                    # Since we're scaling down significantly, we need to be careful with font size
                    avg_scale = (scale_x + scale_y) / 2
                    # Don't apply extra multiplier - just scale proportionally
                    scaled_font_size = max(10, int(base_font_size * avg_scale))

                    # Try to load a font (fallback to default if not available)
                    try:
                        font = ImageFont.truetype("arial.ttf", scaled_font_size)
                        logger.debug(f"Using Arial font with size {scaled_font_size} (base: {base_font_size})")
                    except:
                        try:
                            # Try alternative font paths
                            import os
                            if os.name == 'nt':  # Windows
                                font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", scaled_font_size)
                            else:  # Linux/Mac
                                font = ImageFont.truetype("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", scaled_font_size)
                        except:
                            font = ImageFont.load_default()
                            logger.debug("Using default font")

                    # Log positioning details for debugging
                    logger.debug(f"Text positioning for '{text_content.strip()[:20]}':")
                    logger.debug(f"  CSS position: top={top_str}, left={left_str}")
                    logger.debug(f"  Pixels (adjusted): top_px={top_px}, left_px={left_px}")
                    logger.debug(f"  Font size: base={base_font_size}px, scaled={scaled_font_size}px")

                    # Get text color (from inline_styles first, then computed_styles)
                    color = '#000000'
                    if hasattr(child, 'inline_styles') and child.inline_styles:
                        color = child.inline_styles.get('color', color)
                    if color == '#000000' and hasattr(child, 'computed_styles'):
                        color = child.computed_styles.get('color', '#000000')

                    # Simple color parsing (hex)
                    if color.startswith('#'):
                        color = color[1:]
                        if len(color) == 3:
                            color = ''.join([c*2 for c in color])
                        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
                        color_tuple = (r, g, b)
                    else:
                        color_tuple = (0, 0, 0)  # default black

                    # Draw text with anchor='lt' (left-top) for more predictable positioning
                    draw.text((left_px, top_px), text_content.strip(), fill=color_tuple, font=font, anchor='lt')
                    logger.debug(f"Drew text at ({left_px}, {top_px}): {text_content.strip()[:30]}")

            # Save to PNG bytes
            output = io.BytesIO()
            bg_img.save(output, format='PNG')
            png_data = output.getvalue()

            logger.info(f"Composited background+text image: {len(png_data)} bytes")
            return png_data

        except Exception as e:
            logger.error(f"Failed to composite image: {e}", exc_info=True)
            return None

    def _convert_background_image_element(self, node: DOMNode) -> bool:
        """
        Convert an element with background-image to an embedded image.

        Extracts base64 image data from background-image CSS property and
        inserts it directly into Word document. If the element has child text
        elements, they will be composited onto the image using PIL.

        Args:
            node: DOM node with background-image

        Returns:
            True if conversion succeeded, False otherwise
        """
        try:
            import base64
            import re
            import io
            from docx.shared import Inches
            from html2word.utils.units import UnitConverter

            # Get background-image from computed_styles or inline_styles
            bg_image = node.computed_styles.get('background-image', '')
            if not bg_image and hasattr(node, 'inline_styles'):
                bg_image = node.inline_styles.get('background-image', '')

            if not bg_image or 'url(' not in bg_image:
                logger.warning("No valid background-image found")
                return False

            # Extract data URI from url(data:image/png;base64,...)
            match = re.search(r'url\(["\']?data:image/([^;]+);base64,([^)"\'\s]+)', bg_image)
            if not match:
                logger.warning("Background-image is not a base64 data URI")
                return False

            image_format = match.group(1)  # e.g., 'png', 'jpeg'
            base64_data = match.group(2)

            logger.debug(f"Extracting {image_format} image from background-image ({len(base64_data)} base64 chars)")

            # Decode base64 to binary
            image_data = base64.b64decode(base64_data)
            logger.debug(f"Decoded image size: {len(image_data)} bytes")

            # Get element dimensions, fallback to inline_styles if needed
            width_str = node.computed_styles.get('width') or \
                       (node.inline_styles.get('width') if hasattr(node, 'inline_styles') else None) or \
                       '800px'
            height_str = node.computed_styles.get('height') or \
                        (node.inline_styles.get('height') if hasattr(node, 'inline_styles') else None) or \
                        '400px'

            width_pt = UnitConverter.to_pt(width_str)
            height_pt = UnitConverter.to_pt(height_str)

            # Get actual image dimensions for aspect ratio calculation
            from PIL import Image
            image_stream_temp = io.BytesIO(image_data)
            img = Image.open(image_stream_temp)
            actual_width_px, actual_height_px = img.size
            aspect_ratio = actual_width_px / actual_height_px

            # If width is 0 or auto, calculate from height and aspect ratio
            if width_pt == 0 or width_str.lower() in ('auto', 'none'):
                if height_pt > 0:
                    # Use height and maintain aspect ratio
                    width_pt = height_pt * aspect_ratio
                else:
                    # Both auto, use actual dimensions (converted to pt)
                    width_pt = actual_width_px * 72 / 96
                    height_pt = actual_height_px * 72 / 96

                logger.debug(f"Calculated dimensions from aspect ratio: {width_pt:.1f}x{height_pt:.1f}pt (ratio={aspect_ratio:.2f})")

            # Check if we're in a table cell context
            in_cell = getattr(self, 'in_table_cell', False)

            # Limit to page width or cell width
            if in_cell:
                # In table cells, use reduced width to account for cell padding
                max_width_pt = 432  # 468 - 36 = 432pt (accounting for padding)
                logger.debug("Background image in table cell, using reduced max width")
            else:
                # Standard page width (6.5 inches = 468pt for standard Word page with margins)
                max_width_pt = 468

            if width_pt > max_width_pt:
                # Scale down to fit width while maintaining aspect ratio
                scale_factor = max_width_pt / width_pt
                width_pt = max_width_pt
                height_pt = height_pt * scale_factor
                logger.debug(f"Scaled down to fit {'cell' if in_cell else 'page'}: {width_pt:.1f}x{height_pt:.1f}pt")

            # Try to composite with text overlays first
            logger.debug("Attempting PIL composite method...")
            composited_image_data = self._composite_background_with_text(node, width_pt, height_pt)

            if composited_image_data:
                # Use composited image
                image_data = composited_image_data
            else:
                logger.warning("PIL composite failed, using original background image only")

            # Insert image into document
            image_stream = io.BytesIO(image_data)
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            picture = run.add_picture(
                image_stream,
                width=Inches(width_pt / 72)
            )

            # Apply text alignment
            text_align = node.computed_styles.get('text-align', '')
            if text_align == 'center':
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif text_align == 'right':
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            logger.info(f"Successfully inserted background-image ({width_pt:.1f}x{height_pt:.1f}pt, {len(image_data)} bytes)")

            return True

        except Exception as e:
            logger.error(f"Failed to convert background-image element: {e}", exc_info=True)
            return False

    def _convert_grid_to_table(self, grid_node: DOMNode):
        """
        Convert a grid/flex container to a Word table.

        Args:
            grid_node: DOM node with grid/flex layout
        """
        try:
            # Get grid properties
            display = grid_node.computed_styles.get('display', '')
            grid_template_cols = grid_node.computed_styles.get('grid-template-columns', '')

            # Collect direct children (grid items)
            child_elements = [child for child in grid_node.children if child.is_element]

            if not child_elements:
                # No children, render as paragraph
                self.paragraph_builder.build_paragraph(grid_node)
                return

            # Determine number of columns
            num_cols = self._determine_grid_columns(grid_template_cols, len(child_elements))

            # Calculate number of rows
            num_rows = (len(child_elements) + num_cols - 1) // num_cols

            logger.debug(f"Creating table from grid: {num_rows} rows × {num_cols} columns")

            # Create table
            table = self.document.add_table(rows=num_rows, cols=num_cols)

            # Fill table with grid items
            cell_index = 0
            for row_idx in range(num_rows):
                for col_idx in range(num_cols):
                    if cell_index < len(child_elements):
                        child = child_elements[cell_index]
                        cell = table.rows[row_idx].cells[col_idx]

                        # Add child content to cell
                        # Get text content and add as paragraph
                        text_content = child.get_text_content()
                        if text_content.strip():
                            paragraph = cell.paragraphs[0]
                            paragraph.text = text_content.strip()

                            # Apply styles from child to cell
                            if child.computed_styles:
                                from html2word.word_builder.style_mapper import StyleMapper
                                mapper = StyleMapper()
                                mapper.apply_table_cell_style(cell, child.computed_styles, child.box_model)

                        cell_index += 1

            # Apply grid container styles to table
            if grid_node.computed_styles:
                # Apply background color if present
                bg_color = grid_node.computed_styles.get('background-color')
                if bg_color:
                    # Apply to all cells
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.paragraphs:
                                from html2word.word_builder.style_mapper import StyleMapper
                                mapper = StyleMapper()
                                # Create temporary style dict for background
                                temp_styles = {'background-color': bg_color}
                                mapper.apply_table_cell_style(cell, temp_styles)

        except Exception as e:
            logger.warning(f"Error converting grid to table: {e}, falling back to paragraph")
            # Fallback: render as paragraph
            self.paragraph_builder.build_paragraph(grid_node)

    def _convert_grid_to_table_smart(self, grid_node: DOMNode):
        """
        Intelligently convert grid/flex layout to table, preserving nested structure.

        Args:
            grid_node: DOM node with grid/flex layout
        """
        try:
            grid_template_cols = grid_node.computed_styles.get('grid-template-columns', '')
            flex_wrap = grid_node.computed_styles.get('flex-wrap', 'nowrap')
            display = grid_node.computed_styles.get('display', '')

            # Check if children are chart-panel-wrap sections that should be independent tables
            chart_panel_children = []
            other_children = []

            for child in grid_node.children:
                if child.is_element:
                    if child.tag == 'section' and 'chart-panel-wrap' in child.attributes.get('class', []):
                        chart_panel_children.append(child)
                        logger.info(f"Found chart-panel-wrap in grid, will process as separate table")
                    else:
                        other_children.append(child)

            # If we have chart-panel-wrap sections, process them separately
            if chart_panel_children:
                logger.info(f"Processing {len(chart_panel_children)} chart-panel-wrap sections as independent tables")

                # Process any other content first
                if other_children:
                    # Process other children directly
                    for child in other_children:
                        self._process_node(child)

                # Process each chart panel as a separate table
                for i, chart_panel in enumerate(chart_panel_children):
                    logger.debug(f"Processing chart-panel-wrap as independent table")
                    self._process_node(chart_panel)
                    # Add a paragraph separator between chart panels to ensure they are separate tables
                    if i < len(chart_panel_children) - 1:
                        self.document.add_paragraph()
                        logger.debug("Added paragraph separator between chart panels")

                return

            # FIXED: Count ALL element children (including inline like img, svg)
            child_elements = [c for c in grid_node.children if c.is_element]

            if not child_elements:
                self._process_children(grid_node)
                return

            # Determine columns based on layout type
            if 'flex' in display:
                # For flex layouts, determine columns based on flex-wrap
                num_cols = self._determine_flex_columns(flex_wrap, len(child_elements))
            else:
                # For grid layouts, use grid-template-columns
                num_cols = self._determine_grid_columns(grid_template_cols, len(child_elements))

            num_rows = (len(child_elements) + num_cols - 1) // num_cols

            logger.debug(f"Creating smart grid table: {num_rows}×{num_cols}")

            # Create table (no spacing before - previous element's space_after handles that)
            table = self.document.add_table(rows=num_rows, cols=num_cols)

            # Apply gap as table cell spacing (read from HTML, not hardcoded)
            self._apply_grid_gap_to_table(table, grid_node.computed_styles)

            # Fill cells with grid items
            cell_index = 0
            for row_idx in range(num_rows):
                for col_idx in range(num_cols):
                    if cell_index < len(child_elements):
                        child = child_elements[cell_index]
                        cell = table.rows[row_idx].cells[col_idx]

                        # Apply cell styles from child (including background, border, padding)
                        if child.computed_styles:
                            from html2word.word_builder.style_mapper import StyleMapper
                            from html2word.style.box_model import BoxModel
                            mapper = StyleMapper()
                            # Calculate box model from child element to get padding, borders
                            box_model = BoxModel(child)
                            logger.debug(f"Applying styles for grid child: {child.tag} with classes={child.attributes.get('class', [])}")
                            logger.debug(f"  Border: top={box_model.border.top.color}, right={box_model.border.right.color}")
                            mapper.apply_table_cell_style(cell, child.computed_styles, box_model)

                        # Recursively process child content in cell context
                        # Save current document temporarily
                        original_doc = self.document

                        # Create a temporary processor for cell content
                        # We need to clear the first paragraph if it's empty
                        if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
                            # Remove empty paragraph
                            p = cell.paragraphs[0]._element
                            p.getparent().remove(p)

                        # Process child's children in the cell
                        for grandchild in child.children:
                            self._process_node_in_cell(grandchild, cell, original_doc)

                        # Restore document
                        self.document = original_doc
                        cell_index += 1

            # Apply spacing after table (from grid container's margin-bottom)
            self._apply_spacing_after_table(grid_node)

        except Exception as e:
            logger.warning(f"Error in smart grid conversion: {e}, falling back")
            self._process_children(grid_node)

    def _apply_grid_gap_to_table(self, table, computed_styles: dict):
        """
        Apply grid/flex gap as Word table cell spacing.
        Reads gap values from HTML/CSS (not hardcoded).

        Args:
            table: Word table object
            computed_styles: Computed CSS styles containing gap properties
        """
        if not computed_styles:
            return

        # Read gap values from HTML/CSS (row-gap and column-gap were expanded from gap by CSS parser)
        gap_value = computed_styles.get('gap') or computed_styles.get('row-gap') or computed_styles.get('column-gap')

        if not gap_value:
            return

        try:
            from html2word.utils.units import UnitConverter

            # Convert gap to points
            gap_pt = UnitConverter.to_pt(gap_value)

            if gap_pt <= 0:
                return

            # Convert points to twips (Word's internal unit: 1 pt = 20 twips)
            gap_twips = int(gap_pt * 20)

            logger.debug(f"Applying grid gap: {gap_value} → {gap_pt}pt → {gap_twips} twips")

            # Apply cell spacing to table using OpenXML
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            tbl = table._element
            tblPr = tbl.tblPr

            # Ensure tblPr exists
            if tblPr is None:
                tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
                tbl.insert(0, tblPr)

            # Add or update tblCellSpacing element
            # Remove existing cellSpacing if present
            for child in list(tblPr):
                if child.tag.endswith('tblCellSpacing'):
                    tblPr.remove(child)

            # Add new cellSpacing (applies to all sides)
            cellSpacing = parse_xml(
                f'<w:tblCellSpacing {nsdecls("w")} w:w="{gap_twips}" w:type="dxa"/>'
            )
            tblPr.append(cellSpacing)

        except Exception as e:
            logger.warning(f"Failed to apply grid gap: {e}")

    def _process_node_in_cell(self, node: DOMNode, cell, original_doc):
        """Process a node within a table cell context."""
        # Create a wrapper that redirects document operations to cell
        class CellDocumentWrapper:
            def __init__(self, cell_obj, original):
                self._cell = cell_obj
                self._original = original

            def add_paragraph(self, *args, **kwargs):
                return self._cell.add_paragraph(*args, **kwargs)

            def add_table(self, *args, **kwargs):
                # FIXED: Add table to cell, not to original document
                return self._cell.add_table(*args, **kwargs)

        old_doc = self.paragraph_builder.document
        old_table_doc = self.table_builder.document
        old_image_doc = self.image_builder.document  # Save image builder document
        old_self_doc = self.document
        old_in_cell = self.in_table_cell  # Save previous state
        old_margin_bottom = self.paragraph_builder.last_margin_bottom  # Save margin state

        wrapper = CellDocumentWrapper(cell, original_doc)
        self.paragraph_builder.document = wrapper
        self.table_builder.document = wrapper
        self.image_builder.document = wrapper  # CRITICAL: Also update image_builder to use cell context
        self.document = wrapper  # CRITICAL: Also replace self.document for grid conversion
        self.in_table_cell = True  # Mark that we're inside a cell
        self.paragraph_builder.last_margin_bottom = 0.0  # Reset margin collapse in new context

        try:
            self._process_node(node)
        finally:
            self.paragraph_builder.document = old_doc
            self.table_builder.document = old_table_doc
            self.image_builder.document = old_image_doc  # Restore image builder document
            self.document = old_self_doc
            self.in_table_cell = old_in_cell  # Restore previous state
            self.paragraph_builder.last_margin_bottom = old_margin_bottom  # Restore margin state

    def _wrap_div_in_styled_table(self, node: DOMNode):
        """
        Wrap a styled div/section in a table to preserve background/borders.

        Special handling for sections containing chart-panel-header elements:
        - These get their own row/cell with background color

        Args:
            node: DOM node with styling (div, section, etc.)
        """
        try:
            logger.debug(f"Wrapping styled {node.tag} in table (classes={node.attributes.get('class', [])})")

            # Check if this section contains chart-panel-wrap sections that should be separate tables
            chart_panel_children = []
            header_children = []
            other_children = []

            for child in node.children:
                # Check if this child is a chart-panel-wrap section (should be its own table)
                if child.tag == 'section':
                    classes = child.attributes.get('class', [])
                    logger.debug(f"Checking child section with classes: {classes}")
                    if 'chart-panel-wrap' in classes:
                        chart_panel_children.append(child)
                        logger.info(f"Found chart-panel-wrap section that needs separate table (classes: {classes})")
                        continue

                # Check if this child is or contains a chart-panel-header
                is_header = False
                if child.attributes.get('class'):
                    classes = child.attributes.get('class', [])
                    if 'chart-panel-header' in classes:
                        is_header = True
                        logger.debug(f"Found chart-panel-header child: {child.tag}")

                if is_header:
                    header_children.append(child)
                else:
                    other_children.append(child)

            # If we have chart-panel-wrap sections, process them separately without wrapping in table
            if chart_panel_children:
                logger.info(f"Processing {len(chart_panel_children)} chart-panel-wrap sections as separate tables")

                # Process any other content first (if any) - but skip creating a wrapper table
                # since the chart panels should be independent
                for child in other_children:
                    self._process_node(child)

                for child in header_children:
                    self._process_node(child)

                # Now process each chart panel as a separate table
                for chart_panel in chart_panel_children:
                    logger.debug(f"Processing chart-panel-wrap as separate table")
                    self._process_node(chart_panel)

                return  # Exit early since we've handled everything

            # Determine number of rows needed
            num_rows = 1
            if header_children:
                # Each header gets its own row, plus one for other content
                num_rows = len(header_children) + (1 if other_children else 0)

            # Create table with appropriate number of rows
            table = self.document.add_table(rows=num_rows, cols=1)

            # Add table borders for better visual structure
            # This is especially important for section tables
            self._add_table_borders(table, node)

            current_row = 0

            # Process header children first (each in its own row/cell)
            for header_child in header_children:
                cell = table.rows[current_row].cells[0]

                # Apply header styles to the cell (especially background color)
                if header_child.computed_styles:
                    from html2word.word_builder.style_mapper import StyleMapper
                    from html2word.style.box_model import BoxModel, BoxEdge
                    mapper = StyleMapper()

                    # Use header's styles for the cell
                    cell_styles = header_child.computed_styles.copy()
                    box_model = header_child.layout_info.get('box_model') if hasattr(header_child, 'layout_info') else None

                    # If no box_model or padding, create padding based on height to simulate vertical spacing
                    if 'height' in cell_styles and cell_styles['height']:
                        from html2word.utils.units import UnitConverter
                        height_pt = UnitConverter.to_pt(str(cell_styles['height']))

                        # If box_model doesn't exist or has no padding, create one with appropriate padding
                        if not box_model or (box_model.padding.top == 0 and box_model.padding.bottom == 0):
                            # Calculate padding: for a 28px height with 12px font, add some padding
                            # This simulates the visual effect of the height property
                            font_size_pt = 12  # Default font size
                            if 'font-size' in cell_styles:
                                font_size_pt = UnitConverter.to_pt(str(cell_styles['font-size']))

                            # Add padding to achieve the desired height effect
                            # Typically, (height - font_size) / 2 for top and bottom
                            vertical_padding = max(3, (height_pt - font_size_pt) / 2)  # At least 3pt padding

                            if not box_model:
                                # Create a new box model with padding
                                box_model = BoxModel()
                                box_model.padding = BoxEdge(
                                    top=vertical_padding,
                                    right=6.0,  # Small horizontal padding
                                    bottom=vertical_padding,
                                    left=6.0
                                )
                            else:
                                # Update existing box model's padding
                                box_model.padding.top = max(box_model.padding.top, vertical_padding)
                                box_model.padding.bottom = max(box_model.padding.bottom, vertical_padding)

                            logger.debug(f"Added cell padding: top={vertical_padding:.1f}pt, bottom={vertical_padding:.1f}pt for height={height_pt}pt")

                    # Apply background to cell, not just paragraph
                    mapper.apply_table_cell_style(cell, cell_styles, box_model)

                    # Also set row height if specified
                    if 'height' in cell_styles:
                        from html2word.utils.units import UnitConverter
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls

                        height_pt = UnitConverter.to_pt(str(cell_styles['height']))
                        if height_pt > 0:
                            tr = table.rows[current_row]._element
                            trPr = tr.get_or_add_trPr()
                            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="exact"/>')
                            trPr.append(trHeight)
                            logger.debug(f"Set row height to {height_pt}pt for chart-panel-header")

                # Clear default empty paragraph if present
                if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
                    p = cell.paragraphs[0]._element
                    p.getparent().remove(p)

                # Process the header content
                original_doc = self.document
                self._process_node_in_cell(header_child, cell, original_doc)
                self.document = original_doc

                # CRITICAL FIX: Keep header row with next row (content) to prevent pagination split
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True

                current_row += 1

            # Process other children in remaining cell(s)
            if other_children:
                cell = table.rows[current_row].cells[0]

                # Apply section styles to cell (if not already handled by headers)
                if node.computed_styles and not header_children:
                    from html2word.word_builder.style_mapper import StyleMapper
                    mapper = StyleMapper()
                    box_model = node.layout_info.get('box_model') if hasattr(node, 'layout_info') else None
                    mapper.apply_table_cell_style(cell, node.computed_styles, box_model)

                # Clear default empty paragraph if present
                if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
                    p = cell.paragraphs[0]._element
                    p.getparent().remove(p)

                # Process other children
                original_doc = self.document
                for child in other_children:
                    if child.tag:
                        logger.debug(f"Processing child {child.tag} (classes={child.attributes.get('class', [])}) in table cell")
                    self._process_node_in_cell(child, cell, original_doc)
                self.document = original_doc

            # Apply spacing after table (from div's margin-bottom)
            self._apply_spacing_after_table(node)

            # Special handling for chart-panel-wrap: add a clear separator
            # This ensures the table is truly independent and not merged with adjacent tables
            if 'chart-panel-wrap' in node.attributes.get('class', []):
                # Add a separator paragraph with a non-breaking space to ensure separation
                # Use minimal spacing (half line = ~6pt)
                from docx.shared import Pt
                separator = self.document.add_paragraph("\u00A0")  # Non-breaking space
                separator.paragraph_format.space_before = Pt(0)
                separator.paragraph_format.space_after = Pt(0)
                separator.paragraph_format.line_spacing = Pt(6)  # Half line spacing (~12pt font / 2)
                logger.info("Added minimal separator paragraph after chart-panel-wrap table to prevent merging")

        except Exception as e:
            logger.warning(f"Error wrapping div in table: {e}, falling back")
            self._process_children(node)

    def _add_table_borders(self, table, node=None):
        """
        Add borders to a table using the original CSS border style if available.

        Args:
            table: python-docx Table object
            node: Optional DOM node to extract border styles from
        """
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            # Get the table element
            tbl = table._element

            # Default border style (similar to #dee3ed - light gray, thin)
            border_size = "4"  # 0.5pt (4 = 0.5pt, 8 = 1pt)
            border_color = "DEE3ED"  # Light gray color from original CSS

            # Try to extract border style from node if provided
            if node and node.computed_styles:
                # Check for border styles
                for side in ['top', 'right', 'bottom', 'left']:
                    border_width = node.computed_styles.get(f'border-{side}-width', '')
                    border_color_css = node.computed_styles.get(f'border-{side}-color', '')

                    if border_width and '1px' in str(border_width):
                        border_size = "4"  # 1px = 0.5pt in Word

                    if border_color_css:
                        # Convert CSS color to Word color (remove # and use uppercase)
                        if border_color_css.startswith('#'):
                            border_color = border_color_css[1:].upper()
                        elif border_color_css.startswith('rgb'):
                            # Try to parse RGB color
                            import re
                            rgb_match = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', border_color_css)
                            if rgb_match:
                                r, g, b = rgb_match.groups()
                                border_color = f"{int(r):02X}{int(g):02X}{int(b):02X}"
                        break  # Use first found color

            # Create table borders element with extracted or default styles
            tblBorders = parse_xml(f'''
                <w:tblBorders {nsdecls("w")}>
                    <w:top w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>
                    <w:left w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>
                    <w:bottom w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>
                    <w:right w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>
                    <w:insideH w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>
                    <w:insideV w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>
                </w:tblBorders>
            ''')

            # Add borders to table properties
            tblPr = tbl.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
            if tblPr is not None:
                # Remove existing borders if present
                existing_borders = tblPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
                if existing_borders is not None:
                    tblPr.remove(existing_borders)
                # Add new borders
                tblPr.append(tblBorders)
            else:
                # Create table properties if not present
                tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
                tblPr.append(tblBorders)
                tbl.insert(0, tblPr)

            logger.debug("Added borders to wrapper table")

        except Exception as e:
            logger.warning(f"Failed to add table borders: {e}")

    def _apply_spacing_before_table(self, node: DOMNode):
        """
        DEPRECATED: No longer used. Previous element's space_after handles spacing.

        Kept for compatibility but does nothing.
        This implements CSS margin collapse: we don't add spacing before elements,
        only after (via space_after on paragraphs or spacer after tables).
        """
        pass  # Intentionally do nothing

    def _apply_spacing_after_table(self, node: DOMNode):
        """
        Apply spacing after a table based on the source element's margin-bottom.

        Since tables don't support paragraph_format.space_after, we add a spacer paragraph
        with space_before set to the table's margin-bottom. This implements the "only space_after"
        principle (or in this case, space_before on the spacer represents the table's bottom margin).

        Args:
            node: Source DOM node (div/grid container/table wrapper)
        """
        try:
            # Calculate box model to get margin
            from html2word.style.box_model import BoxModel
            box_model = BoxModel(node)

            # Special handling for chart-panel-wrap: ensure proper separation
            classes = node.attributes.get('class', [])
            is_chart_panel = 'chart-panel-wrap' in classes

            # Apply margin-bottom as space_before on a spacer paragraph
            # This represents the table's margin-bottom and will be the spacing
            # between this table and the next element
            spacing = box_model.margin.bottom

            # For chart panels, don't apply extra spacing here since we add a separator in _wrap_div_in_styled_table
            # Just skip the spacing to avoid double-spacing
            if is_chart_panel:
                logger.debug("Skipping spacing for chart-panel-wrap (separator already added)")
                return

            if spacing > 0:
                from docx.shared import Pt
                spacer = self.document.add_paragraph()
                spacer.paragraph_format.space_before = Pt(spacing)
                logger.debug(f"Applied {spacing}pt spacing after table (from margin-bottom)")
        except Exception as e:
            logger.debug(f"Could not apply spacing after table: {e}")

    def _apply_container_spacing_before(self, node: DOMNode):
        """
        DEPRECATED: No longer used. Implements CSS margin collapse behavior.

        In CSS, vertical margins collapse. To mimic this in Word, we only apply
        space_after on elements, never space_before. This method is kept for
        compatibility but does nothing.
        """
        pass  # Intentionally do nothing - implements margin collapse

    def _apply_container_spacing_after(self, node: DOMNode):
        """
        DEPRECATED: No longer used. Implements CSS margin collapse behavior.

        Container elements (section, article, etc.) don't need explicit spacing.
        Their child elements handle spacing via their own margin-bottom (space_after).
        This method is kept for compatibility but does nothing.
        """
        pass  # Intentionally do nothing - child elements handle their own spacing

    def _determine_flex_columns(self, flex_wrap: str, num_items: int) -> int:
        """
        Determine number of columns for flex layout based on flex-wrap.

        Args:
            flex_wrap: CSS flex-wrap value (nowrap, wrap, wrap-reverse)
            num_items: Number of flex items

        Returns:
            Number of columns
        """
        if flex_wrap in ('wrap', 'wrap-reverse'):
            # For wrapping flex, estimate columns based on square-ish layout
            # This is a heuristic - in reality it depends on item widths
            # Use a reasonable default that works well for most cases
            if num_items <= 3:
                return num_items  # Single row
            elif num_items <= 6:
                return 3  # 2 rows of 3
            else:
                return 4  # Multiple rows of 4
        else:
            # For nowrap (default), all items in one row
            # Limit to 6 columns for readability
            return min(num_items, 6)

    def _determine_grid_columns(self, grid_template_cols: str, num_items: int) -> int:
        """
        Determine number of columns from grid-template-columns or auto-fit/auto-fill.

        Args:
            grid_template_cols: CSS grid-template-columns value
            num_items: Number of grid items

        Returns:
            Number of columns
        """
        if not grid_template_cols:
            # Default: try to make a square-ish grid
            return max(1, int(math.sqrt(num_items)))

        # Count explicit column definitions
        # Handle: "repeat(3, 1fr)", "200px 200px 200px", "repeat(auto-fit, minmax(250px, 1fr))"
        if 'repeat' in grid_template_cols.lower():
            # Try to extract repeat count
            import re
            match = re.search(r'repeat\s*\(\s*(\d+|auto-fit|auto-fill)', grid_template_cols)
            if match:
                count_str = match.group(1)
                if count_str.isdigit():
                    return int(count_str)
                else:
                    # auto-fit/auto-fill: use actual item count for single-row layout
                    # In Word output (print), we want all items in one row
                    # Limit to reasonable max (e.g., 6 columns)
                    return min(6, num_items)

        # Count space-separated values
        parts = grid_template_cols.strip().split()
        if len(parts) > 0:
            return len(parts)

        # Default
        return min(3, num_items)

    def _is_icon_text_pattern(self, node: DOMNode) -> bool:
        """
        Check if a div is an icon + text pattern (like tips/hints).
        Pattern: div containing a small SVG icon followed by text content.

        Args:
            node: DOM node to check

        Returns:
            True if this is an icon+text pattern that should be treated as single paragraph
        """
        if node.tag != 'div':
            return False

        # Check for common tip/hint class names
        classes = node.attributes.get('class', [])
        if isinstance(classes, str):
            classes = classes.split()

        tip_classes = ['overall-tips', 'tips', 'hint', 'tip', 'info-tip', 'warning-tip']
        if any(tc in classes for tc in tip_classes):
            # Verify it has SVG + text structure
            has_svg = False
            has_text = False
            for child in node.children:
                if child.is_element and child.tag == 'svg':
                    has_svg = True
                elif child.is_element and child.tag in ('span', 'p', 'div', 'a'):
                    has_text = True
                elif child.is_text and child.text_content.strip():
                    has_text = True
            return has_svg and has_text

        # Also check for flex display with small SVG
        display = node.computed_styles.get('display', '')
        if display == 'flex':
            children = [c for c in node.children if c.is_element]
            if len(children) >= 2:
                first_child = children[0]
                if first_child.tag == 'svg':
                    # Check if SVG is small (icon-sized)
                    width = first_child.computed_styles.get('width', '')
                    height = first_child.computed_styles.get('height', '')
                    if width and height:
                        try:
                            w = float(width.replace('px', ''))
                            h = float(height.replace('px', ''))
                            if w <= 24 and h <= 24:  # Small icon
                                return True
                        except:
                            pass

        return False

    def _contains_svg(self, node: DOMNode) -> bool:
        """
        Recursively check if a node contains SVG elements.
        使用缓存避免重复遍历。

        Args:
            node: DOM node to check

        Returns:
            True if node or any descendants contain SVG
        """
        node_id = id(node)
        if node_id in self._svg_cache:
            return self._svg_cache[node_id]

        result = False
        for child in node.children:
            if child.is_element:
                if child.tag == 'svg':
                    result = True
                    break
                if self._contains_svg(child):
                    result = True
                    break

        self._svg_cache[node_id] = result
        return result

    def _is_root_layout_container(self, node: DOMNode) -> bool:
        """
        Check if node is a root layout container (like .container) that should not be wrapped.

        Args:
            node: DOM node

        Returns:
            True if it's a root layout container
        """
        # Check if parent is body or a direct child of a major container
        is_child_of_body = node.parent and node.parent.tag == 'body'

        # Check styles:
        # - Has white/transparent background
        # - Has box-shadow (likely for visual separation, not content styling)
        # - Likely has max-width, margin: 0 auto (centered layout)
        if not node.computed_styles:
            return False

        styles = node.computed_styles

        # Check if it has white/transparent background
        bg_color = styles.get('background-color', '')
        bg_prop = styles.get('background', '')

        has_white_bg = False
        if bg_color:
            bg_lower = bg_color.lower()
            has_white_bg = bg_lower in ('#fff', '#ffffff', 'white', 'rgb(255,255,255)', 'rgb(255, 255, 255)')
        elif bg_prop:
            bg_lower = bg_prop.lower()
            has_white_bg = bg_lower in ('white', '#fff', '#ffffff')

        # If not white background, it's not a layout container
        if not has_white_bg:
            return False

        # Check if has box-shadow (layout visual effect)
        has_shadow = 'box-shadow' in styles and styles['box-shadow'] not in ('none', '')

        # Check if has max-width (typical of layout containers)
        has_max_width = 'max-width' in styles

        # It's a root layout container if:
        # 1. Direct child of body (or very close to root)
        # 2. Has white background
        # 3. Has box-shadow (for visual separation)
        # 4. Probably has max-width for centering
        return is_child_of_body and has_shadow and (has_max_width or True)

    def _process_svg(self, svg_node: DOMNode):
        """
        Process SVG element and convert to image.

        Args:
            svg_node: SVG DOM node
        """
        try:
            # Get SVG dimensions from attributes or CSS
            width = svg_node.get_attribute('width') or svg_node.computed_styles.get('width', '100')
            height = svg_node.get_attribute('height') or svg_node.computed_styles.get('height', '100')

            logger.info(f"Processing SVG element: width={width}, height={height}")

            # Check if we're in a table cell context
            in_cell = getattr(self, 'in_table_cell', False)
            if in_cell:
                logger.debug("SVG is being inserted in a table cell, will use reduced max width")

            # Convert SVG to image using ImageBuilder's SVG support
            result = self.image_builder.build_svg(svg_node, width, height, in_cell)

            if result:
                logger.info(f"Successfully processed SVG element")
            else:
                logger.warning(f"SVG processing returned None")

        except Exception as e:
            logger.error(f"Failed to process SVG element: {e}", exc_info=True)
            # Fallback: skip the SVG element
            pass

    def _is_el_table_header(self, node: DOMNode) -> bool:
        """
        Check if a table node is an Element UI table header.

        Element UI tables split the header and body into separate <table> elements:
        - <table class="el-table__header"> contains only <thead>
        - <table class="el-table__body"> contains only <tbody>

        Args:
            node: Table DOM node

        Returns:
            True if this is an el-table__header
        """
        if node.tag != 'table':
            return False

        # Use attributes.get() for consistency with rest of codebase
        classes = node.attributes.get('class', '')

        # Handle both list and string class attributes
        if isinstance(classes, list):
            return 'el-table__header' in classes
        elif isinstance(classes, str):
            # Check if 'el-table__header' is a complete class name (not substring)
            return 'el-table__header' in classes.split()
        return False

    def _find_el_table_body(self, header_node: DOMNode) -> Optional[DOMNode]:
        """
        Find the corresponding el-table__body for an el-table__header.

        Element UI often wraps el-table__header and el-table__body in separate divs.
        This method searches the parent's siblings (i.e., at the grandparent level)
        to find the el-table__body table.

        Args:
            header_node: The el-table__header node

        Returns:
            The el-table__body node if found, None otherwise
        """
        # 性能优化：先检查预索引缓存
        cached = self._el_table_pairs.get(id(header_node))
        if cached is not None:
            logger.debug("Using cached el-table pair")
            return cached

        if not header_node.parent:
            logger.debug("Header node has no parent, cannot find body")
            return None

        # Get the parent node (typically div.el-table__header-wrapper)
        parent = header_node.parent

        # Check parent's class to verify it's a header wrapper
        parent_classes = parent.attributes.get('class', '')
        if isinstance(parent_classes, list):
            parent_classes = ' '.join(parent_classes)
        logger.debug(f"Header parent classes: {parent_classes}")

        if not parent.parent:
            logger.debug("Header parent has no grandparent, cannot find body")
            return None

        # Get the grandparent node
        grandparent = parent.parent

        # Find the index of parent in grandparent's children
        try:
            parent_idx = grandparent.children.index(parent)
        except ValueError:
            logger.debug("Could not find parent index in grandparent")
            return None

        # Helper function to check if a node has el-table__body class
        def has_body_class(node):
            if node.tag != 'table':
                return False
            classes = node.attributes.get('class', '')
            if isinstance(classes, list):
                return 'el-table__body' in classes
            elif isinstance(classes, str):
                # Check if 'el-table__body' is a complete class name (not substring)
                return 'el-table__body' in classes.split()
            return False

        # Search for el-table__body in subsequent siblings of parent
        logger.debug(f"Searching for el-table__body in {len(grandparent.children) - parent_idx - 1} siblings")
        for i in range(parent_idx + 1, len(grandparent.children)):
            sibling = grandparent.children[i]

            # Check if sibling itself is el-table__body
            if has_body_class(sibling):
                logger.info(f"Found el-table__body as direct sibling")
                return sibling

            # Check if sibling contains el-table__body (e.g., wrapped in div)
            if sibling.is_element:
                # Log sibling info
                sibling_classes = sibling.attributes.get('class', '')
                if isinstance(sibling_classes, list):
                    sibling_classes = ' '.join(sibling_classes)
                logger.debug(f"Checking sibling {sibling.tag} with classes: {sibling_classes}")

                for child in sibling.children:
                    if has_body_class(child):
                        logger.info(f"Found el-table__body wrapped in sibling {sibling.tag}")
                        return child

        logger.warning("Could not find el-table__body after searching all siblings")
        return None

    def _build_merged_el_table(self, header_node: DOMNode, body_node: DOMNode):
        """
        Build a Word table by merging el-table__header and el-table__body.

        Creates a single Word table with rows from both header and body tables.

        Args:
            header_node: The el-table__header node (contains <thead>)
            body_node: The el-table__body node (contains <tbody>)
        """
        # Extract rows from both tables
        header_rows = self.table_builder._extract_rows(header_node)
        body_rows = self.table_builder._extract_rows(body_node)

        if not header_rows and not body_rows:
            logger.warning("No rows found in merged el-table")
            return

        # Merge rows
        all_rows = header_rows + body_rows
        num_rows = len(all_rows)
        num_cols = self.table_builder._calculate_columns(all_rows)

        if num_rows == 0 or num_cols == 0:
            return

        logger.debug(f"Building merged el-table: {num_rows} rows × {num_cols} columns")

        # Create Word table
        table = self.document.add_table(rows=num_rows, cols=num_cols)

        # Fill table with merged rows
        self.table_builder._fill_table(table, all_rows)

        # Apply table-level styles (prefer header_node styles, as it usually has the main table styling)
        self.table_builder._apply_table_style(table, header_node)

    def save(self, output_path: str):
        """
        Save document to file.

        Args:
            output_path: Output file path
        """
        self.document.save(output_path)
        logger.info(f"Document saved to: {output_path}")

    def configure_header_footer(self, **kwargs):
        """
        Configure header and footer settings.

        This method allows users to customize header/footer settings programmatically.

        Args:
            **kwargs: Configuration parameters to update

        Example:
            # Change footer text
            builder.configure_header_footer(
                FOOTER_LEFT_TEXT="New contact information",
                FOOTER_FONT_SIZE=10
            )

            # Change header images
            builder.configure_header_footer(
                HEADER_LEFT_IMAGE="path/to/new/image.png",
                HEADER_IMAGE_MAX_HEIGHT=1.0
            )
        """
        if self.header_footer_builder:
            self.header_footer_builder.update_config(**kwargs)
            logger.info(f"Updated header/footer configuration with {len(kwargs)} settings")

    def disable_header_footer(self):
        """Disable header and footer for this document."""
        self.enable_header_footer = False
        logger.info("Headers and footers disabled")

    def enable_header_footer_feature(self):
        """Enable header and footer for this document."""
        self.enable_header_footer = True
        logger.info("Headers and footers enabled")
