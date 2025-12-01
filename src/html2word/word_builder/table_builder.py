"""
Table builder for Word documents.

Builds Word tables from HTML table elements with proper cell merging.
"""

import logging
from typing import List, Optional, Tuple, Dict, Any
from docx.shared import Inches, Pt

from html2word.parser.dom_tree import DOMNode
from html2word.word_builder.style_mapper import StyleMapper
from html2word.word_builder.paragraph_builder import ParagraphBuilder

logger = logging.getLogger(__name__)


class TableBuilder:
    """Builds Word tables from HTML table nodes."""

    def __init__(self, document):
        """
        Initialize table builder.

        Args:
            document: python-docx Document object
        """
        self.document = document
        self.style_mapper = StyleMapper()
        self.paragraph_builder = ParagraphBuilder(document)

    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        """
        Normalize whitespace in text while preserving leading/trailing spaces.

        This function collapses multiple consecutive whitespace characters (spaces, tabs,
        newlines) into a single space, but preserves single leading and trailing spaces.
        This is important for HTML rendering where spaces between inline elements are significant.

        Examples:
            "  hello  world  " -> " hello world "
            "hello\n\n  world" -> "hello world"
            "  text  " -> " text "

        Args:
            text: Input text to normalize

        Returns:
            Normalized text with collapsed whitespace but preserved boundaries
        """
        if not text:
            return ""

        # Check if text has leading/trailing whitespace before normalization
        has_leading_space = text and text[0].isspace()
        has_trailing_space = text and text[-1].isspace()

        # Collapse all whitespace
        normalized = ' '.join(text.split())

        # Restore leading/trailing space if originally present
        if has_leading_space and normalized:
            normalized = ' ' + normalized
        if has_trailing_space and normalized:
            normalized = normalized + ' '

        return normalized

    def build_table(self, table_node: DOMNode) -> Optional[object]:
        """
        Build a Word table from HTML table node.

        Args:
            table_node: Table DOM node

        Returns:
            python-docx Table object or None
        """
        if table_node.tag != 'table':
            return None

        # Extract table structure
        rows = self._extract_rows(table_node)
        if not rows:
            return None

        # Calculate dimensions
        num_rows = len(rows)
        num_cols = self._calculate_columns(rows)

        if num_rows == 0 or num_cols == 0:
            return None

        # Create table
        table = self.document.add_table(rows=num_rows, cols=num_cols)

        # Fill table cells
        self._fill_table(table, rows)

        # Apply table-level styles
        self._apply_table_style(table, table_node)

        return table

    def _extract_rows(self, table_node: DOMNode) -> List[DOMNode]:
        """
        Extract row nodes from table.

        Args:
            table_node: Table DOM node

        Returns:
            List of row nodes (tr)
        """
        rows = []

        # Process children to find rows
        for child in table_node.children:
            if child.tag == 'tr':
                rows.append(child)
            elif child.tag in ('thead', 'tbody', 'tfoot'):
                # Process section
                for section_child in child.children:
                    if section_child.tag == 'tr':
                        rows.append(section_child)

        return rows

    def _calculate_columns(self, rows: List[DOMNode]) -> int:
        """
        Calculate number of columns in table, excluding Element UI gutter columns.

        Element UI tables often have a "gutter" column (for scrollbar space) which should be ignored.
        This is typically:
        - The last column
        - Has empty content
        - Has class containing "gutter" or col name containing "gutter"

        Args:
            rows: List of row nodes

        Returns:
            Number of columns (excluding gutter)
        """
        if not rows:
            return 0

        max_cols = 0

        for row in rows:
            col_count = 0
            for cell in row.children:
                if cell.tag in ('td', 'th'):
                    colspan = int(cell.get_attribute('colspan', '1'))
                    col_count += colspan
            max_cols = max(max_cols, col_count)

        # Check if the last column in ALL rows is a gutter column
        # If so, exclude it from the column count
        if max_cols > 0 and self._has_gutter_column(rows, max_cols):
            logger.debug(f"Detected gutter column, reducing column count from {max_cols} to {max_cols - 1}")
            return max_cols - 1

        return max_cols

    def _has_gutter_column(self, rows: List[DOMNode], max_cols: int) -> bool:
        """
        Check if there is an Element UI gutter column.

        A gutter column is identified by:
        1. Any cell has class="gutter" (most reliable, check this FIRST)
        2. Element UI pattern: some rows have N+1 cells, others have N cells
        3. The extra cell is empty

        Args:
            rows: List of row nodes
            max_cols: Maximum number of cells in any row

        Returns:
            True if gutter column detected
        """
        if max_cols <= 1:
            return False

        # PRIORITY 1: Check if ANY cell has class="gutter"
        # This is the most reliable indicator and works even with single-row tables
        for row in rows:
            cells = [c for c in row.children if c.tag in ('td', 'th')]
            for cell in cells:
                cell_classes = cell.attributes.get('class', '')
                if isinstance(cell_classes, list):
                    cell_classes = ' '.join(cell_classes)
                elif not isinstance(cell_classes, str):
                    cell_classes = ''

                if 'gutter' in cell_classes.lower():
                    logger.debug(f"Detected gutter column via class='gutter'")
                    return True

        # PRIORITY 2: Check Element UI pattern (some rows have more cells than others)
        row_cell_counts = []
        for row in rows:
            cells = [c for c in row.children if c.tag in ('td', 'th')]
            row_cell_counts.append(len(cells))

        if not row_cell_counts:
            return False

        min_count = min(row_cell_counts)
        max_count = max(row_cell_counts)

        if max_count == min_count + 1:
            # Some rows have one extra column
            # Check if rows with max_count have empty last cell
            for row in rows:
                cells = [c for c in row.children if c.tag in ('td', 'th')]
                if len(cells) == max_count:
                    last_cell = cells[-1]

                    # Check if cell has empty content
                    text_content = last_cell.get_text_content() if hasattr(last_cell, 'get_text_content') else ''
                    if not text_content.strip():
                        logger.debug(f"Detected gutter column via empty last cell in longer row")
                        return True

        return False

    def _fill_table(self, table, rows: List[DOMNode]):
        """
        Fill table with content and handle cell merging.

        Args:
            table: python-docx Table object
            rows: List of row nodes
        """
        # Build cell matrix to track merged cells
        matrix = self._build_cell_matrix(rows)

        # Detect gutter column (for Element UI tables)
        # Identify which rows have gutter column
        rows_with_gutter = set()
        max_html_cols = 0
        min_html_cols = float('inf')

        for idx, row in enumerate(rows):
            col_count = sum(1 for c in row.children if c.tag in ('td', 'th'))
            max_html_cols = max(max_html_cols, col_count)
            min_html_cols = min(min_html_cols, col_count)

        # Element UI pattern: some rows have more cells than others
        # Rows with extra cells have gutter
        if max_html_cols > min_html_cols:
            for idx, row in enumerate(rows):
                cells = [c for c in row.children if c.tag in ('td', 'th')]
                if len(cells) == max_html_cols:
                    # This row has extra cells, check if last is gutter
                    last_cell = cells[-1]
                    cell_classes = last_cell.attributes.get('class', '')
                    if isinstance(cell_classes, list):
                        cell_classes = ' '.join(cell_classes)

                    text_content = last_cell.get_text_content() if hasattr(last_cell, 'get_text_content') else ''

                    # If last cell is empty or has "gutter" class, mark this row
                    if 'gutter' in cell_classes.lower() or not text_content.strip():
                        rows_with_gutter.add(idx)
                        logger.debug(f"Row {idx} has gutter column")

        # Fill cells
        for row_idx, row_node in enumerate(rows):
            word_row = table.rows[row_idx]

            # Set row height based on CSS height property
            self._apply_row_height(word_row, row_node)

            # Set row to not split across pages
            self._set_row_cant_split(word_row)

            # Check if this is a header row
            is_header = False
            if row_node.parent and row_node.parent.tag == 'thead':
                is_header = True
            elif any(c.tag == 'th' for c in row_node.children if c.is_element):
                is_header = True
            
            if is_header:
                self._set_row_as_header(word_row)
                self._set_row_keep_with_next(word_row)

            col_idx = 0

            # Get cells for this row
            cells_in_row = [c for c in row_node.children if c.tag in ('td', 'th')]

            for cell_idx, cell_node in enumerate(cells_in_row):
                if cell_node.tag not in ('td', 'th'):
                    continue

                # Skip gutter column ONLY if this row has gutter and it's the last cell
                if row_idx in rows_with_gutter and cell_idx == len(cells_in_row) - 1:
                    logger.debug(f"Skipping gutter cell in row {row_idx}")
                    continue

                # Skip cells that are occupied by a previous rowspan/colspan
                # (matrix[row_idx][col_idx] points to a different cell_node)
                while col_idx < len(matrix[row_idx]) and matrix[row_idx][col_idx] != cell_node:
                    col_idx += 1

                if col_idx >= len(word_row.cells):
                    break

                word_cell = word_row.cells[col_idx]

                # Get colspan and rowspan
                colspan = int(cell_node.get_attribute('colspan', '1'))
                rowspan = int(cell_node.get_attribute('rowspan', '1'))

                # Prepare cell styles - merge row styles with cell styles
                cell_styles = self._get_merged_cell_styles(row_node, cell_node)
                box_model = cell_node.layout_info.get('box_model')

                # Handle cell merging
                if colspan > 1 or rowspan > 1:
                    # CRITICAL: Apply borders to ALL cells BEFORE merging
                    # This ensures consistent border appearance across merged cells
                    self._apply_borders_before_merge(table, row_idx, col_idx, rowspan, colspan, box_model)
                    # Now perform the merge
                    self._merge_cells(table, row_idx, col_idx, rowspan, colspan)

                # Apply cell-level styles (background, borders, alignment)
                # For merged cells, this re-applies to ensure consistency
                self.style_mapper.apply_table_cell_style(
                    word_cell,
                    cell_styles,
                    box_model
                )

                # Fill cell content with merged styles
                self._fill_cell(word_cell, cell_node, cell_styles)

                col_idx += colspan

    def _get_merged_cell_styles(self, row_node: DOMNode, cell_node: DOMNode) -> Dict[str, Any]:
        """
        Merge row styles with cell styles.

        Args:
            row_node: Row DOM node
            cell_node: Cell DOM node

        Returns:
            Merged styles dictionary
        """
        merged_styles = {}

        # Start with row styles
        if row_node.computed_styles:
            merged_styles.update(row_node.computed_styles)

        # Override with cell styles
        if cell_node.computed_styles:
            merged_styles.update(cell_node.computed_styles)

        # SPECIAL CASE: If cell has no background but contains a single main element with background,
        # promote that background to the cell level for better visual effect
        # This handles cases like <td><div class="chart-panel-header">...</div></td>
        if not merged_styles.get('background-color') or merged_styles.get('background-color') in ('transparent', 'rgba(0,0,0,0)'):
            # Check if cell has a single main content element
            content_children = [child for child in cell_node.children
                              if child.tag in ('div', 'section', 'article', 'header', 'footer', 'aside', 'nav')]

            # If there's exactly one main content element, check its background
            if len(content_children) == 1:
                main_child = content_children[0]
                child_bg = main_child.computed_styles.get('background-color')

                # Also check for the special case of height=line-height (vertical centering)
                child_height = main_child.computed_styles.get('height')
                child_line_height = main_child.computed_styles.get('line-height')

                if child_bg and child_bg not in ('transparent', 'rgba(0,0,0,0)'):
                    # Promote the background to cell level
                    merged_styles['background-color'] = child_bg

                    # If this is a vertical centering case (height=line-height),
                    # also promote the height to ensure proper cell height
                    if child_height and child_line_height:
                        try:
                            from html2word.utils.units import UnitConverter
                            height_pt = UnitConverter.to_pt(str(child_height))
                            line_height_pt = UnitConverter.to_pt(str(child_line_height))

                            if abs(height_pt - line_height_pt) < 2.0:
                                # This is vertical centering, promote height too
                                merged_styles['height'] = child_height
                                logger.debug(f"Promoted background #{child_bg} and height {child_height} from child element to cell")
                        except:
                            pass
                    else:
                        logger.debug(f"Promoted background #{child_bg} from child element to cell")

        return merged_styles

    def _build_cell_matrix(self, rows: List[DOMNode]) -> List[List[Optional[DOMNode]]]:
        """
        Build matrix to track cell positions including merged cells.

        Args:
            rows: List of row nodes

        Returns:
            Matrix where None indicates a merged cell
        """
        if not rows:
            return []

        num_rows = len(rows)
        num_cols = self._calculate_columns(rows)
        matrix = [[None for _ in range(num_cols)] for _ in range(num_rows)]

        # Detect which rows have gutter columns
        rows_with_gutter = set()
        max_html_cols = 0
        min_html_cols = float('inf')

        for idx, row in enumerate(rows):
            col_count = sum(1 for c in row.children if c.tag in ('td', 'th'))
            max_html_cols = max(max_html_cols, col_count)
            min_html_cols = min(min_html_cols, col_count)

        if max_html_cols > min_html_cols:
            for idx, row in enumerate(rows):
                cells = [c for c in row.children if c.tag in ('td', 'th')]
                if len(cells) == max_html_cols:
                    last_cell = cells[-1]
                    cell_classes = last_cell.attributes.get('class', '')
                    if isinstance(cell_classes, list):
                        cell_classes = ' '.join(cell_classes)

                    text_content = last_cell.get_text_content() if hasattr(last_cell, 'get_text_content') else ''

                    if 'gutter' in cell_classes.lower() or not text_content.strip():
                        rows_with_gutter.add(idx)

        for row_idx, row_node in enumerate(rows):
            col_idx = 0

            # Get cells for this row
            cells_in_row = [c for c in row_node.children if c.tag in ('td', 'th')]

            for cell_idx, cell_node in enumerate(cells_in_row):
                if cell_node.tag not in ('td', 'th'):
                    continue

                # Skip gutter column ONLY if this row has gutter and it's the last cell
                if row_idx in rows_with_gutter and cell_idx == len(cells_in_row) - 1:
                    continue

                # Find next available column
                while col_idx < num_cols and matrix[row_idx][col_idx] is not None:
                    col_idx += 1

                if col_idx >= num_cols:
                    break

                # Get span values
                colspan = int(cell_node.get_attribute('colspan', '1'))
                rowspan = int(cell_node.get_attribute('rowspan', '1'))

                # Mark cells as occupied
                for r in range(row_idx, min(row_idx + rowspan, num_rows)):
                    for c in range(col_idx, min(col_idx + colspan, num_cols)):
                        matrix[r][c] = cell_node

                col_idx += colspan

        return matrix

    def _merge_cells(self, table, row_idx: int, col_idx: int, rowspan: int, colspan: int):
        """
        Merge table cells.

        Args:
            table: python-docx Table object
            row_idx: Starting row index
            col_idx: Starting column index
            rowspan: Number of rows to span
            colspan: Number of columns to span
        """
        try:
            # Get starting cell
            start_cell = table.rows[row_idx].cells[col_idx]

            # Get ending cell
            end_row_idx = min(row_idx + rowspan - 1, len(table.rows) - 1)
            end_col_idx = min(col_idx + colspan - 1, len(table.rows[0].cells) - 1)
            end_cell = table.rows[end_row_idx].cells[end_col_idx]

            # Merge
            if start_cell != end_cell:
                start_cell.merge(end_cell)
        except Exception as e:
            logger.warning(f"Error merging cells: {e}")

    def _apply_borders_before_merge(self, table, row_idx: int, col_idx: int,
                                      rowspan: int, colspan: int, box_model):
        """
        Apply border styles to all cells BEFORE merging.

        This is critical for preventing border color inconsistencies in merged cells.
        When cells are merged in Word, each cell in the merged range retains its own
        border settings. If these borders differ, you'll see inconsistent colors/styles.

        Solution: Apply the same border to ALL cells in the range BEFORE merging.

        Args:
            table: python-docx Table object
            row_idx: Starting row index
            col_idx: Starting column index
            rowspan: Number of rows to span
            colspan: Number of columns to span
            box_model: BoxModel object with border information
        """
        if not box_model or not box_model.border.has_border():
            return

        try:
            # Calculate the range of cells to be merged
            end_row_idx = min(row_idx + rowspan, len(table.rows))

            # Apply the same border to ALL cells in the merge range
            for r in range(row_idx, end_row_idx):
                # Get the number of cells in this row
                num_cells = len(table.rows[r].cells)
                end_col_idx = min(col_idx + colspan, num_cells)

                for c in range(col_idx, end_col_idx):
                    try:
                        cell = table.rows[r].cells[c]
                        # Apply the source cell's borders to this cell
                        self.style_mapper._apply_cell_borders(cell, box_model.border)
                        logger.debug(f"Applied border to cell [{r},{c}] before merge")
                    except Exception as e:
                        logger.debug(f"Could not apply border to cell [{r},{c}]: {e}")
                        continue

        except Exception as e:
            logger.warning(f"Error applying borders before merge: {e}")

    def _fill_cell(self, word_cell, cell_node: DOMNode, cell_styles: Dict[str, Any]):
        """
        Fill cell with content.

        Args:
            word_cell: python-docx Cell object
            cell_node: Cell DOM node
            cell_styles: Merged styles (row + cell)
        """
        # Make a copy to avoid modifying the original dict
        cell_styles = cell_styles.copy()

        # Apply default font and size for table cells if not specified or using browser defaults
        # This ensures tables without explicit font styles still render with appropriate fonts
        # Check if font-family is missing, empty, or a browser default (Segoe UI, Arial without explicit CSS)
        current_font = cell_styles.get('font-family', '')

        if not current_font or current_font in ('Segoe UI', 'sans-serif'):
            cell_styles['font-family'] = 'Microsoft YaHei, Arial, sans-serif'
            logger.debug(f"Applied default font-family (was: {repr(current_font)})")

        # Check if font-size is missing, empty, or too small (likely browser default)
        current_size = cell_styles.get('font-size', '')
        # Handle both string ('12px') and numeric (9.0, 12.0, 16.0) formats
        # Replace small sizes (9px, 10px, 11px, 12px) with Chinese document standard (14px/10.5pt)
        should_replace_size = (
            not current_size or
            current_size in ('9px', '10px', '11px', '12px', '16px') or
            (isinstance(current_size, (int, float)) and current_size <= 12.0)
        )
        if should_replace_size:
            cell_styles['font-size'] = '14px'  # Common size for Chinese documents (10.5pt)
            logger.debug(f"Applied default font-size (was: {repr(current_size)})")

        # Get or create first paragraph
        if word_cell.paragraphs:
            paragraph = word_cell.paragraphs[0]
            # Clear default content
            paragraph.clear()
        else:
            paragraph = word_cell.add_paragraph()

        # CRITICAL FIX: In table cells, do NOT apply box_model margins as paragraph spacing
        # Table cell vertical spacing should be controlled by cell padding ONLY
        # If we apply box_model (which includes margins), it creates extra space_after
        # that causes table rows to be too tall
        # Solution: Pass None for box_model to disable margin-based spacing

        # CRITICAL FIX: In table cells, do NOT apply cell background-color to paragraphs
        # This creates the "sandwich" effect where paragraph background overlays cell background
        # Cell background is already applied via apply_table_cell_style
        # Solution: Remove background-color from paragraph styles
        para_styles = cell_styles.copy()
        if 'background-color' in para_styles:
            del para_styles['background-color']

        box_model = cell_node.layout_info.get('box_model')
        # Table cells: limit line spacing to 1.2 for compact layout, auto-detect HTML line-height
        # FIXED: Respect Element UI alignment classes for table headers and cells
        # Check for Element UI alignment classes first
        cell_classes = cell_node.attributes.get('class', [])
        if isinstance(cell_classes, str):
            cell_classes = cell_classes.split()

        # Determine alignment based on Element UI classes or existing styles
        if 'is-center' in cell_classes:
            para_styles['text-align'] = 'center'
            logger.debug("Applied center alignment for Element UI 'is-center' class")
        elif 'is-right' in cell_classes:
            para_styles['text-align'] = 'right'
            logger.debug("Applied right alignment for Element UI 'is-right' class")
        elif 'is-left' in cell_classes:
            para_styles['text-align'] = 'left'
            logger.debug("Applied left alignment for Element UI 'is-left' class")
        else:
            # Keep table cell content left-aligned by default for better readability
            # Left alignment is more appropriate for tabular data
            text_align = para_styles.get('text-align')
            if text_align in ('left', 'start', None):
                para_styles['text-align'] = 'left'
                logger.debug("Keeping text-align as 'left' for table cell content")

        # Apply paragraph styles
        self.style_mapper.apply_paragraph_style(paragraph, para_styles, box_model=None, max_line_spacing=1.2)

        # Check if we've added any content
        has_content = False

        # Process all children
        for child in cell_node.children:
            # Skip Element UI hidden-columns div (contains column templates, not visible content)
            if child.is_element and self._should_skip_element(child):
                logger.debug(f"Skipping element with class '{child.attributes.get('class', '')}' in table cell")
                continue

            if child.is_text:
                text = child.text or ""
                # Normalize whitespace while preserving leading/trailing spaces
                text = self._normalize_whitespace(text)
                if text.strip():  # Only add if there's non-whitespace content
                    # Add text with merged styles
                    run = paragraph.add_run(text)
                    self.style_mapper.apply_run_style(run, cell_styles)
                    has_content = True

            elif child.is_element:
                if child.tag == 'br':
                    # Line break in table cell - add line break (soft return), not new paragraph
                    # This preserves HTML <br/> behavior in table cells
                    logger.debug(f"Found <br> tag in table cell, adding line break")
                    paragraph.add_run('\n')
                    has_content = True

                elif child.tag == 'img':
                    # FIXED: Handle images in table cells
                    self._add_cell_image(paragraph, child)
                    has_content = True

                elif child.tag == 'svg':
                    # FIXED: Handle SVG in table cells
                    self._add_cell_svg(paragraph, child)
                    has_content = True

                elif child.tag in ('p', 'div'):
                    # Block element
                    if has_content:
                        # Create new paragraph for block element
                        paragraph = word_cell.add_paragraph()
                        # Apply paragraph spacing for multiple paragraphs in table cells
                        # This ensures proper spacing between <p> tags in HTML
                        if child.tag == 'p':
                            # Add space before paragraph for separation (except for first paragraph)
                            # Calculate spacing based on font size (20% of font size, 2-3pt range)
                            font_size_pt = self._get_font_size_pt(cell_styles.get('font-size', '14px'))
                            spacing_pt = max(2.0, min(3.0, font_size_pt * 0.2))
                            paragraph.paragraph_format.space_before = Pt(spacing_pt)
                    # Process block content and update paragraph reference
                    paragraph = self._add_cell_block_content(paragraph, child, cell_styles)
                    has_content = True

                elif child.tag in ('strong', 'b', 'em', 'i', 'u', 'span', 'a'):
                    # Inline formatting
                    self._add_cell_inline_content(paragraph, child, cell_styles)
                    has_content = True

                elif child.tag == 'table':
                    # Nested table - Word supports tables within table cells
                    nested_table = self._build_nested_table(word_cell, child)
                    if nested_table:
                        has_content = True

                else:
                    # Other elements - process recursively to preserve nested styles
                    # IMPORTANT: Don't just extract text, process children to preserve their styles
                    self._process_cell_element_children(paragraph, child, cell_styles)
                    has_content = True


        # FINAL FIX: Clean up extra line breaks and force left alignment for multiline content
        # This addresses the Major Types column issues with extra line breaks
        for para in word_cell.paragraphs:
            if para.text or para.runs:  # Process all paragraphs with content
                # Clean up extra line breaks in runs
                if para.runs:
                    # Remove leading newlines from first run
                    if para.runs[0].text and para.runs[0].text.startswith('\n'):
                        para.runs[0].text = para.runs[0].text.lstrip('\n')
                        logger.debug("Removed leading line breaks from table cell")

                    # Remove trailing newlines from last run
                    if para.runs[-1].text and para.runs[-1].text.endswith('\n'):
                        para.runs[-1].text = para.runs[-1].text.rstrip('\n')
                        logger.debug("Removed trailing line breaks from table cell")

                # Force left alignment ONLY for multiline content (contains actual line breaks)
                # This is needed for content with <br> tags that should maintain line structure
                if para.text.strip() and '\n' in para.text:
                    current_alignment = para.paragraph_format.alignment
                    if current_alignment not in [1, 2]:  # Not CENTER or RIGHT
                        para.paragraph_format.alignment = 0  # Force LEFT
                        logger.debug("Forced left alignment for multiline paragraph with line breaks")


    def _process_cell_element_children(self, paragraph, element: DOMNode, base_styles: Dict[str, Any]):
        """
        Process children of an element in a table cell, preserving nested styles.

        Args:
            paragraph: Current Word paragraph
            element: Element whose children to process
            base_styles: Base styles from parent
        """
        # Merge element's styles with base styles
        element_styles = base_styles.copy()
        element_styles.update(element.computed_styles)

        logger.debug(f"_process_cell_element_children: {element.tag}, base color: {base_styles.get('color')}, element color: {element.computed_styles.get('color')}, merged color: {element_styles.get('color')}")

        for child in element.children:
            if child.is_text:
                text = child.text or ""
                text = self._normalize_whitespace(text)
                if text.strip():
                    run = paragraph.add_run(text)
                    self.style_mapper.apply_run_style(run, element_styles)

            elif child.is_element:
                if child.tag in ('strong', 'b', 'em', 'i', 'u', 'span', 'a'):
                    # Inline elements - process with their own styles
                    self._add_cell_inline_content(paragraph, child, element_styles)
                elif child.tag == 'br':
                    # Smart handling of <br> tags to avoid extra line breaks
                    # Check if we already have content and if it doesn't end with newline
                    existing_text = ''.join(run.text or '' for run in paragraph.runs)
                    # Only add line break if we have content before and no trailing newline
                    if existing_text and not existing_text.endswith('\n'):
                        paragraph.add_run('\n')
                    elif not existing_text:
                        # Skip br at the beginning of content
                        logger.debug("Skipped <br> at beginning of table cell")
                else:
                    # Recursively process other elements
                    self._process_cell_element_children(paragraph, child, element_styles)

    def _should_skip_element(self, node: DOMNode) -> bool:
        """
        Check if an element should be skipped during table cell content processing.

        Element UI tables have hidden template elements that should not be rendered:
        - hidden-columns: Contains column definition templates
        - Other hidden/template elements marked by specific classes

        Args:
            node: DOM node to check

        Returns:
            True if element should be skipped, False otherwise
        """
        if not node.is_element:
            return False

        # Check for hidden-columns class (Element UI column templates)
        classes = node.attributes.get('class', '')
        if isinstance(classes, list):
            classes = ' '.join(classes)
        elif not isinstance(classes, str):
            classes = ''

        # Skip hidden-columns div (Element UI table column templates)
        if 'hidden-columns' in classes:
            return True

        # Skip elements with display:none or visibility:hidden
        computed_styles = node.computed_styles or {}
        display = computed_styles.get('display', '')
        visibility = computed_styles.get('visibility', '')

        if display == 'none' or visibility == 'hidden':
            return True

        return False

    def _add_cell_block_content(self, paragraph, node: DOMNode, base_styles: Dict[str, Any]):
        """
        Add content from a block element to paragraph.

        Args:
            paragraph: Word paragraph (the current paragraph to write to)
            node: Block element node
            base_styles: Base styles from cell/row

        Returns:
            The last paragraph used (may create new paragraphs for nested blocks)
        """
        # Merge styles - node styles override base
        merged_styles = base_styles.copy()
        merged_styles.update(node.computed_styles)

        # CRITICAL FIX: In table cells, do NOT apply cell background-color to paragraphs
        # This creates the "sandwich" effect where paragraph background overlays cell background
        if 'background-color' in merged_styles:
            del merged_styles['background-color']

        # FIXED: Respect Element UI alignment classes for table headers and cells
        # If no explicit alignment set, default to left for table cells
        text_align = merged_styles.get('text-align')
        if text_align in ('left', 'start', None):
            merged_styles['text-align'] = 'left'

        # Apply paragraph-level styles
        # CRITICAL FIX: In table cells, do NOT apply box_model to avoid extra spacing
        # Table cells: limit line spacing to 1.2 for compact layout, auto-detect HTML line-height
        self.style_mapper.apply_paragraph_style(paragraph, merged_styles, box_model=None, max_line_spacing=1.2)

        # Get the parent cell to add new paragraphs if needed
        word_cell = paragraph._parent

        # Track if current paragraph has content
        has_content_in_paragraph = False

        # Process children
        for child in node.children:
            if child.is_text:
                text = child.text or ""
                text = self._normalize_whitespace(text)
                if text.strip():
                    run = paragraph.add_run(text)
                    self.style_mapper.apply_run_style(run, merged_styles)
                    has_content_in_paragraph = True

            elif child.is_element and child.tag == 'br':
                # Handle <br> tags inside block elements like <p>
                # Add line break (soft return) within the current paragraph
                logger.debug(f"Found <br> tag inside {node.tag}, adding line break")
                paragraph.add_run('\n')
                has_content_in_paragraph = True

            elif child.is_inline:
                self._add_cell_inline_content(paragraph, child, merged_styles)
                has_content_in_paragraph = True

            elif child.is_element and child.tag in ('p', 'div'):
                # Nested block element (p, div) - each should be on a new paragraph
                if has_content_in_paragraph:
                    # Current paragraph already has content, create new one
                    paragraph = word_cell.add_paragraph()
                    para_styles = merged_styles.copy()
                    if 'background-color' in para_styles:
                        del para_styles['background-color']
                    # FIXED: Keep table cell content left-aligned by default for better readability
                    # Left alignment is more appropriate for tabular data
                    text_align_nested = para_styles.get('text-align')
                    if text_align_nested in ('left', 'start', None):
                        para_styles['text-align'] = 'left'

                    self.style_mapper.apply_paragraph_style(paragraph, para_styles, box_model=None, max_line_spacing=1.2)
                    # Add spacing for nested <p> tags
                    if child.tag == 'p':
                        # Calculate spacing based on font size (20% of font size, 2-3pt range)
                        font_size_pt = self._get_font_size_pt(merged_styles.get('font-size', '14px'))
                        spacing_pt = max(2.0, min(3.0, font_size_pt * 0.2))
                        paragraph.paragraph_format.space_before = Pt(spacing_pt)
                    has_content_in_paragraph = False

                # Recursively process the nested block element
                # IMPORTANT: Pass merged_styles instead of base_styles to preserve style inheritance
                paragraph = self._add_cell_block_content(paragraph, child, merged_styles)
                has_content_in_paragraph = True

            else:
                # Other nested block elements - extract text
                text = child.get_text_content()
                text = self._normalize_whitespace(text)
                if text.strip():
                    run = paragraph.add_run(text)
                    self.style_mapper.apply_run_style(run, merged_styles)
                    has_content_in_paragraph = True

        return paragraph

    def _add_cell_inline_content(self, paragraph, node: DOMNode, base_styles: Dict[str, Any]):
        """
        Add inline element content to paragraph.

        Args:
            paragraph: Word paragraph
            node: Inline element node
            base_styles: Base styles from parent
        """
        # Merge styles - CRITICAL: node styles override base styles
        merged_styles = base_styles.copy()
        merged_styles.update(node.computed_styles)

        # DEBUG: Log color merging for span elements
        if node.tag == 'span' and 'color' in node.computed_styles:
            logger.debug(f"Span color merging: base={base_styles.get('color')}, node={node.computed_styles.get('color')}, merged={merged_styles.get('color')}")

        for child in node.children:
            if child.is_text:
                text = child.text or ""
                text = self._normalize_whitespace(text)
                if text.strip():
                    run = paragraph.add_run(text)
                    self.style_mapper.apply_run_style(run, merged_styles)

            elif child.is_element and child.tag == 'br':
                # Handle <br> tags inside inline elements
                logger.debug(f"Found <br> tag inside inline element {node.tag}, adding line break")
                paragraph.add_run('\n')

            elif child.is_inline:
                # Nested inline
                self._add_cell_inline_content(paragraph, child, merged_styles)

            else:
                # Block inside inline - extract text
                text = child.get_text_content()
                text = self._normalize_whitespace(text)
                if text.strip():
                    run = paragraph.add_run(text)
                    self.style_mapper.apply_run_style(run, merged_styles)

    def _build_nested_table(self, word_cell, table_node: DOMNode) -> Optional[object]:
        """
        Build a nested table within a table cell.

        Args:
            word_cell: python-docx Cell object (parent cell)
            table_node: Table DOM node (nested table)

        Returns:
            python-docx Table object or None
        """
        try:
            # Extract table structure
            rows = self._extract_rows(table_node)
            if not rows:
                return None

            # Calculate dimensions
            num_rows = len(rows)
            num_cols = self._calculate_columns(rows)

            if num_rows == 0 or num_cols == 0:
                return None

            # Create nested table in the cell
            # Note: We need to clear the default paragraph first
            if word_cell.paragraphs:
                # Clear the first paragraph but keep it
                word_cell.paragraphs[0].clear()

            # Add table to cell
            nested_table = word_cell.add_table(rows=num_rows, cols=num_cols)

            # Fill nested table cells
            self._fill_table(nested_table, rows)

            # Apply table-level styles
            self._apply_table_style(nested_table, table_node)

            logger.debug(f"Built nested table: {num_rows}x{num_cols}")
            return nested_table

        except Exception as e:
            logger.warning(f"Error building nested table: {e}")
            # Fallback: just extract text content
            text = table_node.get_text_content()
            text = self._normalize_whitespace(text)
            if text.strip() and word_cell.paragraphs:
                para = word_cell.paragraphs[0]
                para.add_run(text)
            return None

    def _add_cell_image(self, paragraph, img_node: DOMNode):
        """
        Add image to table cell paragraph.

        Args:
            paragraph: Word paragraph
            img_node: Image DOM node
        """
        try:
            from html2word.word_builder.image_builder import ImageBuilder
            from html2word.utils.image_utils import ImageProcessor
            from docx.shared import Inches

            image_builder = ImageBuilder(self.document)
            src = image_builder._get_best_image_src(img_node)
            if not src:
                return

            # Get CSS dimensions
            css_width = img_node.computed_styles.get('width')
            css_height = img_node.computed_styles.get('height')
            transform = img_node.computed_styles.get('transform')
            filter_value = img_node.computed_styles.get('filter')

            # Process image
            image_processor = ImageProcessor()
            result = image_processor.process_image(src, transform=transform, filter_css=filter_value)
            if not result:
                return

            image_stream, image_size = result

            # Calculate size
            max_width_inches = 2.0  # Default max for table images
            max_height_inches = 2.0

            if css_width:
                from html2word.utils.units import UnitConverter
                width_pt = UnitConverter.to_pt(css_width)
                max_width_inches = width_pt / 72

            if css_height:
                from html2word.utils.units import UnitConverter
                height_pt = UnitConverter.to_pt(css_height)
                max_height_inches = height_pt / 72

            display_width, display_height = image_processor.calculate_display_size(
                image_size, css_width, css_height, max_width_inches, max_height_inches
            )

            # Add image to paragraph run
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(display_width), height=Inches(display_height))

            logger.debug(f"Added image to table cell: {src}")

        except Exception as e:
            logger.warning(f"Failed to add image to table cell: {e}")

    def _add_cell_svg(self, paragraph, svg_node: DOMNode):
        """
        Add SVG to table cell paragraph (converted to image).

        Args:
            paragraph: Word paragraph
            svg_node: SVG DOM node
        """
        try:
            from html2word.word_builder.image_builder import ImageBuilder
            import io
            from docx.shared import Inches

            image_builder = ImageBuilder(self.document)
            width = svg_node.get_attribute('width') or svg_node.computed_styles.get('width', '100')
            height = svg_node.get_attribute('height') or svg_node.computed_styles.get('height', '100')

            svg_content = image_builder._serialize_svg_node(svg_node, width, height)
            if not svg_content:
                return

            try:
                import cairosvg
                png_data = cairosvg.svg2png(bytestring=svg_content.encode('utf-8'))
                image_stream = io.BytesIO(png_data)

                width_pt = image_builder._parse_dimension(width)
                height_pt = image_builder._parse_dimension(height)

                run = paragraph.add_run()
                run.add_picture(
                    image_stream,
                    width=Inches(width_pt / 72),
                    height=Inches(height_pt / 72)
                )

                logger.debug(f"Added SVG to table cell ({width}x{height})")

            except ImportError:
                logger.warning("cairosvg not available, SVG in table cell skipped")

        except Exception as e:
            logger.warning(f"Failed to add SVG to table cell: {e}")

    def _apply_table_style(self, table, table_node: DOMNode):
        """
        Apply table-level styles.

        Args:
            table: python-docx Table object
            table_node: Table DOM node
        """
        # Set table width if specified
        box_model = table_node.layout_info.get('box_model')
        if box_model and box_model.width:
            try:
                from docx.oxml.ns import qn

                # Limit table width to maximum page width (6.5 inches = 468pt)
                max_table_width_pt = 468
                width_pt = min(box_model.width, max_table_width_pt)  # Ensure it doesn't exceed page width
                width_dxa = int(width_pt / 72 * 1440)  # pt to DXA (twentieths of a point)

                # Get or create tblW element
                tbl = table._element
                tblPr = tbl.tblPr
                tblW = tblPr.find(qn('w:tblW'))

                if tblW is not None:
                    # Modify existing tblW
                    tblW.set(qn('w:w'), str(width_dxa))
                    tblW.set(qn('w:type'), 'dxa')
                else:
                    # Create new tblW (should not happen, but just in case)
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
                    tblPr.append(tblW)

                logger.debug(f"Set table width to {width_pt}pt = {width_pt/72:.2f}\" = {width_dxa} dxa")
            except Exception as e:
                logger.warning(f"Failed to set table width: {e}")

        # Apply column widths
        self._apply_column_widths(table, table_node)

        # Apply table style
        try:
            table.style = 'Table Grid'
        except:
            pass

    def _apply_column_widths(self, table, table_node: DOMNode):
        """
        Extract and apply column widths from HTML table with proportional scaling.

        This method ensures table columns fit within the Word page width while
        maintaining the relative proportions from the HTML source.

        Args:
            table: python-docx Table object
            table_node: Table DOM node
        """
        from html2word.utils.units import UnitConverter

        # Try to extract column widths from <col> tags or first row cells
        col_widths = self._extract_column_widths(table_node)

        if not col_widths:
            return  # No width information, use default

        # Get total table width - default to 6.5 inches for standard Word page with margins
        box_model = table_node.layout_info.get('box_model')
        # Maximum usable width in Word (6.5 inches = 468 pt for standard page)
        max_table_width_pt = 468

        # If box model specifies a width, use the smaller of specified width and max width
        if box_model and box_model.width:
            table_width_pt = min(box_model.width, max_table_width_pt)
        else:
            table_width_pt = max_table_width_pt

        try:
            # First pass: categorize columns and collect widths
            html_widths = []  # HTML widths in px
            is_percentage = []  # Track if width is percentage
            has_absolute_widths = False
            has_percentage_widths = False
            total_absolute_width_px = 0.0

            for col_idx, width_info in enumerate(col_widths):
                if col_idx >= len(table.columns):
                    break

                if width_info:
                    width_value = None
                    is_pct = False

                    # Parse width (can be px, %, pt, or just a number)
                    if isinstance(width_info, str):
                        if '%' in width_info:
                            # Percentage width
                            percent = float(width_info.rstrip('%'))
                            width_value = percent  # Store as percentage
                            is_pct = True
                            has_percentage_widths = True
                        else:
                            # Absolute width (px, pt, etc.)
                            width_pt = UnitConverter.to_pt(width_info)
                            width_value = width_pt / 0.75  # Convert pt to px for calculation
                            has_absolute_widths = True
                            total_absolute_width_px += width_value
                    elif isinstance(width_info, (int, float)):
                        # Numeric width, assume pixels
                        width_value = float(width_info)
                        has_absolute_widths = True
                        total_absolute_width_px += width_value

                    html_widths.append(width_value)
                    is_percentage.append(is_pct)
                else:
                    html_widths.append(None)
                    is_percentage.append(False)

            # If no valid widths found, return
            if not has_absolute_widths and not has_percentage_widths:
                logger.debug("No valid column widths found in HTML table")
                return

            # Determine column width strategy
            cols_with_no_width = sum(1 for w in html_widths if w is None)

            # Second pass: calculate Word column widths
            word_widths = []
            assigned_width = 0.0

            if has_absolute_widths and not has_percentage_widths:
                # Strategy 1: Mix of absolute and auto widths
                # Absolute widths should use their actual size (converted to pt)
                # Remaining space distributed to auto columns

                logger.debug(f"Using absolute width strategy: {total_absolute_width_px}px absolute, {cols_with_no_width} auto columns")

                # Convert total absolute width to pt
                total_absolute_width_pt = total_absolute_width_px * 0.75

                # Check if absolute widths exceed table width
                if total_absolute_width_pt >= table_width_pt - (cols_with_no_width * 30):
                    # Need to scale down - use proportional strategy instead
                    logger.debug("Absolute widths too large, falling back to proportional strategy")
                    total_html_width = total_absolute_width_px
                    for col_idx, html_width in enumerate(html_widths):
                        if html_width is not None:
                            proportion = html_width / total_html_width
                            word_width_pt = table_width_pt * proportion
                            word_width_pt = max(word_width_pt, 30)
                            word_widths.append(word_width_pt)
                            assigned_width += word_width_pt
                        else:
                            word_widths.append(None)
                else:
                    # Use absolute widths directly
                    for col_idx, html_width in enumerate(html_widths):
                        if html_width is not None:
                            # Convert px to pt
                            word_width_pt = html_width * 0.75
                            word_width_pt = max(word_width_pt, 30)
                            word_widths.append(word_width_pt)
                            assigned_width += word_width_pt
                        else:
                            word_widths.append(None)
            else:
                # Strategy 2: All percentages or mixed percentages
                # Calculate proportional widths based on total
                total_html_width = sum(w for w in html_widths if w is not None)

                for col_idx, html_width in enumerate(html_widths):
                    if html_width is not None:
                        proportion = html_width / total_html_width
                        word_width_pt = table_width_pt * proportion
                        word_width_pt = max(word_width_pt, 30)
                        word_widths.append(word_width_pt)
                        assigned_width += word_width_pt
                    else:
                        word_widths.append(None)

            # If total assigned width exceeds table width, scale down proportionally
            if assigned_width > table_width_pt:
                scale_factor = table_width_pt / assigned_width
                for i in range(len(word_widths)):
                    if word_widths[i] is not None:
                        word_widths[i] *= scale_factor
                assigned_width = table_width_pt
                logger.debug(f"Scaled down column widths by factor {scale_factor:.2f} to fit table width")

            # Distribute remaining width to auto columns
            auto_columns = sum(1 for w in word_widths if w is None)
            if auto_columns > 0:
                remaining_width = table_width_pt - assigned_width
                auto_width = max(remaining_width / auto_columns, 30)  # Min 30pt per column
                for i in range(len(word_widths)):
                    if word_widths[i] is None:
                        word_widths[i] = auto_width
                logger.debug(f"Distributed {remaining_width:.1f}pt to {auto_columns} auto columns")

            # Third pass: apply calculated widths to Word table columns
            for col_idx, width_pt in enumerate(word_widths):
                if col_idx >= len(table.columns):
                    break

                if width_pt:
                    # Convert pt to inches and apply
                    width_inches = width_pt / 72
                    table.columns[col_idx].width = Inches(width_inches)

                    # Calculate percentage for logging
                    percent = (width_pt / table_width_pt) * 100
                    logger.debug(f"Column {col_idx}: {width_pt:.1f}pt ({percent:.1f}% of table width)")

            # CRITICAL: Also set cell widths for each cell to ensure proper rendering
            # Setting only table.columns[].width updates tblGrid but not individual cell tcW
            # Word needs both for correct proportional column display
            self._apply_cell_widths(table, word_widths)

            logger.info(f"Applied proportional column widths to table: total width {table_width_pt:.1f}pt ({table_width_pt/72:.2f} inches)")

        except Exception as e:
            logger.warning(f"Error applying column widths: {e}")

    def _apply_cell_widths(self, table, col_widths_pt: List[float]):
        """
        Apply column widths to individual cells.

        Word tables need both tblGrid (column definitions) AND tcW (cell widths) set
        to properly display proportional columns. This method sets the tcW for each cell.

        Args:
            table: python-docx Table object
            col_widths_pt: List of column widths in pt
        """
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn

        try:
            for row_idx, row in enumerate(table.rows):
                col_idx = 0

                for cell_idx, cell in enumerate(row.cells):
                    if col_idx >= len(col_widths_pt):
                        break

                    width_pt = col_widths_pt[col_idx]
                    if width_pt:
                        # Convert pt to DXA (twentieths of a point)
                        width_dxa = int(width_pt * 20)

                        # Get or create cell properties
                        tc = cell._element
                        tcPr = tc.get_or_add_tcPr()

                        # Remove existing width if present
                        for tcW in tcPr.findall(qn('w:tcW')):
                            tcPr.remove(tcW)

                        # Set new cell width
                        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
                        tcPr.append(tcW)

                        logger.debug(f"Set cell [row {row_idx}, col {cell_idx}] width to {width_pt:.1f}pt ({width_dxa} DXA)")

                    col_idx += 1

            logger.debug(f"Applied cell widths to {len(table.rows)} rows")

        except Exception as e:
            logger.warning(f"Error applying cell widths: {e}")

    def _extract_column_widths(self, table_node: DOMNode) -> List[Optional[str]]:
        """
        Extract column widths from <col> tags or first row cells.

        Prioritizes <col> tags within <colgroup> as they provide the most
        accurate width information, especially for Element UI tables.

        Args:
            table_node: Table DOM node

        Returns:
            List of width values (can be None for auto-width columns)
        """
        widths = []

        # First, try to find <col> or <colgroup> tags
        for child in table_node.children:
            if child.tag == 'colgroup':
                for col in child.children:
                    if col.tag == 'col':
                        # Try multiple sources for width
                        width = (
                            col.get_attribute('width') or
                            col.computed_styles.get('width') or
                            None
                        )

                        # Convert numeric width attribute to string with 'px'
                        if width and isinstance(width, (int, float)):
                            width = f"{width}px"

                        widths.append(width)
            elif child.tag == 'col':
                width = (
                    child.get_attribute('width') or
                    child.computed_styles.get('width') or
                    None
                )

                # Convert numeric width attribute to string with 'px'
                if width and isinstance(width, (int, float)):
                    width = f"{width}px"

                widths.append(width)

        # If we found column definitions, log them and return
        if widths:
            logger.debug(f"Found {len(widths)} column definitions from <col> tags: {widths}")
            return widths

        # Otherwise, extract from first row cells
        rows = self._extract_rows(table_node)
        if not rows:
            return []

        first_row = rows[0]
        for cell in first_row.children:
            if cell.tag in ('td', 'th'):
                # Only use explicitly set widths (from width attribute or inline style)
                # Don't use computed_styles.get('width') as it may contain browser defaults
                width_attr = cell.get_attribute('width')

                # Convert numeric width attribute to string with 'px'
                if width_attr and isinstance(width_attr, (int, float)):
                    width_attr = f"{width_attr}px"

                # Check inline style for width
                style_attr = cell.get_attribute('style')
                width_from_style = None
                if style_attr and isinstance(style_attr, str) and 'width' in style_attr:
                    # Parse inline style to extract width
                    import re
                    match = re.search(r'width\s*:\s*([^;]+)', style_attr)
                    if match:
                        width_from_style = match.group(1).strip()

                width = width_attr or width_from_style
                widths.append(width)

        if widths:
            logger.debug(f"Extracted {len(widths)} column widths from first row cells: {widths}")

        return widths

    def _apply_row_height(self, word_row, row_node: DOMNode):
        """
        Apply row height from CSS styles.

        Args:
            word_row: python-docx Row object
            row_node: Row DOM node
        """
        from docx.shared import Pt
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        # Get height from CSS
        height_str = row_node.computed_styles.get('height')
        if not height_str:
            # Also check cells for height
            for cell in row_node.children:
                if cell.tag in ('td', 'th'):
                    height_str = cell.computed_styles.get('height')
                    if height_str:
                        break

        if height_str and height_str != 'auto':
            try:
                # Parse height value
                from html2word.utils.units import UnitConverter
                height_pt = UnitConverter.to_pt(height_str)

                if height_pt > 0:
                    # Determine the height rule
                    # Check if this is a vertical centering case (height == line-height)
                    height_rule = "atLeast"  # Default: allow expansion

                    # Check cells for line-height matching height (vertical centering pattern)
                    for cell in row_node.children:
                        if cell.tag in ('td', 'th'):
                            # Check direct cell content or first child
                            check_nodes = [cell]
                            if cell.children:
                                check_nodes.extend(cell.children[:3])  # Check first few children

                            for node in check_nodes:
                                line_height_str = node.computed_styles.get('line-height')
                                node_height_str = node.computed_styles.get('height')

                                if line_height_str and node_height_str:
                                    try:
                                        line_height_pt = UnitConverter.to_pt(line_height_str)
                                        node_height_pt = UnitConverter.to_pt(node_height_str)

                                        # If height and line-height are equal (within tolerance),
                                        # this is vertical centering - use exact height
                                        if abs(node_height_pt - line_height_pt) < 2.0:
                                            height_rule = "exact"
                                            logger.debug(f"Detected vertical centering pattern (height={node_height_pt}pt ≈ line-height={line_height_pt}pt), using exact rule")
                                            break
                                    except:
                                        pass

                            if height_rule == "exact":
                                break

                    trPr = word_row._element.get_or_add_trPr()

                    # Remove existing height if present
                    for trHeight in trPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trHeight'):
                        trPr.remove(trHeight)

                    # Add new height with determined rule
                    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="{height_rule}"/>')
                    trPr.append(trHeight)

                    logger.debug(f"Applied row height: {height_pt}pt ({height_rule})")
            except Exception as e:
                logger.warning(f"Failed to apply row height: {e}")

    def _get_font_size_pt(self, font_size) -> float:
        """
        Get font size in pt for spacing calculations.

        Args:
            font_size: Font size value (can be str like '14px' or numeric)

        Returns:
            Font size in pt (float)
        """
        if isinstance(font_size, (int, float)):
            return float(font_size)

        try:
            from html2word.utils.units import UnitConverter
            return UnitConverter.to_pt(str(font_size))
        except:
            return 10.5  # Default: 14px = 10.5pt (common for Chinese documents)

    def _set_row_cant_split(self, row):
        """
        Set row property to prevent splitting across pages.

        Args:
            row: python-docx Row object
        """
        try:
            from docx.oxml.shared import OxmlElement
            from docx.oxml.ns import qn

            tr = row._tr
            trPr = tr.get_or_add_trPr()
            
            # Check if cantSplit element already exists
            cantSplit = trPr.find(qn('w:cantSplit'))
            if cantSplit is None:
                cantSplit = OxmlElement('w:cantSplit')
                trPr.append(cantSplit)
                
            logger.debug("Set row cantSplit property")
        except Exception as e:
            logger.warning(f"Failed to set row cantSplit property: {e}")

    def _set_row_as_header(self, row):
        """
        Set row as header row (repeats at the top of each page).

        Args:
            row: python-docx Row object
        """
        try:
            from docx.oxml.shared import OxmlElement
            from docx.oxml.ns import qn

            tr = row._tr
            trPr = tr.get_or_add_trPr()
            
            # Check if tblHeader element already exists
            tblHeader = trPr.find(qn('w:tblHeader'))
            if tblHeader is None:
                tblHeader = OxmlElement('w:tblHeader')
                tblHeader.set(qn('w:val'), 'true')
                trPr.append(tblHeader)
            
            logger.debug("Set row as header")
        except Exception as e:
            logger.warning(f"Failed to set row as header: {e}")

    def _set_row_keep_with_next(self, row):
        """
        Set keep_with_next=True for all paragraphs in the row.
        This ensures the row sticks to the next row (preventing pagination split).

        Args:
            row: python-docx Row object
        """
        try:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
            
            logger.debug("Set row keep_with_next property")
        except Exception as e:
            logger.warning(f"Failed to set row keep_with_next: {e}")
