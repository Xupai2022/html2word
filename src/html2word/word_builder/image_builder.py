"""
Image builder for Word documents.

Inserts images into Word documents with proper sizing.
"""

import logging
import io
import re
from typing import Optional
from docx.shared import Inches

from html2word.parser.dom_tree import DOMNode
from html2word.utils.image_utils import ImageProcessor

logger = logging.getLogger(__name__)


class ImageBuilder:
    """Builds image insertions for Word documents."""

    def __init__(self, document, base_path: Optional[str] = None):
        """
        Initialize image builder.

        Args:
            document: python-docx Document object
            base_path: Base path for resolving relative image paths
        """
        self.document = document
        self.image_processor = ImageProcessor(base_path)

    def build_image(self, img_node: DOMNode) -> Optional[object]:
        """
        Insert image from img node.

        Args:
            img_node: Image DOM node

        Returns:
            python-docx InlineShape object or None
        """
        if img_node.tag != 'img':
            return None

        # Get image source (handle srcset if present)
        src = self._get_best_image_src(img_node)
        if not src:
            logger.warning("Image node has no src attribute")
            return None

        # Get CSS dimensions
        css_width = img_node.computed_styles.get('width')
        css_height = img_node.computed_styles.get('height')
        max_width = img_node.computed_styles.get('max-width')
        max_height = img_node.computed_styles.get('max-height')
        
        # If img has very small dimensions (like 14px from font-size inheritance)
        # but parent container has explicit dimensions, use parent's dimensions
        if css_width and css_height:
            try:
                width_val = float(css_width.replace('px', ''))
                height_val = float(css_height.replace('px', ''))
                
                # If dimensions are suspiciously small (< 20px) and parent has explicit size, use parent's
                if (width_val < 20 or height_val < 20) and img_node.parent:
                    parent_width = img_node.parent.computed_styles.get('width')
                    parent_height = img_node.parent.computed_styles.get('height')
                    
                    if parent_width and parent_height:
                        try:
                            parent_width_val = float(parent_width.replace('px', ''))
                            parent_height_val = float(parent_height.replace('px', ''))
                            
                            # Use parent dimensions if they're reasonable (> 50px)
                            if parent_width_val > 50 and parent_height_val > 50:
                                logger.debug(f"Using parent dimensions ({parent_width}x{parent_height}) instead of img's tiny dimensions ({css_width}x{css_height})")
                                css_width = parent_width
                                css_height = parent_height
                        except:
                            pass
            except:
                pass

        # Get CSS transform and filter for pre-processing
        transform = img_node.computed_styles.get('transform')
        filter_value = img_node.computed_styles.get('filter')

        # Process image (with optional filters/transforms)
        result = self.image_processor.process_image(
            src,
            transform=transform,
            filter_css=filter_value
        )
        if result is None:
            logger.error(f"Failed to process image: {src}")
            # Add placeholder text for failed images
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(f"[Image failed to load: {src[:100]}{'...' if len(src) > 100 else ''}]")
            run.font.color.rgb = None  # Use theme color
            run.italic = True
            return paragraph

        image_stream, image_size = result

        # Calculate display size in inches
        # Priority: width/height > max-width/max-height > defaults
        # Maximum usable width: 6.5 inches (468pt) for standard Word page with margins
        max_width_inches = 6.5  # Adjusted to match actual page width
        max_height_inches = 9.0  # Default max height

        # Only use max-width/max-height if width/height are not specified
        if not css_width and max_width:
            from html2word.utils.units import UnitConverter
            max_width_pt = UnitConverter.to_pt(max_width)
            max_width_inches = max_width_pt / 72

        if not css_height and max_height:
            from html2word.utils.units import UnitConverter
            max_height_pt = UnitConverter.to_pt(max_height)
            max_height_inches = max_height_pt / 72

        display_width, display_height = self.image_processor.calculate_display_size(
            image_size,
            css_width,
            css_height,
            max_width_inches,
            max_height_inches
        )

        # Insert image
        try:
            # Create paragraph for image
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()

            # Add picture
            picture = run.add_picture(
                image_stream,
                width=Inches(display_width),
                height=Inches(display_height)
            )

            # Apply text alignment if specified (check both img and parent container)
            text_align = img_node.computed_styles.get('text-align')

            # Check for margin: auto (CSS horizontal centering) on img or parent
            is_margin_centered = self._check_margin_centering(img_node)

            # If not already aligned by text-align or margin, check parent
            if not text_align and not is_margin_centered and img_node.parent:
                # Check parent's text-align
                text_align = img_node.parent.computed_styles.get('text-align')

            # Apply alignment based on detected styles
            if is_margin_centered or text_align == 'center':
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.debug(f"Image centered via {'margin:auto' if is_margin_centered else 'text-align:center'}")
            elif text_align == 'right':
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif text_align == 'left':
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # TODO: Handle float (requires more complex wrapping setup)
            # float_val = img_node.computed_styles.get('float')
            # if float_val in ('left', 'right'):
            #     # Would need to set wrapping style on picture

            logger.debug(f"Inserted image: {src} ({display_width:.2f}x{display_height:.2f} in)")
            return picture

        except Exception as e:
            logger.error(f"Error inserting image {src}: {e}")
            return None

    def _check_margin_centering(self, img_node: DOMNode) -> bool:
        """
        Check if image or its parent has margin: auto centering.

        Args:
            img_node: Image DOM node

        Returns:
            True if centered via margin: auto
        """
        # Check img node itself
        margin = img_node.computed_styles.get('margin', '')
        margin_left = img_node.computed_styles.get('margin-left', '')
        margin_right = img_node.computed_styles.get('margin-right', '')

        # Check for margin shorthand with auto
        if margin:
            parts = str(margin).split()
            # Pattern: "16px auto 12px" or "0 auto" etc
            # In shorthand: margin: top right bottom left or margin: top horizontal bottom
            if 'auto' in parts:
                # If auto appears in margin shorthand, it's for horizontal centering
                if len(parts) == 2 and parts[1] == 'auto':  # "16px auto"
                    return True
                elif len(parts) == 3 and parts[1] == 'auto':  # "16px auto 12px"
                    return True
                elif len(parts) == 4 and (parts[1] == 'auto' or parts[3] == 'auto'):  # "16px auto 12px auto"
                    return True if parts[1] == 'auto' and parts[3] == 'auto' else False

        # Check explicit margin-left/right
        if (margin_left == 'auto' and margin_right == 'auto'):
            return True

        # Check parent node (often the container has margin: auto)
        if img_node.parent:
            parent_margin = img_node.parent.computed_styles.get('margin', '')
            parent_margin_left = img_node.parent.computed_styles.get('margin-left', '')
            parent_margin_right = img_node.parent.computed_styles.get('margin-right', '')

            if parent_margin:
                parts = str(parent_margin).split()
                if 'auto' in parts:
                    if len(parts) == 2 and parts[1] == 'auto':
                        return True
                    elif len(parts) == 3 and parts[1] == 'auto':
                        return True
                    elif len(parts) == 4 and parts[1] == 'auto' and parts[3] == 'auto':
                        return True

            if (parent_margin_left == 'auto' and parent_margin_right == 'auto'):
                return True

        return False

    def _get_best_image_src(self, img_node: DOMNode) -> Optional[str]:
        """
        Get the best image source from src or srcset attribute.

        Args:
            img_node: Image DOM node

        Returns:
            Image URL string or None
        """
        # Check for srcset first
        srcset = img_node.get_attribute('srcset')
        if srcset:
            # Parse srcset: "url1 1x, url2 2x, url3 3x" or "url1 100w, url2 200w"
            candidates = []
            for item in srcset.split(','):
                item = item.strip()
                parts = item.rsplit(None, 1)  # Split from right, max 1 split

                if len(parts) == 2:
                    url, descriptor = parts
                    # Extract multiplier or width
                    if descriptor.endswith('x'):
                        try:
                            multiplier = float(descriptor[:-1])
                            candidates.append((url, multiplier, 'density'))
                        except:
                            pass
                    elif descriptor.endswith('w'):
                        try:
                            width = int(descriptor[:-1])
                            candidates.append((url, width, 'width'))
                        except:
                            pass
                elif len(parts) == 1:
                    # No descriptor, assume 1x
                    candidates.append((parts[0], 1.0, 'density'))

            # Select best candidate
            if candidates:
                # Prefer 2x or 3x for better quality in print
                density_candidates = [c for c in candidates if c[2] == 'density']
                if density_candidates:
                    # Sort by multiplier descending, pick highest (up to 3x)
                    density_candidates.sort(key=lambda x: x[1], reverse=True)
                    for url, mult, _ in density_candidates:
                        if mult <= 3.0:  # Don't go too high
                            logger.debug(f"Selected srcset image: {url} ({mult}x)")
                            return url
                    return density_candidates[0][0]  # Return highest if all > 3x

                # Fall back to width-based, select largest
                width_candidates = [c for c in candidates if c[2] == 'width']
                if width_candidates:
                    width_candidates.sort(key=lambda x: x[1], reverse=True)
                    logger.debug(f"Selected srcset image: {width_candidates[0][0]} ({width_candidates[0][1]}w)")
                    return width_candidates[0][0]

        # Fall back to src attribute
        return img_node.get_attribute('src')

    def build_svg(self, svg_node: DOMNode, width: str, height: str, in_table_cell: bool = False) -> Optional[object]:
        """
        Convert inline SVG to image and insert.

        Tries multiple conversion methods in order:
        1. cairosvg (best quality)
        2. svglib + reportlab (good compatibility)
        3. PIL placeholder (fallback)

        Args:
            svg_node: SVG DOM node
            width: SVG width (from attribute or CSS)
            height: SVG height (from attribute or CSS)
            in_table_cell: Whether the SVG is being inserted in a table cell

        Returns:
            python-docx InlineShape object or None
        """
        try:
            # Check if this is an icon SVG with use element referencing missing symbol
            use_element = self._find_use_element(svg_node)
            if use_element and self._is_missing_symbol(use_element, svg_node):
                logger.debug("SVG icon references missing symbol, creating fallback")
                return self._create_icon_fallback(use_element, width, height)

            # Get SVG content: use preprocessed content if available (ensures cache consistency)
            if hasattr(svg_node, '_preprocessed_svg_content'):
                svg_content = svg_node._preprocessed_svg_content
                logger.debug("Using preprocessed SVG content")
            else:
                # Serialize the node
                svg_content = self.serialize_svg_node(svg_node, width, height)

            if not svg_content:
                logger.warning("Empty SVG content")
                return None

            # Parse dimensions
            width_val = self._parse_dimension(width)
            height_val = self._parse_dimension(height)

            # Skip SVGs with zero or very small dimensions
            if width_val < 1 or height_val < 1:
                logger.debug(f"Skipping SVG with zero or very small dimensions: {width_val}x{height_val}")
                return None

            # Try method 1: BrowserSVGConverter (best for complex charts)
            png_data = self._convert_svg_with_browser(svg_content, width_val, height_val)
            if png_data:
                return self._insert_svg_as_image(png_data, width_val, height_val, "Browser", in_table_cell)

            # Try method 2: cairosvg (high quality)
            png_data = self._convert_svg_with_cairosvg(svg_content)
            if png_data:
                return self._insert_svg_as_image(png_data, width_val, height_val, "cairosvg", in_table_cell)

            # Try method 3: svglib + reportlab
            png_data = self._convert_svg_with_svglib(svg_content, width_val, height_val)
            if png_data:
                return self._insert_svg_as_image(png_data, width_val, height_val, "svglib", in_table_cell)

            # Fallback: cairosvg and svglib failed
            logger.warning("SVG conversion failed: Both cairosvg and svglib not available or failed")
            return self._create_svg_fallback_placeholder(svg_node, width, height)

        except Exception as e:
            logger.error(f"Error processing SVG: {e}")
            return None

    def _find_use_element(self, svg_node: DOMNode) -> Optional[DOMNode]:
        """Find the first use element in the SVG node."""
        for child in svg_node.children:
            if child.is_element and child.tag == 'use':
                return child
        return None

    def _is_missing_symbol(self, use_element: DOMNode, svg_node: DOMNode) -> bool:
        """
        Check if the use element references a symbol that doesn't exist in the document.

        Args:
            use_element: The use element to check
            svg_node: The parent SVG node

        Returns:
            True if the referenced symbol is missing
        """
        href = use_element.get_attribute('xlink:href') or use_element.get_attribute('href')
        if not href or not href.startswith('#'):
            return False

        symbol_id = href[1:]

        # Check if symbol exists in the document by looking for it in the root node
        root = svg_node
        while root.parent:
            root = root.parent

        # Search for the symbol in the entire document
        def find_symbol(node: DOMNode, symbol_id: str) -> bool:
            if node.is_element and node.tag == 'symbol' and node.get_attribute('id') == symbol_id:
                return True
            for child in node.children:
                if find_symbol(child, symbol_id):
                    return True
            return False

        return not find_symbol(root, symbol_id)

    def _create_icon_fallback(self, use_element: DOMNode, width: str, height: str) -> Optional[object]:
        """
        Create a fallback icon for missing SVG symbols.

        Args:
            use_element: The use element that references a missing symbol
            width: Width of the icon
            height: Height of the icon

        Returns:
            python-docx InlineShape object or None
        """
        try:
            from PIL import Image, ImageDraw, ImageFont
            import io
            from docx.shared import Inches

            # Parse dimensions
            width_val = int(self._parse_dimension(width))
            height_val = int(self._parse_dimension(height))

            # Ensure minimum size for icon
            width_val = max(width_val, 16)
            height_val = max(height_val, 16)

            # Get href to determine icon type
            href = use_element.get_attribute('xlink:href') or use_element.get_attribute('href')
            icon_type = 'info'  # Default

            if href:
                if 'tishi' in href or 'info' in href:
                    icon_type = 'info'
                elif 'warning' in href or 'jinggao' in href:
                    icon_type = 'warning'
                elif 'error' in href or 'cuowu' in href:
                    icon_type = 'error'
                elif 'success' in href or 'zhengque' in href:
                    icon_type = 'success'

            # Create icon based on type
            img = Image.new('RGBA', (width_val, height_val), (0, 0, 0, 0))
            draw = ImageDraw.Draw(img)

            # Colors for different icon types
            colors = {
                'info': (86, 125, 245, 255),      # Blue from the original style
                'warning': (255, 193, 7, 255),    # Amber
                'error': (244, 67, 54, 255),      # Red
                'success': (76, 175, 80, 255)     # Green
            }

            color = colors.get(icon_type, colors['info'])

            # Draw appropriate icon shape
            if icon_type == 'info':
                # Draw a circle with 'i' in it - outlined style (white background, blue border, blue text)
                margin = 1
                border_width = max(2, width_val // 8)
                # Draw white filled circle with blue outline
                draw.ellipse([(margin, margin), (width_val - margin - 1, height_val - margin - 1)],
                           fill=(255, 255, 255, 255), outline=color, width=border_width)

                # Draw 'i' character in blue
                try:
                    # Try to use a font
                    font_size = int(width_val * 0.6)
                    font = ImageFont.truetype("arial.ttf", font_size)
                    bbox = draw.textbbox((0, 0), "i", font=font)
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1]
                    x = (width_val - text_width) // 2
                    y = (height_val - text_height) // 2 - font_size // 8  # Slight adjustment
                    draw.text((x, y), "i", fill=color, font=font)
                except:
                    # Fallback: draw a simple line for the dot and stem in blue
                    dot_size = max(1, width_val // 12)
                    stem_width = max(1, width_val // 16)
                    # Dot
                    draw.ellipse([(width_val//2 - dot_size, height_val//4 - dot_size),
                                (width_val//2 + dot_size, height_val//4 + dot_size)],
                               fill=color)
                    # Stem
                    draw.rectangle([(width_val//2 - stem_width, height_val//4 + dot_size*2),
                                  (width_val//2 + stem_width, height_val*3//4)],
                                 fill=color)

            # Save to bytes
            img_stream = io.BytesIO()
            img.save(img_stream, format='PNG')
            img_stream.seek(0)

            # Insert image
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            picture = run.add_picture(
                img_stream,
                width=Inches(width_val / 72),
                height=Inches(height_val / 72)
            )

            logger.info(f"Created {icon_type} icon fallback for missing SVG symbol")
            return picture

        except Exception as e:
            logger.error(f"Failed to create icon fallback: {e}")
            return None

    def _convert_svg_with_browser(self, svg_content: str, width_pt: float, height_pt: float) -> Optional[bytes]:
        """
        Convert SVG to PNG using browser rendering (best quality for complex charts).

        Args:
            svg_content: SVG XML string
            width_pt: Target width in points
            height_pt: Target height in points

        Returns:
            PNG data as bytes or None
        """
        try:
            from html2word.utils.browser_svg_converter import get_browser_converter

            converter = get_browser_converter()

            # Convert points to pixels (assuming 96 DPI)
            width_px = int(width_pt * 96 / 72)
            height_px = int(height_pt * 96 / 72)

            png_data = converter.convert(svg_content, width_px, height_px)

            if png_data:
                logger.debug("SVG converted using Browser")
                return png_data
            else:
                logger.debug("Browser conversion returned None")
                return None

        except ImportError:
            logger.debug("BrowserSVGConverter not available")
            return None
        except Exception as e:
            logger.warning(f"Browser conversion failed: {e}")
            return None

    def _convert_svg_with_cairosvg(self, svg_content: str) -> Optional[bytes]:
        """
        Convert SVG to PNG using cairosvg.

        Args:
            svg_content: SVG XML string

        Returns:
            PNG data as bytes or None
        """
        try:
            import cairosvg
            png_data = cairosvg.svg2png(bytestring=svg_content.encode('utf-8'))
            logger.debug("SVG converted using cairosvg")
            return png_data
        except ImportError:
            logger.debug("cairosvg not available")
            return None
        except Exception as e:
            logger.warning(f"cairosvg conversion failed: {e}")
            return None

    def _convert_svg_with_svglib(self, svg_content: str, width_pt: float, height_pt: float) -> Optional[bytes]:
        """
        Convert SVG to PNG using svglib + reportlab.

        Args:
            svg_content: SVG XML string
            width_pt: Target width in points
            height_pt: Target height in points

        Returns:
            PNG data as bytes or None
        """
        try:
            from svglib.svglib import svg2rlg
            from reportlab.graphics import renderPM
            from io import StringIO

            # Parse SVG
            svg_stream = StringIO(svg_content)
            drawing = svg2rlg(svg_stream)

            if not drawing:
                return None

            # Set dimensions
            drawing.width = width_pt
            drawing.height = height_pt
            drawing.scale(width_pt / drawing.width, height_pt / drawing.height)

            # Render to PNG
            png_data = renderPM.drawToString(drawing, fmt='PNG')
            logger.debug("SVG converted using svglib")
            return png_data

        except ImportError:
            logger.debug("svglib/reportlab not available")
            return None
        except Exception as e:
            logger.warning(f"svglib conversion failed: {e}")
            return None

    def _insert_svg_as_image(self, png_data: bytes, width_pt: float, height_pt: float, method: str, in_table_cell: bool = False) -> Optional[object]:
        """
        Insert PNG data as image in document with automatic scaling to fit page width.

        Args:
            png_data: PNG image data
            width_pt: Width in points
            height_pt: Height in points
            method: Conversion method used
            in_table_cell: Whether the image is being inserted in a table cell

        Returns:
            python-docx InlineShape object or None
        """
        try:
            image_stream = io.BytesIO(png_data)

            # Maximum usable width depends on context
            # In table cells, we need to account for cell padding (typically 5-10pt on each side)
            # Use a more conservative limit to ensure images don't get cut off
            if in_table_cell:
                # Reduce max width by 36pt (18pt padding on each side) to ensure image fits
                max_width_pt = 432  # 468 - 36 = 432pt (6 inches)
                logger.debug(f"Image in table cell, using reduced max width: {max_width_pt}pt")
            else:
                # Standard page width (6.5 inches = 468pt for standard page with margins)
                max_width_pt = 468

            # Scale down if image exceeds max width
            if width_pt > max_width_pt:
                scale_factor = max_width_pt / width_pt
                original_width = width_pt
                original_height = height_pt
                width_pt = max_width_pt
                height_pt = height_pt * scale_factor
                logger.info(f"Scaling SVG from {original_width:.1f}x{original_height:.1f}pt to {width_pt:.1f}x{height_pt:.1f}pt to fit {'cell' if in_table_cell else 'page'} width")

            # Insert image
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            picture = run.add_picture(
                image_stream,
                width=Inches(width_pt / 72),
                height=Inches(height_pt / 72)
            )

            logger.info(f"Inserted SVG as image using {method} ({width_pt:.1f}x{height_pt:.1f}pt)")
            return picture

        except Exception as e:
            logger.error(f"Failed to insert SVG image: {e}")
            return None

    def serialize_svg_node(self, svg_node: DOMNode, width: str, height: str) -> str:
        """
        Serialize SVG DOM node to SVG string.

        Args:
            svg_node: SVG DOM node
            width: SVG width
            height: SVG height

        Returns:
            SVG XML string
        """
        # Build SVG opening tag with attributes
        attrs = []
        attrs.append(f'width="{width}"')
        attrs.append(f'height="{height}"')
        attrs.append('xmlns="http://www.w3.org/2000/svg"')

        # Add viewBox if present
        viewbox = svg_node.get_attribute('viewBox')
        if viewbox:
            attrs.append(f'viewBox="{viewbox}"')

        svg_str = f"<svg {' '.join(attrs)}>"

        # Serialize children
        svg_str += self._serialize_svg_children(svg_node)

        svg_str += "</svg>"
        return svg_str

    def _serialize_svg_children(self, node: DOMNode) -> str:
        """Recursively serialize SVG child elements."""
        result = []

        for child in node.children:
            if child.is_text:
                result.append(child.text or "")
            elif child.is_element:
                # Serialize element
                attrs = []
                for key, value in child.attributes.items():
                    attrs.append(f'{key}="{value}"')

                # Check for computed styles to add as inline style
                if child.computed_styles:
                    style_parts = []
                    for key, value in child.computed_styles.items():
                        style_parts.append(f"{key}:{value}")
                    if style_parts:
                        attrs.append(f'style="{";".join(style_parts)}"')

                attrs_str = f" {' '.join(attrs)}" if attrs else ""

                # Handle self-closing tags
                if not child.children:
                    result.append(f"<{child.tag}{attrs_str} />")
                else:
                    result.append(f"<{child.tag}{attrs_str}>")
                    result.append(self._serialize_svg_children(child))
                    result.append(f"</{child.tag}>")

        return "".join(result)

    def _parse_dimension(self, value: str) -> float:
        """
        Parse dimension string to points.

        Args:
            value: Dimension string (e.g., "100", "100px", "2in")

        Returns:
            Value in points
        """
        if not value:
            return 100.0  # Default

        # Remove whitespace
        value = str(value).strip()

        # If just a number, assume pixels
        try:
            return float(value) * 0.75  # px to pt conversion (96 DPI)
        except:
            pass

        # Parse with unit
        match = re.match(r'([\d.]+)(px|pt|in|cm|mm)?', value)
        if match:
            num = float(match.group(1))
            unit = match.group(2) or 'px'

            # Convert to points
            if unit == 'px':
                return num * 0.75
            elif unit == 'pt':
                return num
            elif unit == 'in':
                return num * 72
            elif unit == 'cm':
                return num * 28.35
            elif unit == 'mm':
                return num * 2.835

        return 100.0  # Default

    def _create_svg_fallback_placeholder(self, svg_node: DOMNode, width: str, height: str) -> Optional[object]:
        """
        Create a fallback placeholder for SVG when cairosvg is not available.

        Args:
            svg_node: SVG DOM node
            width: SVG width
            height: SVG height

        Returns:
            Placeholder paragraph or None
        """
        try:
            from PIL import Image, ImageDraw

            # Parse dimensions
            width_val = int(self._parse_dimension(width))
            height_val = int(self._parse_dimension(height))

            # Skip SVGs with zero or very small dimensions
            if width_val < 1 or height_val < 1:
                logger.debug(f"Skipping fallback placeholder for zero-dimension SVG: {width_val}x{height_val}")
                return None

            # Create a simple placeholder image
            img = Image.new('RGB', (width_val, height_val), color='#F0F0F0')
            draw = ImageDraw.Draw(img)

            # Draw a border
            draw.rectangle([(0, 0), (width_val-1, height_val-1)], outline='#CCCCCC', width=2)

            # Add text indicating it's a chart
            try:
                # Try to add text
                text = "[Chart/SVG]"
                bbox = draw.textbbox((0, 0), text)
                text_width = bbox[2] - bbox[0]
                text_height = bbox[3] - bbox[1]
                x = (width_val - text_width) // 2
                y = (height_val - text_height) // 2
                draw.text((x, y), text, fill='#999999')
            except:
                pass  # Skip text if font not available

            # Save to bytes
            img_stream = io.BytesIO()
            img.save(img_stream, format='PNG')
            img_stream.seek(0)

            # Insert image
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            picture = run.add_picture(
                img_stream,
                width=Inches(width_val / 72),
                height=Inches(height_val / 72)
            )

            logger.debug(f"Inserted SVG fallback placeholder ({width}x{height})")
            return picture

        except Exception as e:
            logger.warning(f"Failed to create SVG fallback placeholder: {e}")
            # Last resort: just add a text note
            paragraph = self.document.add_paragraph(f"[Chart/SVG - {width}x{height}]")
            return None
