"""
Paragraph builder for Word documents.

Builds Word paragraphs from HTML elements with styles.
"""

import logging
from typing import Optional, List
from docx.shared import Pt

from html2word.parser.dom_tree import DOMNode
from html2word.word_builder.style_mapper import StyleMapper
from html2word.style.style_normalizer import StyleNormalizer

logger = logging.getLogger(__name__)


class ParagraphBuilder:
    """Builds Word paragraphs from HTML nodes."""

    def __init__(self, document, document_builder=None):
        """
        Initialize paragraph builder.

        Args:
            document: python-docx Document object
            document_builder: Reference to parent DocumentBuilder (optional, for context tracking)
        """
        self.document = document
        self.document_builder = document_builder
        self.style_mapper = StyleMapper()
        # Track previous element's margin-bottom for true margin collapse
        self.last_margin_bottom = 0.0

    def build_paragraph(self, node: DOMNode) -> Optional[object]:
        """
        Build a Word paragraph from HTML node with proper margin collapse.

        Args:
            node: DOM node (typically p, h1-h6, div, etc.)

        Returns:
            python-docx Paragraph object or None
        """
        if not node.is_element:
            return None

        # Create paragraph
        paragraph = self.document.add_paragraph()

        # Apply paragraph-level styles with margin collapse
        box_model = node.layout_info.get('box_model')

        # Check if we're in a table cell (to avoid redundant white backgrounds)
        in_table_cell = getattr(self.document_builder, 'in_table_cell', False) if self.document_builder else False

        # Convert left-aligned text to justified for better readability
        # Create a copy of computed_styles if we need to modify it
        styles = node.computed_styles
        text_align = styles.get('text-align') if styles else None

        # Convert left-aligned or default (no text-align) to justified
        # This ensures better readability for body text
        # Keep explicit center/right/justify alignments as-is
        if text_align in ('left', 'start', None):
            styles = dict(styles) if styles else {}  # Create a copy
            styles['text-align'] = 'justify'

        self.style_mapper.apply_paragraph_style(
            paragraph,
            styles,
            box_model,
            prev_margin_bottom=self.last_margin_bottom,  # Pass for margin collapse
            in_table_cell=in_table_cell  # Pass table cell context
        )

        # Update last_margin_bottom for next element
        if box_model:
            self.last_margin_bottom = box_model.margin.bottom
        else:
            self.last_margin_bottom = 0.0

        # Process content (text and inline elements)
        self._process_content(node, paragraph)

        # CRITICAL: For div.square-title elements, keep with next content to prevent pagination splits
        # This ensures titles like "Assets Requiring Attention" stay with their tables
        if node.tag == 'div' and 'square-title' in (node.attributes.get('class') or []):
            paragraph.paragraph_format.keep_with_next = True
            logger.debug(f"Set keep_with_next=True for square-title: {node.get_text_content()[:30]}...")

        return paragraph

    def _process_content(self, node: DOMNode, paragraph):
        """
        Process node content and add runs to paragraph.

        Args:
            node: DOM node
            paragraph: python-docx Paragraph object
        """
        for child in node.children:
            if child.is_text:
                # Add text run
                self._add_text_run(child, paragraph)
            elif child.tag == 'img':
                # FIXED: Handle inline images
                self._add_inline_image(child, paragraph)
            elif child.tag == 'svg':
                # FIXED: Handle inline SVG
                self._add_inline_svg(child, paragraph)
            elif child.is_inline:
                # Process inline element
                self._process_inline_element(child, paragraph)
            else:
                # Block element inside paragraph - this shouldn't happen
                # but handle it by adding text content
                text = child.get_text_content()
                if text.strip():
                    run = paragraph.add_run(text)

    def _add_text_run(self, text_node: DOMNode, paragraph):
        """
        Add text run to paragraph.

        Args:
            text_node: Text DOM node
            paragraph: python-docx Paragraph object
        """
        text = text_node.text or ""
        if not text:
            return

        # Check if whitespace should be preserved
        from html2word.layout.inline_layout import InlineLayout
        preserve_whitespace = InlineLayout.should_preserve_whitespace(text_node)

        # Normalize whitespace unless it should be preserved
        if not preserve_whitespace:
            text = InlineLayout.normalize_whitespace(text, preserve=False)

        # Skip if text is only whitespace after normalization
        if not text.strip():
            return

        # Apply text-transform if needed
        if text_node.computed_styles.get('text-transform'):
            text = StyleNormalizer.apply_text_transform(
                text,
                text_node.computed_styles['text-transform']
            )

        run = paragraph.add_run(text)

        # Apply run styles from computed styles
        self.style_mapper.apply_run_style(run, text_node.computed_styles)

    def _process_inline_element(self, node: DOMNode, paragraph):
        """
        Process inline element (span, strong, em, etc.).

        Args:
            node: Inline element node
            paragraph: python-docx Paragraph object
        """
        # Get text content
        for child in node.children:
            if child.is_text:
                text = child.text or ""
                if not text:
                    continue

                # Normalize whitespace
                from html2word.layout.inline_layout import InlineLayout
                preserve = InlineLayout.should_preserve_whitespace(child)
                text = InlineLayout.normalize_whitespace(text, preserve=preserve)

                # Skip pure whitespace
                if not text.strip():
                    continue

                # Apply text-transform if needed
                if node.computed_styles.get('text-transform'):
                    text = StyleNormalizer.apply_text_transform(
                        text,
                        node.computed_styles['text-transform']
                    )

                run = paragraph.add_run(text)
                # Apply styles from parent inline element
                self.style_mapper.apply_run_style(run, node.computed_styles)
            elif child.is_inline:
                # Nested inline element
                self._process_inline_element(child, paragraph)
            else:
                # Unexpected block element
                text = child.get_text_content()
                if text.strip():
                    run = paragraph.add_run(text)

    def build_heading(self, node: DOMNode, level: int = 1) -> Optional[object]:
        """
        Build a Word heading from HTML heading node.

        Args:
            node: Heading node (h1-h6)
            level: Heading level (1-6)

        Returns:
            python-docx Paragraph object
        """
        paragraph = self.build_paragraph(node)

        if paragraph:
            # Apply heading style if available
            try:
                paragraph.style = f'Heading {level}'
            except:
                # Heading style not available, keep as normal paragraph
                # Make it bold and larger
                for run in paragraph.runs:
                    run.font.bold = True
                    if run.font.size:
                        run.font.size = Pt(run.font.size.pt * 1.2)

            # CRITICAL: Always keep headings with next content to prevent pagination splits
            # This ensures headings like "Assets Requiring Attention" stay attached to their tables
            paragraph.paragraph_format.keep_with_next = True
            logger.debug(f"Set keep_with_next=True for heading {level}: {node.get_text_content()[:30]}...")

        return paragraph

    def _add_inline_image(self, img_node: DOMNode, paragraph):
        """
        Add inline image to paragraph run.

        Args:
            img_node: Image DOM node
            paragraph: python-docx Paragraph object
        """
        try:
            # Import image builder utilities
            from html2word.word_builder.image_builder import ImageBuilder
            from html2word.utils.image_utils import ImageProcessor
            from docx.shared import Inches

            # Get image source
            image_builder = ImageBuilder(self.document)
            src = image_builder._get_best_image_src(img_node)
            if not src:
                return

            # Get CSS dimensions
            css_width = img_node.computed_styles.get('width')
            css_height = img_node.computed_styles.get('height')
            transform = img_node.computed_styles.get('transform')
            filter_value = img_node.computed_styles.get('filter')

            # Process image
            image_processor = ImageProcessor()
            result = image_processor.process_image(src, transform=transform, filter_css=filter_value)
            if not result:
                return

            image_stream, image_size = result

            # Calculate size (smaller for inline images)
            max_width_inches = 1.0  # Default max for inline images
            max_height_inches = 0.5

            if css_width:
                from html2word.utils.units import UnitConverter
                width_pt = UnitConverter.to_pt(css_width)
                max_width_inches = width_pt / 72

            if css_height:
                from html2word.utils.units import UnitConverter
                height_pt = UnitConverter.to_pt(css_height)
                max_height_inches = height_pt / 72

            display_width, display_height = image_processor.calculate_display_size(
                image_size, css_width, css_height, max_width_inches, max_height_inches
            )

            # Add inline image to paragraph run
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(display_width), height=Inches(display_height))

            logger.debug(f"Added inline image: {src} ({display_width:.2f}x{display_height:.2f} in)")

        except Exception as e:
            logger.warning(f"Failed to add inline image: {e}")

    def _add_inline_svg(self, svg_node: DOMNode, paragraph):
        """
        Add inline SVG to paragraph run (converted to image).

        Args:
            svg_node: SVG DOM node
            paragraph: python-docx Paragraph object
        """
        try:
            from html2word.word_builder.image_builder import ImageBuilder
            import io
            from docx.shared import Inches

            image_builder = ImageBuilder(self.document)

            # Get dimensions - for icons, ensure square aspect ratio
            width = svg_node.get_attribute('width') or svg_node.computed_styles.get('width', '14')
            height = svg_node.get_attribute('height') or svg_node.computed_styles.get('height', '14')

            # Parse dimensions
            width_pt = image_builder._parse_dimension(width)
            height_pt = image_builder._parse_dimension(height)

            # Check if this is an icon SVG with use element referencing missing symbol
            use_element = image_builder._find_use_element(svg_node)

            # For icons, match size to adjacent text font size for visual consistency
            if use_element:
                # Try to get font size from sibling text elements or parent
                icon_size = self._get_icon_size_from_context(svg_node)
                width_pt = icon_size
                height_pt = icon_size
            if use_element and image_builder._is_missing_symbol(use_element, svg_node):
                # Create inline icon fallback
                png_data = self._create_inline_icon_fallback(use_element, width_pt, height_pt, svg_node)
                if png_data:
                    image_stream = io.BytesIO(png_data)
                    run = paragraph.add_run()
                    run.add_picture(
                        image_stream,
                        width=Inches(width_pt / 72),
                        height=Inches(height_pt / 72)
                    )
                    logger.debug(f"Added inline icon fallback ({width}x{height})")
                return

            # Try to convert SVG to PNG using various methods
            svg_content = image_builder.serialize_svg_node(svg_node, width, height)
            if not svg_content:
                return

            # Try cairosvg first
            png_data = None
            try:
                import cairosvg
                png_data = cairosvg.svg2png(bytestring=svg_content.encode('utf-8'))
            except ImportError:
                logger.debug("cairosvg not available for inline SVG")
            except Exception as e:
                logger.debug(f"cairosvg failed: {e}")

            # Try browser converter if cairosvg failed
            if not png_data:
                try:
                    from html2word.utils.browser_svg_converter import get_browser_converter
                    converter = get_browser_converter()
                    width_px = int(width_pt * 96 / 72)
                    height_px = int(height_pt * 96 / 72)
                    png_data = converter.convert(svg_content, width_px, height_px)
                except Exception as e:
                    logger.debug(f"Browser converter failed: {e}")

            if png_data:
                image_stream = io.BytesIO(png_data)
                run = paragraph.add_run()
                run.add_picture(
                    image_stream,
                    width=Inches(width_pt / 72),
                    height=Inches(height_pt / 72)
                )
                logger.debug(f"Added inline SVG as image ({width}x{height})")
            else:
                logger.warning("Could not convert inline SVG, skipping")

        except Exception as e:
            logger.warning(f"Failed to add inline SVG: {e}")

    def _get_icon_size_from_context(self, svg_node: DOMNode) -> float:
        """
        Get icon size based on adjacent text font size for visual consistency.

        Args:
            svg_node: The SVG icon node

        Returns:
            Icon size in points (matching text font size)
        """
        default_size = 12.0  # Default to 12pt (common text size)

        # Check parent node for font-size
        parent = svg_node.parent
        if parent:
            # Check siblings for font-size
            for sibling in parent.children:
                if sibling.is_element and sibling.tag in ('span', 'p', 'div', 'a'):
                    font_size = sibling.computed_styles.get('font-size')
                    if font_size:
                        from html2word.word_builder.image_builder import ImageBuilder
                        ib = ImageBuilder(self.document)
                        size = ib._parse_dimension(font_size)
                        if size > 0:
                            return size

            # Check parent's font-size
            font_size = parent.computed_styles.get('font-size')
            if font_size:
                from html2word.word_builder.image_builder import ImageBuilder
                ib = ImageBuilder(self.document)
                size = ib._parse_dimension(font_size)
                if size > 0:
                    return size

        # Check SVG's own computed font-size
        font_size = svg_node.computed_styles.get('font-size')
        if font_size:
            from html2word.word_builder.image_builder import ImageBuilder
            ib = ImageBuilder(self.document)
            size = ib._parse_dimension(font_size)
            if size > 0:
                return size

        return default_size

    def _create_inline_icon_fallback(self, use_element: DOMNode, width_pt: float, height_pt: float, svg_node: DOMNode) -> bytes:
        """
        Create a fallback icon PNG for inline SVG with missing symbol.
        Uses browser rendering with icon font to get exact icon appearance.

        Args:
            use_element: The use element referencing a missing symbol
            width_pt: Width in points
            height_pt: Height in points
            svg_node: The parent SVG node (for color info)

        Returns:
            PNG data as bytes or None
        """
        # First try browser rendering with icon font
        png_data = self._render_icon_with_browser(use_element, width_pt, height_pt, svg_node)
        if png_data:
            return png_data

        # Fallback to PIL-based icon
        return self._create_pil_icon_fallback(width_pt, height_pt, svg_node)

    def _render_icon_with_browser(self, _use_element: DOMNode, width_pt: float, height_pt: float, svg_node: DOMNode) -> bytes:
        """
        Render icon using Chrome headless to get exact icon appearance.
        Creates a simple info icon that matches the HTML style.
        """
        try:
            import subprocess
            import tempfile
            import os
            import re
            import shutil

            # Get color from SVG style
            color = 'rgb(86, 125, 245)'  # Default blue
            style = svg_node.get_attribute('style') or ''
            color_match = re.search(r'color:\s*(rgb\([^)]+\))', style)
            if color_match:
                color = color_match.group(1)

            # Convert to pixels
            width_px = max(int(width_pt * 96 / 72), 16)
            height_px = max(int(height_pt * 96 / 72), 16)

            # Create HTML that renders an info icon similar to the original
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        html, body {{
            width: {width_px}px;
            height: {height_px}px;
            overflow: hidden;
            background: transparent;
        }}
        .icon {{
            width: {width_px}px;
            height: {height_px}px;
            background: white;
            border: 2px solid {color};
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            color: {color};
            font-family: Arial, sans-serif;
            font-size: {int(height_px * 0.6)}px;
            font-weight: bold;
            font-style: italic;
        }}
    </style>
</head>
<body>
    <span class="icon">i</span>
</body>
</html>"""

            # Find Chrome
            chrome_paths = [
                r"C:\Program Files\Google\Chrome\Application\chrome.exe",
                r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
                os.path.expandvars(r"%LOCALAPPDATA%\Google\Chrome\Application\chrome.exe"),
            ]

            chrome_exe = None
            for path in chrome_paths:
                if os.path.exists(path):
                    chrome_exe = path
                    break

            if not chrome_exe:
                logger.debug("Chrome not found for icon rendering")
                return None

            # Create temp files and user data dir
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
                f.write(html_content)
                html_file = f.name

            screenshot_file = html_file.replace('.html', '.png')
            user_data_dir = tempfile.mkdtemp(prefix='chrome_icon_')

            try:
                # Run Chrome headless with separate user data dir to avoid conflicts
                cmd = [
                    chrome_exe,
                    '--headless=new',
                    '--disable-gpu',
                    '--no-sandbox',
                    f'--user-data-dir={user_data_dir}',
                    f'--window-size={width_px},{height_px}',
                    '--default-background-color=00000000',
                    f'--screenshot={screenshot_file}',
                    f'file:///{html_file}'
                ]

                subprocess.run(cmd, capture_output=True, timeout=10)

                if os.path.exists(screenshot_file):
                    with open(screenshot_file, 'rb') as f:
                        png_data = f.read()
                    logger.debug(f"Successfully rendered icon with browser ({width_px}x{height_px})")
                    return png_data
                else:
                    logger.debug("Chrome screenshot not created for icon")
                    return None

            finally:
                # Cleanup
                try:
                    os.unlink(html_file)
                    if os.path.exists(screenshot_file):
                        os.unlink(screenshot_file)
                    shutil.rmtree(user_data_dir, ignore_errors=True)
                except:
                    pass

        except Exception as e:
            logger.debug(f"Browser icon rendering failed: {e}")
            return None

    def _create_pil_icon_fallback(self, width_pt: float, height_pt: float, svg_node: DOMNode) -> bytes:
        """
        Create a simple PIL-based info icon as last resort fallback.
        """
        try:
            from PIL import Image, ImageDraw, ImageFont
            import re

            # Convert to pixels
            width_px = max(int(width_pt * 96 / 72), 16)
            height_px = max(int(height_pt * 96 / 72), 16)

            # Get color from SVG style
            color = (86, 125, 245, 255)  # Default blue
            style = svg_node.get_attribute('style') or ''
            if 'color:' in style:
                match = re.search(r'color:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)', style)
                if match:
                    color = (int(match.group(1)), int(match.group(2)), int(match.group(3)), 255)

            # Create transparent image
            img = Image.new('RGBA', (width_px, height_px), (0, 0, 0, 0))
            draw = ImageDraw.Draw(img)

            # Draw info icon (circle with 'i')
            margin = width_px // 8
            draw.ellipse([(margin, margin), (width_px - margin, height_px - margin)],
                       fill=color, outline=color)

            # Draw 'i' character
            try:
                font_size = width_px // 2
                font = ImageFont.truetype("arial.ttf", font_size)
                bbox = draw.textbbox((0, 0), "i", font=font)
                text_width = bbox[2] - bbox[0]
                text_height = bbox[3] - bbox[1]
                x = (width_px - text_width) // 2
                y = (height_px - text_height) // 2 - font_size // 8
                draw.text((x, y), "i", fill=(255, 255, 255, 255), font=font)
            except:
                # Fallback: draw simple 'i' shape
                center_x = width_px // 2
                dot_r = max(1, width_px // 10)
                stem_w = max(1, width_px // 12)
                draw.ellipse([(center_x - dot_r, height_px // 4 - dot_r),
                            (center_x + dot_r, height_px // 4 + dot_r)],
                           fill=(255, 255, 255, 255))
                draw.rectangle([(center_x - stem_w, height_px // 4 + dot_r * 2),
                              (center_x + stem_w, height_px * 3 // 4)],
                             fill=(255, 255, 255, 255))

            # Save to bytes
            import io
            img_stream = io.BytesIO()
            img.save(img_stream, format='PNG')
            img_stream.seek(0)
            return img_stream.getvalue()

        except Exception as e:
            logger.warning(f"Failed to create PIL icon fallback: {e}")
            return None
