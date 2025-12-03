"""
Header and Footer Builder for Word documents.

Handles creation and configuration of headers and footers in Word documents.
"""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from html2word.config.header_footer_config import HeaderFooterConfig

logger = logging.getLogger(__name__)


class HeaderFooterBuilder:
    """Builder class for Word document headers and footers."""

    def __init__(self, document: Document, base_path: str = None):
        """
        Initialize header/footer builder.

        Args:
            document: python-docx Document object
            base_path: Base path for resolving relative paths
        """
        self.document = document
        self.base_path = base_path
        self.config = HeaderFooterConfig()

    def apply_cover_image(self):
        """
        Apply cover image at the beginning of the document.

        This should be called BEFORE building the main content.
        封面图片会插入到文档的最开头位置。
        """
        if not self.config.ENABLE_COVER_IMAGE:
            logger.info("Cover image is disabled in configuration")
            return

        cover_image_path = self.config.get_cover_image_path(self.base_path)

        if not cover_image_path.exists():
            logger.warning(f"Cover image not found: {cover_image_path}")
            return

        try:
            logger.info(f"Adding cover image: {cover_image_path}")

            # Create a paragraph for the cover image
            paragraph = self.document.add_paragraph()

            # Set alignment
            alignment = self.config.COVER_IMAGE_ALIGNMENT
            if alignment == "CENTER":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "RIGHT":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add the image
            run = paragraph.add_run()

            # Determine image dimensions
            width = None
            height = None

            if self.config.COVER_IMAGE_WIDTH is not None:
                width = Inches(self.config.COVER_IMAGE_WIDTH)

            if self.config.COVER_IMAGE_HEIGHT is not None:
                height = Inches(self.config.COVER_IMAGE_HEIGHT)

            # Add picture with specified dimensions
            if width and height:
                run.add_picture(str(cover_image_path), width=width, height=height)
            elif width:
                run.add_picture(str(cover_image_path), width=width)
            elif height:
                run.add_picture(str(cover_image_path), height=height)
            else:
                run.add_picture(str(cover_image_path))

            # Set space after the cover image
            paragraph.paragraph_format.space_after = Pt(self.config.COVER_IMAGE_SPACE_AFTER)

            # Add page break after cover if configured
            if self.config.COVER_ADD_PAGE_BREAK:
                self.document.add_page_break()

            logger.info("Cover image added successfully")

        except Exception as e:
            logger.error(f"Failed to add cover image: {e}", exc_info=True)

    def apply_headers_footers(self):
        """
        Apply headers and footers to the document based on configuration.

        This is the main entry point for adding headers and footers.
        """
        # Check master switch
        if not self.config.ENABLE_HEADER_FOOTER:
            logger.info("Headers and footers are disabled in configuration")
            return

        logger.info("Applying headers and footers to document")

        # Validate configuration first
        if not self.config.validate_config(self.base_path):
            logger.warning("Header/Footer configuration validation failed, some images may be missing")

        # Access document sections
        sections = self.document.sections
        if not sections:
            logger.warning("No sections found in document")
            return

        # Apply to each section
        for section in sections:
            # Configure section settings
            self._configure_section_settings(section)

            # Apply headers if enabled
            if self.config.ENABLE_HEADER:
                self._apply_header(section)

            # Apply footers if enabled
            if self.config.ENABLE_FOOTER:
                self._apply_footer(section)

        logger.info("Headers and footers applied successfully")

    def _configure_section_settings(self, section):
        """
        Configure section settings for headers/footers.

        Args:
            section: Document section object
        """
        # Set whether to use different first page header/footer
        section.different_first_page_header_footer = self.config.DIFFERENT_FIRST_PAGE

        # Set whether to use different odd/even page headers/footers
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        # 重要！实际应用页眉页脚距离设置
        # IMPORTANT! Apply header/footer distance settings
        section.header_distance = Inches(self.config.HEADER_DISTANCE_FROM_TOP)
        section.footer_distance = Inches(self.config.FOOTER_DISTANCE_FROM_BOTTOM)

        # 同时设置页边距（可选）
        # Also set page margins (optional)
        section.top_margin = Inches(self.config.TOP_MARGIN)
        section.bottom_margin = Inches(self.config.BOTTOM_MARGIN)
        section.left_margin = Inches(self.config.LEFT_MARGIN)
        section.right_margin = Inches(self.config.RIGHT_MARGIN)

    def _apply_header(self, section):
        """
        Apply header with left and right images.

        Args:
            section: Document section object
        """
        try:
            # Get the header for this section
            header = section.header

            # Clear existing content (if any)
            for paragraph in header.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)

            # Create a new paragraph for the header
            header_paragraph = header.add_paragraph()

            # Create a table for positioning images
            # We use a table with 1 row and 2 columns to position images left and right
            table = header.add_table(rows=1, cols=2, width=Inches(6.5))

            # Remove table borders
            self._remove_table_borders(table)

            # Get table cells
            left_cell = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]

            # Add left image if enabled
            if self.config.SHOW_HEADER_LEFT_IMAGE:
                left_image_path = self.config.get_header_left_image_path(self.base_path)
                if left_image_path.exists():
                    left_paragraph = left_cell.paragraphs[0]
                    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    left_run = left_paragraph.add_run()

                    # Add image with size constraints
                    try:
                        picture = left_run.add_picture(
                            str(left_image_path),
                            width=Inches(self.config.HEADER_LEFT_IMAGE_MAX_WIDTH)
                        )
                        # Check and adjust height if needed
                        if picture.height > Inches(self.config.HEADER_IMAGE_MAX_HEIGHT):
                            # Scale down to fit height
                            aspect_ratio = picture.width / picture.height
                            picture.height = Inches(self.config.HEADER_IMAGE_MAX_HEIGHT)
                            picture.width = picture.height * aspect_ratio
                        logger.debug(f"Added left header image: {left_image_path}")
                    except Exception as e:
                        logger.error(f"Failed to add left header image: {e}")

            # Add right image if enabled
            if self.config.SHOW_HEADER_RIGHT_IMAGE:
                right_image_path = self.config.get_header_right_image_path(self.base_path)
                if right_image_path.exists():
                    right_paragraph = right_cell.paragraphs[0]
                    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    right_run = right_paragraph.add_run()

                    # Add image with size constraints
                    try:
                        picture = right_run.add_picture(
                            str(right_image_path),
                            width=Inches(self.config.HEADER_RIGHT_IMAGE_MAX_WIDTH)
                        )
                        # Check and adjust height if needed
                        if picture.height > Inches(self.config.HEADER_IMAGE_MAX_HEIGHT):
                            # Scale down to fit height
                            aspect_ratio = picture.width / picture.height
                            picture.height = Inches(self.config.HEADER_IMAGE_MAX_HEIGHT)
                            picture.width = picture.height * aspect_ratio
                        logger.debug(f"Added right header image: {right_image_path}")
                    except Exception as e:
                        logger.error(f"Failed to add right header image: {e}")

            # Set table cell widths for proper alignment
            # Left cell takes most of the space, right cell is narrower
            self._set_column_widths(table, [4.5, 2.0])  # in inches

        except Exception as e:
            logger.error(f"Failed to apply header: {e}", exc_info=True)

    def _apply_footer(self, section):
        """
        Apply footer with left text and right page number.

        Args:
            section: Document section object
        """
        try:
            # Get the footer for this section
            footer = section.footer

            # Clear existing content (if any)
            for paragraph in footer.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)

            # Create a table for positioning text left and page number right
            table = footer.add_table(rows=1, cols=2, width=Inches(6.5))

            # Remove table borders
            self._remove_table_borders(table)

            # Get table cells
            left_cell = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]

            # Add left text (contact information) if enabled
            if self.config.SHOW_FOOTER_LEFT_TEXT:
                left_paragraph = left_cell.paragraphs[0]

                # Set alignment based on config
                if self.config.FOOTER_LEFT_ALIGN == "CENTER":
                    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif self.config.FOOTER_LEFT_ALIGN == "RIGHT":
                    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                left_run = left_paragraph.add_run(self.config.FOOTER_LEFT_TEXT)

                # Apply text formatting
                left_run.font.name = self.config.FOOTER_FONT_NAME
                left_run.font.size = Pt(self.config.FOOTER_FONT_SIZE)
                left_run.font.bold = self.config.FOOTER_TEXT_BOLD
                left_run.font.italic = self.config.FOOTER_TEXT_ITALIC
                # Set text color
                r, g, b = self.config.FOOTER_TEXT_COLOR
                left_run.font.color.rgb = RGBColor(r, g, b)

            # Add page number on the right
            if self.config.SHOW_PAGE_NUMBERS:
                right_paragraph = right_cell.paragraphs[0]
                right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # Get page number color
                page_r, page_g, page_b = self.config.PAGE_NUMBER_COLOR

                # Check format type
                if self.config.PAGE_NUMBER_FORMAT == "X/Y":
                    # Add current page / total pages format
                    self._add_page_number_with_total(right_paragraph)
                else:
                    # Legacy format support (e.g., "Page X")
                    format_run = right_paragraph.add_run(self.config.PAGE_NUMBER_FORMAT)
                    format_run.font.name = self.config.FOOTER_FONT_NAME
                    format_run.font.size = Pt(self.config.PAGE_NUMBER_FONT_SIZE)
                    format_run.font.color.rgb = RGBColor(page_r, page_g, page_b)

                    # Add page number field
                    self._add_page_number_field(right_paragraph)

                # Apply formatting to page number runs
                # Note: The page number field itself needs special handling
                for run in right_paragraph.runs:
                    if not run.text or run.text.isspace():  # Skip empty runs
                        continue
                    run.font.name = self.config.FOOTER_FONT_NAME
                    run.font.size = Pt(self.config.PAGE_NUMBER_FONT_SIZE)
                    run.font.color.rgb = RGBColor(page_r, page_g, page_b)

            # Set table cell widths for proper alignment
            self._set_column_widths(table, [5.0, 1.5])  # in inches

        except Exception as e:
            logger.error(f"Failed to apply footer: {e}", exc_info=True)

    def _remove_table_borders(self, table):
        """
        Remove all borders from a table.

        Args:
            table: python-docx Table object
        """
        try:
            tbl = table._element
            tblPr = tbl.tblPr

            # Create table borders element with no borders
            tblBorders = OxmlElement('w:tblBorders')

            # Define all border positions
            border_positions = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']

            for position in border_positions:
                border = OxmlElement(f'w:{position}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)

            tblPr.append(tblBorders)

        except Exception as e:
            logger.warning(f"Failed to remove table borders: {e}")

    def _set_column_widths(self, table, widths_inches):
        """
        Set column widths for a table.

        Args:
            table: python-docx Table object
            widths_inches: List of column widths in inches
        """
        try:
            for i, width in enumerate(widths_inches):
                for cell in table.columns[i].cells:
                    cell.width = Inches(width)
        except Exception as e:
            logger.warning(f"Failed to set column widths: {e}")

    def _add_page_number_field(self, paragraph):
        """
        Add a page number field to a paragraph.

        Args:
            paragraph: python-docx Paragraph object
        """
        try:
            # This adds a simple page number field
            run = paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._element.append(fldChar1)
            run._element.append(instrText)
            run._element.append(fldChar2)

        except Exception as e:
            logger.warning(f"Failed to add page number field: {e}")
            # Fallback: just add placeholder text
            paragraph.add_run("1")

    def _add_page_number_with_total(self, paragraph):
        """
        Add page number in format "current/total" to a paragraph.

        Args:
            paragraph: python-docx Paragraph object
        """
        try:
            # Get page number color
            page_r, page_g, page_b = self.config.PAGE_NUMBER_COLOR

            # Add current page number field
            run1 = paragraph.add_run()
            run1.font.name = self.config.FOOTER_FONT_NAME
            run1.font.size = Pt(self.config.PAGE_NUMBER_FONT_SIZE)
            run1.font.color.rgb = RGBColor(page_r, page_g, page_b)

            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText1 = OxmlElement('w:instrText')
            instrText1.text = 'PAGE'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run1._element.append(fldChar1)
            run1._element.append(instrText1)
            run1._element.append(fldChar2)

            # Add separator "/"
            separator_run = paragraph.add_run("/")
            separator_run.font.name = self.config.FOOTER_FONT_NAME
            separator_run.font.size = Pt(self.config.PAGE_NUMBER_FONT_SIZE)
            separator_run.font.color.rgb = RGBColor(page_r, page_g, page_b)

            # Add total pages field
            run2 = paragraph.add_run()
            run2.font.name = self.config.FOOTER_FONT_NAME
            run2.font.size = Pt(self.config.PAGE_NUMBER_FONT_SIZE)
            run2.font.color.rgb = RGBColor(page_r, page_g, page_b)

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'begin')

            instrText2 = OxmlElement('w:instrText')
            instrText2.text = 'NUMPAGES'

            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')

            run2._element.append(fldChar3)
            run2._element.append(instrText2)
            run2._element.append(fldChar4)

        except Exception as e:
            logger.warning(f"Failed to add page number with total: {e}")
            # Fallback: just add placeholder text
            paragraph.add_run("1/1")

    def update_config(self, **kwargs):
        """
        Update configuration dynamically.

        This allows programmatic configuration changes without editing the config file.

        Args:
            **kwargs: Configuration parameters to update

        Example:
            builder.update_config(
                FOOTER_LEFT_TEXT="New contact info",
                HEADER_IMAGE_MAX_HEIGHT=1.0
            )
        """
        for key, value in kwargs.items():
            if hasattr(self.config, key):
                setattr(self.config, key, value)
                logger.debug(f"Updated config: {key} = {value}")
            else:
                logger.warning(f"Unknown config parameter: {key}")