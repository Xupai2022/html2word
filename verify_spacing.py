#!/usr/bin/env python
"""Verify spacing between elements in the converted Word document."""

import sys
from pathlib import Path

# Add src to path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from docx import Document
from docx.shared import Pt

def analyze_spacing(docx_path):
    """Analyze spacing in Word document."""
    doc = Document(docx_path)

    print(f"\n{'='*80}")
    print(f"Analyzing: {docx_path}")
    print(f"{'='*80}\n")

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text and len(doc.tables) > 0:
            # Check if this is a spacer paragraph
            fmt = para.paragraph_format
            space_before = fmt.space_before
            space_after = fmt.space_after

            if space_before and space_before.pt > 0:
                print(f"Paragraph {i}: [SPACER - space_before={space_before.pt}pt]")
            elif space_after and space_after.pt > 0:
                print(f"Paragraph {i}: [SPACER - space_after={space_after.pt}pt]")
        elif text:
            # Show first 60 chars of text
            preview = text[:60] + "..." if len(text) > 60 else text
            print(f"Paragraph {i}: {preview}")

    print(f"\n{'-'*80}")
    print(f"Tables found: {len(doc.tables)}")
    print(f"{'-'*80}\n")

    # Analyze table positions and surrounding spacing
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        print(f"  Rows: {len(table.rows)}, Cols: {len(table.columns)}")

        # Get first cell content as identifier
        first_cell = table.rows[0].cells[0]
        cell_text = first_cell.text.strip()[:40]
        print(f"  First cell: {cell_text}")

        # Check for spacing paragraphs before/after this table
        # This is approximate since we're looking at document structure
        print(f"  (Spacing is applied via empty paragraphs with spacing)")

def compare_html_margins(html_path):
    """Extract margin information from HTML for comparison."""
    from html2word.parser.html_parser import HTMLParser
    from html2word.style.box_model import BoxModel

    print(f"\n{'='*80}")
    print(f"HTML Margin Analysis: {html_path}")
    print(f"{'='*80}\n")

    parser = HTMLParser()
    tree = parser.parse(html_path)
    body = parser.get_body_content(tree)

    # Find key elements with margins
    elements_of_interest = [
        'executive-summary',
        'risk-overview',
        'section',
        'metric-grid',
        'table-container'
    ]

    def find_elements_by_class(node, class_name):
        """Recursively find elements with specific class."""
        results = []
        if node.is_element:
            elem_class = node.get_attribute('class', '')
            if class_name in elem_class:
                results.append(node)
            for child in node.children:
                results.extend(find_elements_by_class(child, class_name))
        return results

    for class_name in elements_of_interest:
        elements = find_elements_by_class(body, class_name)
        if elements:
            elem = elements[0]  # First occurrence
            box_model = BoxModel(elem)
            print(f"{class_name}:")
            print(f"  margin-top: {box_model.margin.top}pt")
            print(f"  margin-bottom: {box_model.margin.bottom}pt")
            print(f"  display: {elem.computed_styles.get('display', 'block')}")
            print()

def main():
    html_file = 'security_quarterly_report.html'
    docx_file = 'security_quarterly_report_fixed.docx'

    if Path(html_file).exists():
        compare_html_margins(html_file)

    if Path(docx_file).exists():
        analyze_spacing(docx_file)
    else:
        print(f"Error: {docx_file} not found")
        return 1

    print("\n" + "="*80)
    print("EXPECTED BEHAVIOR:")
    print("="*80)
    print("1. executive-summary (margin-bottom: 30px) should have spacer after it")
    print("2. risk-overview (margin-bottom: 30px) should have spacer after it")
    print("3. section (margin-bottom: 40px) should have spacer after it")
    print("4. metric-grid (margin: 20px 0) should have spacers before AND after")
    print("5. These spacers should create visible gaps in the Word document")
    print("="*80 + "\n")

    return 0

if __name__ == '__main__':
    sys.exit(main())
